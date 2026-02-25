"""
Generate Claude Code Agent Orchestration Guide for Frontera Platform Development.

Best practices for using Claude Code agents, the Agent SDK, and
multi-agent orchestration to build out the Frontera platform.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

NAVY = RGBColor(0x1a, 0x1f, 0x3a)
GOLD = RGBColor(0xfb, 0xbf, 0x24)
CYAN = RGBColor(0x08, 0x91, 0xb2)
WHITE = RGBColor(0xff, 0xff, 0xff)
SLATE = RGBColor(0x47, 0x55, 0x69)


def set_cell_shading(cell, color_hex):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(s)


def H(text, level=1):
    h = doc.add_heading(text, level)
    for r in h.runs:
        r.font.color.rgb = NAVY
    return h


def T(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hc = t.rows[0].cells
    for i, header in enumerate(headers):
        hc[i].text = header
        hc[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hc[i], '1a1f3a')
        hc[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        hc[i].paragraphs[0].runs[0].font.size = Pt(9)
    for rd in rows:
        r = t.add_row()
        for i, ct in enumerate(rd):
            r.cells[i].text = str(ct)
            for p in r.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return t


def B(items):
    for i in items:
        p = doc.add_paragraph(i, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)


def Q(text):
    p = doc.add_paragraph(text, style='Intense Quote')
    for run in p.runs:
        run.font.size = Pt(10)


def P(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p


def para(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)


def code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.name = 'Consolas'
    r.font.color.rgb = SLATE
    return p


def link_ref(label, url):
    p = doc.add_paragraph()
    r = p.add_run('\u2192 ')
    r.font.size = Pt(9)
    r.font.color.rgb = CYAN
    r2 = p.add_run(f'{label}')
    r2.font.size = Pt(9)
    r2.bold = True
    r2.font.color.rgb = CYAN
    r3 = p.add_run(f'\n   {url}')
    r3.font.size = Pt(8)
    r3.font.color.rgb = SLATE


# ============================================================================
# TITLE PAGE
# ============================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Building Frontera with Claude Code Agents', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY

sub = doc.add_paragraph('Agent Architecture, Orchestration & Best Practices')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(18)
sub.runs[0].font.color.rgb = CYAN

doc.add_paragraph()
tag = doc.add_paragraph('A Practical Guide for Using Claude Code, the Agent SDK,')
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.runs[0].font.size = Pt(12)
tag.runs[0].italic = True
tag2 = doc.add_paragraph('and Multi-Agent Orchestration to Build the Frontera Platform')
tag2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag2.runs[0].font.size = Pt(12)
tag2.runs[0].italic = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Date: February 2026\n').font.size = Pt(10)
info.add_run('Sources: Anthropic Official Docs, Agent SDK, Claude Code Docs,\n').font.size = Pt(10)
info.add_run('Community Best Practices, Enterprise Development Patterns\n').font.size = Pt(10)

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
H('Table of Contents', 1)
toc = [
    ('1.', 'Executive Summary: The Agent-First Development Model'),
    ('2.', 'The Claude Code Agent Landscape'),
    ('3.', 'Anthropic\'s 5 Workflow Patterns & When to Use Each'),
    ('4.', 'CLAUDE.md: Your Project\'s Agent Constitution'),
    ('5.', 'Skills, Hooks & Custom Commands'),
    ('6.', 'Multi-Agent Orchestration: Agent Teams & Swarms'),
    ('7.', 'The Claude Agent SDK: Building Custom Agents'),
    ('8.', 'Frontera Build Plan: Agent Strategy by Phase'),
    ('9.', 'Recommended Agent Architecture for Frontera'),
    ('10.', 'Reference Directory'),
]
for n, i in toc:
    p = doc.add_paragraph()
    r = p.add_run(f'{n} ')
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(i)
    r2.font.size = Pt(11)

doc.add_page_break()

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================
H('1. Executive Summary: The Agent-First Development Model', 1)

para(
    'Claude Code has evolved from a single AI coding assistant into a multi-agent development platform. '
    'For a complex platform like Frontera \u2014 with its Next.js frontend, Supabase backend, AI coaching '
    'system, and growing module architecture \u2014 the question is no longer "should I use Claude Code?" '
    'but "how should I orchestrate Claude Code agents to maximise development velocity while maintaining '
    'quality?"'
)

para(
    'This guide provides a practical framework for using Claude Code agents to build the next phases '
    'of Frontera. It covers the three layers of agent-assisted development:'
)

B([
    'Layer 1: Single-Agent Development \u2014 Using Claude Code interactively for feature development, bug fixing, and code review (what you\'re doing now)',
    'Layer 2: Subagent Parallelisation \u2014 Using Claude Code\'s built-in Task tool to parallelise research, exploration, and independent work units',
    'Layer 3: Multi-Agent Teams \u2014 Using Agent Teams (Swarms) to coordinate multiple Claude Code instances working on different aspects of the same feature',
])

H('The Core Principle', 2)
Q('"Success in the LLM space isn\'t about building the most sophisticated system. It\'s about building the right system for your needs." \u2014 Anthropic, Building Effective Agents')

para(
    'Anthropic\'s own guidance is clear: start simple, add complexity only when needed. For Frontera\'s '
    'current stage, the right approach is a blend of enhanced single-agent development (Layer 1) with '
    'strategic subagent parallelisation (Layer 2). Agent Teams (Layer 3) should be reserved for the '
    'largest features that can genuinely benefit from parallel development \u2014 such as Phase 5: '
    'Strategic Activation, which spans frontend, backend, and AI agent changes simultaneously.'
)

H('What This Means for Frontera', 2)
T(
    ['Development Phase', 'Recommended Agent Approach', 'Why'],
    [
        ['PSC Enhancements (Express Mode, micro-synthesis)', 'Layer 1 + Layer 2 (single agent with research subagents)', 'Sequential changes within existing architecture; subagents for research and testing'],
        ['Phase 5: Strategic Activation', 'Layer 2 + Layer 3 (subagents + agent team)', 'Cross-layer feature: new API routes, new UI components, new AI prompts, database changes'],
        ['Expert Perspectives (UC1)', 'Layer 2 (transcript ingestion pipeline)', 'Data processing pipeline; parallelise transcript ingestion across subagents'],
        ['Sparring Partners (UC2)', 'Layer 1 (single agent)', 'Persona definitions and UI updates; contained, sequential work'],
        ['Case Study Engine (UC3)', 'Layer 2 (extraction pipeline)', 'Bulk extraction from 301 transcripts; highly parallelisable'],
        ['Tension Simulator (UC4)', 'Layer 1 + Layer 2', 'UI component + data model; subagents for tension mapping research'],
        ['Leadership Playbook (UC5)', 'Layer 1', 'Single feature with AI prompt engineering; iterative development'],
        ['Integration Layer (Slack, Jira)', 'Layer 3 (agent team)', 'Multiple independent integrations; each agent owns one integration'],
    ]
)

doc.add_page_break()

# ============================================================================
# 2. THE CLAUDE CODE AGENT LANDSCAPE
# ============================================================================
H('2. The Claude Code Agent Landscape', 1)

para(
    'Understanding the different agent types and mechanisms available in Claude Code is essential '
    'for choosing the right approach for each development task.'
)

H('Agent Types Available', 2)
T(
    ['Agent Type', 'Context', 'Trigger', 'Best For'],
    [
        ['CLAUDE.md', 'Main conversation (always loaded)', 'Automatic at session start', 'Project conventions, architecture rules, coding standards'],
        ['Slash Commands', 'Main conversation', 'Manual (/command)', 'Repeatable workflows: /commit, /review-pr, /test'],
        ['Skills (SKILL.md)', 'Main conversation', 'Auto-discovered by Claude', 'Rich, multi-step capabilities with supporting files'],
        ['Subagents (Task tool)', 'Separate context window', 'Delegated by main agent', 'Research, exploration, parallel work units'],
        ['Agent Teams (Swarms)', 'Multiple separate sessions', 'Explicit team launch', 'Large parallel features, cross-layer development'],
    ]
)

H('Subagent Types (Built-in)', 2)
T(
    ['Subagent Type', 'Tools Available', 'Best For', 'Frontera Use Case'],
    [
        ['Bash', 'Bash commands only', 'Git operations, command execution, terminal tasks', 'Running tests, builds, database migrations'],
        ['general-purpose', 'All tools', 'Complex multi-step research, code search, execution', 'Feature research, multi-file analysis, prototype testing'],
        ['Explore', 'Read-only tools (no Edit/Write)', 'Fast codebase exploration, file search, keyword search', 'Understanding existing patterns before modifying code'],
        ['Plan', 'Read-only tools', 'Designing implementation strategies', 'Planning Phase 5 architecture, module design decisions'],
        ['claude-code-guide', 'Read, Glob, Grep, WebFetch, WebSearch', 'Questions about Claude Code, Agent SDK, API usage', 'Learning new Claude Code features, SDK patterns'],
    ]
)

P('Key insight: ', 'Subagents are the workhorse of efficient Claude Code usage. They keep the main '
  'context window clean by performing research, exploration, and heavy analysis in separate context '
  'windows. Only the distilled results come back to the main conversation.')

H('When to Use Each Mechanism', 2)
T(
    ['Scenario', 'Mechanism', 'Why Not Others?'],
    [
        ['Set coding conventions for all sessions', 'CLAUDE.md', 'Always loaded; no manual trigger needed'],
        ['Run tests after code changes', 'Skill or Hook', 'Skill for optional; Hook for mandatory enforcement'],
        ['Research a codebase pattern before modifying', 'Explore subagent', 'Keeps research out of main context; fast, read-only'],
        ['Build a feature spanning 5+ files', 'Single agent (Layer 1)', 'Sequential dependencies between files; needs full context'],
        ['Build frontend + backend + tests for a new module', 'Agent Team', 'Each layer is independent; parallel development'],
        ['Process 301 podcast transcripts', 'Subagent pipeline', 'Each transcript is independent; highly parallelisable'],
        ['Plan architecture for a new module', 'Plan subagent', 'Architectural exploration shouldn\'t consume main context'],
    ]
)

doc.add_page_break()

# ============================================================================
# 3. ANTHROPIC'S 5 WORKFLOW PATTERNS
# ============================================================================
H('3. Anthropic\'s 5 Workflow Patterns & When to Use Each', 1)

Q('"Finding the simplest solution possible, and only increasing complexity when needed." \u2014 Anthropic')

para(
    'Anthropic identifies 5 distinct workflow patterns for agent orchestration. Each has specific '
    'strengths and appropriate use cases for Frontera development.'
)

H('Pattern 1: Prompt Chaining', 2)
para(
    'Sequential steps where each LLM call processes the previous output. Includes programmatic '
    '"gates" for quality checks between steps.'
)
P('Best for: ', 'Fixed subtasks where trading latency for accuracy is beneficial.')
P('Frontera Example: ', 'Database migration \u2192 API route creation \u2192 TypeScript type generation \u2192 '
  'Component scaffold. Each step depends on the previous output, with validation gates between.')

H('Pattern 2: Routing', 2)
para(
    'Classifies inputs and directs them to specialised downstream tasks with tailored prompts.'
)
P('Best for: ', 'Distinct input categories requiring different handling.')
P('Frontera Example: ', 'Feature request triage: "Is this a UI change, API change, AI prompt change, '
  'or database change?" \u2192 routes to the appropriate skill or agent configuration.')

H('Pattern 3: Parallelisation', 2)
para(
    'Executes multiple LLM calls simultaneously \u2014 either breaking tasks into independent sections '
    'or running identical tasks multiple times for diverse outputs.'
)
P('Best for: ', 'Speed improvements or higher-confidence results through multiple perspectives.')
P('Frontera Example: ', 'Processing 301 podcast transcripts for Expert Perspectives. Each transcript '
  'is independent. Launch 5-10 subagents simultaneously, each processing a batch.')

H('Pattern 4: Orchestrator-Workers', 2)
para(
    'A central LLM dynamically breaks tasks into subtasks, delegates to worker LLMs, then '
    'synthesises results. This is the Agent Teams pattern.'
)
P('Best for: ', 'Complex, unpredictable subtask requirements.')
P('Frontera Example: ', 'Phase 5 development: orchestrator identifies subtasks (API route, UI component, '
  'database migration, AI prompt, test suite), assigns each to a worker agent, synthesises into '
  'a cohesive feature.')

H('Pattern 5: Evaluator-Optimiser', 2)
para(
    'One LLM generates responses while another evaluates and provides iterative feedback in a loop.'
)
P('Best for: ', 'Tasks with clear success criteria where iterative refinement improves output.')
P('Frontera Example: ', 'AI prompt engineering for the Strategy Coach. One agent writes the system prompt; '
  'another runs the eval suite against it; the first refines based on eval results. Loop until '
  'quality metrics pass.')

doc.add_page_break()

H('Pattern Selection Guide for Frontera', 2)
T(
    ['Frontera Task', 'Pattern', 'Rationale'],
    [
        ['Single feature (Express Mode)', 'Prompt Chaining', 'Sequential: UI \u2192 API \u2192 Agent logic, each building on previous'],
        ['Micro-synthesis implementation', 'Prompt Chaining', 'API route \u2192 synthesis helper \u2192 UI component \u2192 tests'],
        ['Phase 5: Strategic Activation', 'Orchestrator-Workers', 'Multiple independent subtasks across layers'],
        ['Transcript ingestion (UC1, UC3)', 'Parallelisation', '301 independent transcripts; maximise throughput'],
        ['AI prompt refinement', 'Evaluator-Optimiser', 'Clear eval metrics; iterative improvement loop'],
        ['Bug triage and fixing', 'Routing', 'Classify bug type \u2192 route to appropriate fix pattern'],
        ['Integration development (Slack, Jira)', 'Parallelisation', 'Each integration is independent; build simultaneously'],
        ['Persona development (UC2)', 'Prompt Chaining', 'Define persona \u2192 implement \u2192 integrate \u2192 test'],
        ['Test suite expansion', 'Parallelisation', 'Unit, integration, component tests can be written independently'],
    ]
)

doc.add_page_break()

# ============================================================================
# 4. CLAUDE.MD
# ============================================================================
H('4. CLAUDE.md: Your Project\'s Agent Constitution', 1)

Q('"The single most important file in your codebase for using Claude Code effectively." \u2014 Builder.io')

para(
    'Frontera already has a comprehensive CLAUDE.md. This section covers best practices for '
    'maintaining and extending it as the platform grows.'
)

H('Current Strengths of Frontera\'s CLAUDE.md', 2)
B([
    'Comprehensive tech stack documentation (Next.js 15, TypeScript, Tailwind v4, Clerk, Supabase, Anthropic)',
    'Clear project structure with file-level detail',
    'Design system with colour palette, typography, component patterns',
    'Database schema documentation with key table relationships',
    'Testing framework documentation across all 8 phases',
    'API patterns and authentication flow documentation',
    'PDF generation pattern with PDFKit',
])

H('Recommendations for Improvement', 2)
P('1. Keep it under 200 lines in the root file: ', 'Frontera\'s CLAUDE.md is currently very long. '
  'Research shows that Claude starts ignoring instructions beyond ~80 lines. Move detailed sections '
  'to subdirectory CLAUDE.md files.')

P('Recommended split:', '')
B([
    'Root CLAUDE.md (80 lines max): Tech stack, project structure, key commands, core patterns, design palette summary',
    'src/CLAUDE.md: Component patterns, coding conventions, TypeScript rules',
    'src/app/api/CLAUDE.md: API route patterns, authentication, error handling, Supabase patterns',
    'src/lib/agents/CLAUDE.md: AI agent patterns, streaming, prompt engineering guidelines',
    'src/components/CLAUDE.md: Design system details, component library, accessibility rules',
    'tests/CLAUDE.md: Test framework patterns, mock utilities, eval patterns',
    'Background/CLAUDE.md: Document generation patterns, python-docx usage',
])

P('2. Add a "What NOT to Do" section: ', 'Explicit anti-patterns are often more useful than positive guidelines.')
B([
    'Do NOT use CSS modules (Tailwind only)',
    'Do NOT use default exports (named exports only)',
    'Do NOT add new dependencies without checking for React 19 compatibility',
    'Do NOT modify the Clerk webhook handler without understanding the org sync flow',
    'Do NOT use gray-* colours (use slate-* per design system)',
])

P('3. Add "Common Gotchas": ', 'Capture hard-won knowledge from previous development sessions.')
B([
    'Binary .docx files cannot be read with the Read tool \u2014 use Python with python-docx',
    'WebFetch fails on JavaScript-rendered SPAs (Maze, etc.) \u2014 use WebSearch instead',
    'Tailwind v4 has different configuration syntax from v3',
    'Supabase service role key is required for server-side operations',
    'Clerk auth() is async in Next.js 15 App Router',
])

doc.add_page_break()

# ============================================================================
# 5. SKILLS, HOOKS & CUSTOM COMMANDS
# ============================================================================
H('5. Skills, Hooks & Custom Commands', 1)

H('Skills You Should Create for Frontera', 2)

para(
    'Frontera already has several custom skills. Based on the platform\'s development patterns, '
    'here are additional skills that would accelerate development:'
)

T(
    ['Skill Name', 'Trigger', 'What It Does', 'Priority'],
    [
        ['agent-extend', 'User says "extend agent"', 'Scaffolds new coaching capabilities following existing patterns', 'Exists'],
        ['api-route', 'User says "create api"', 'Creates API routes following Frontera auth/Supabase patterns', 'Exists'],
        ['db-migration', 'User says "add table"', 'Creates Supabase migration following schema conventions', 'Exists'],
        ['feature', 'User says "add feature"', 'Scaffolds full feature (API + UI + types + tests)', 'Exists'],
        ['test-gen', 'User says "write tests"', 'Generates tests following Frontera test patterns', 'Exists'],
        ['frontend-design', 'User says "build component"', 'Creates UI following Frontera design system', 'Exists'],
        ['prompt-engineer', 'User says "improve prompt"', 'Iterates on AI prompts using eval suite results', 'Create'],
        ['canvas-section', 'User says "new canvas section"', 'Scaffolds new CanvasPanel section following existing patterns', 'Create'],
        ['export-pdf', 'User says "add PDF export"', 'Creates PDFKit-based PDF export following established pattern', 'Create'],
        ['eval-test', 'User says "add eval"', 'Creates AI evaluation test following Phase 8 patterns', 'Create'],
        ['module-plan', 'User says "plan module"', 'Plans new module architecture using Plan subagent', 'Create'],
    ]
)

H('Hooks for Quality Enforcement', 2)

para(
    'Hooks are deterministic rules that complement the "should-do" guidance in CLAUDE.md. '
    'They are critical for enterprise development because they enforce standards automatically.'
)

P('Recommended Hooks:', '')
T(
    ['Hook Type', 'Trigger', 'Action', 'Purpose'],
    [
        ['PreToolUse', 'Bash(git commit)', 'Block unless tests pass', 'Prevent commits with failing tests'],
        ['PreToolUse', 'Write(*.tsx)', 'Lint check on file', 'Enforce TypeScript/ESLint standards'],
        ['PostToolUse', 'Edit(src/lib/agents/*)', 'Run eval suite', 'Catch AI quality regressions after prompt changes'],
        ['PreToolUse', 'Bash(git push)', 'Require user confirmation', 'Prevent accidental pushes to master'],
        ['PostToolUse', 'Write(src/app/api/*)', 'Verify auth() pattern', 'Ensure all API routes have authentication'],
    ]
)

P('Implementation pattern: ', 'Create a .claude/hooks/ directory with shell scripts that enforce each rule. '
  'Configure in .claude/settings.json.')

doc.add_page_break()

# ============================================================================
# 6. MULTI-AGENT ORCHESTRATION
# ============================================================================
H('6. Multi-Agent Orchestration: Agent Teams & Swarms', 1)

Q('"LLMs perform worse as context expands. The more information in the context window, the harder it is for the model to focus. Teams solve this by giving each agent narrow scope and clean context." \u2014 Addy Osmani')

H('How Agent Teams Work', 2)
para(
    'Agent Teams (also called Swarms) let you orchestrate multiple Claude Code sessions working '
    'together on a shared project. The architecture is:'
)

B([
    'Lead Agent: coordinates work, assigns tasks, and synthesises results. Restricted to coordination tools only (no direct code editing)',
    'Teammate Agents: work independently, each with their own context window, on assigned tasks. They can communicate with each other and the lead',
    'Git Worktree Isolation: each teammate works in a separate Git worktree, preventing file conflicts during parallel development',
    'Inbox-Based Communication: agents communicate through an inbox system, not shared context',
    'Self-Claiming: when a teammate completes a task, they can self-claim the next available task',
])

H('The 4 Orchestration Models', 2)
T(
    ['Model', 'Structure', 'Best For', 'Frontera Use Case'],
    [
        ['Leader', 'Hierarchical: lead assigns, teammates execute', 'Well-defined tasks with clear deliverables', 'Phase 5: lead assigns API, UI, tests, AI prompt to separate teammates'],
        ['Swarm', 'Parallel: all agents work on similar tasks simultaneously', 'Batch processing of independent items', 'Transcript ingestion: each agent processes a batch of 30 transcripts'],
        ['Pipeline', 'Sequential: output of one agent feeds into the next', 'Multi-stage processing workflows', 'Data pipeline: extract \u2192 transform \u2192 load \u2192 validate'],
        ['Watchdog', 'Quality monitoring: one agent reviews others\' work', 'Critical code that needs review', 'Security-sensitive changes: one agent writes, another reviews for vulnerabilities'],
    ]
)

H('When to Use Agent Teams for Frontera', 2)
T(
    ['Use Agent Teams When...', 'Use Single Agent When...'],
    [
        ['Feature spans 3+ independent layers (frontend, backend, AI, tests)', 'Changes are sequential and interdependent'],
        ['Work can genuinely be parallelised across file boundaries', 'Changes affect the same files'],
        ['Each subtask takes 30+ minutes independently', 'Total task is under 30 minutes'],
        ['Clear file ownership boundaries exist', 'Tight coupling between components'],
        ['Development time matters more than token cost', 'Token efficiency is a concern'],
    ]
)

H('Agent Team Best Practices', 2)
P('1. Specification-Driven: ', '80% planning, 20% execution. Detailed upfront specifications make '
  'agent delegation work well. Write comprehensive task descriptions before spawning teammates.')

P('2. File Ownership: ', 'Assign distinct file sets to each teammate. Never let two agents edit the same file. '
  'Example: Agent A owns src/app/api/*, Agent B owns src/components/*, Agent C owns tests/*.')

P('3. Delegate Mode: ', 'Use Shift+Tab to restrict the lead agent to coordination-only tools. '
  'Prevents the lead from implementing instead of delegating.')

P('4. Context Inheritance: ', 'Teammates automatically receive CLAUDE.md, MCP servers, and skills, '
  'but NOT conversation history. Include task-specific details in spawn prompts.')

P('5. Documentation as Compounding: ', 'Create an AGENTS.md file that codifies learnings from agent '
  'sessions. Future agents benefit from past work, creating compounding efficiency.')

P('6. 5-6 Tasks per Teammate: ', 'Optimal balance between coordination overhead and productive work. '
  'Too few tasks waste the agent spawn cost; too many risk context degradation.')

doc.add_page_break()

# ============================================================================
# 7. CLAUDE AGENT SDK
# ============================================================================
H('7. The Claude Agent SDK: Building Custom Agents', 1)

para(
    'The Claude Agent SDK (available in Python and TypeScript) gives you the same tools, agent loop, '
    'and context management that power Claude Code, but programmable. This is relevant for Frontera '
    'because the platform already uses the Anthropic API for its Strategy Coach \u2014 the Agent SDK '
    'could enhance or replace the current streaming implementation.'
)

H('When to Use the Agent SDK vs. Direct API', 2)
T(
    ['Use Agent SDK When...', 'Use Direct API When...'],
    [
        ['You need tool use with automatic orchestration', 'Simple prompt \u2192 response (no tools)'],
        ['The agent needs to make multi-step decisions', 'Single-shot generation (synthesis, bets)'],
        ['You want file system access, bash execution', 'Output is text only (coaching responses)'],
        ['Building a backend processing pipeline', 'Streaming UI responses to the user'],
        ['Complex workflows with verification loops', 'Simple API endpoints'],
    ]
)

H('Agent SDK Architecture', 2)
para('The SDK operates through a four-phase feedback loop:')
B([
    'Gather Context: Search files, fetch data via APIs, read relevant information',
    'Take Action: Execute tools, write code, run bash commands',
    'Verify Work: Check output against rules, visual feedback, or LLM-as-judge evaluation',
    'Iterate: Refine based on verification results, loop until quality passes',
])

H('Frontera Use Cases for the Agent SDK', 2)

P('1. Enhanced AI Research Assistant: ', 'Currently, the AI Research Assistant calls Claude once to generate '
  'synthetic documents. With the Agent SDK, it could: search the web, download actual documents, '
  'extract content, synthesise findings, and verify against the company\'s context \u2014 all autonomously.')

P('2. Territory Research Auto-Population: ', 'An Agent SDK pipeline could read uploaded documents, '
  'identify relevant content for each research area, and pre-populate draft answers. The user then '
  'reviews and refines rather than writing from scratch.')

P('3. Synthesis Refinement Loop: ', 'Instead of one-shot synthesis generation, an SDK agent could '
  'generate initial synthesis, evaluate it against quality criteria, identify weak evidence, '
  'request clarification from the coaching conversation, and iterate until quality thresholds are met.')

P('4. Transcript Ingestion Pipeline (UC1, UC3): ', 'Processing 301 transcripts is a batch job perfectly '
  'suited to an Agent SDK pipeline with parallelisation, error handling, and progress tracking.')

P('5. Integration Orchestrator: ', 'An Agent SDK orchestrator could manage data flow between Frontera '
  'and external tools (Jira, Slack, Notion), handling authentication, data transformation, and '
  'error recovery automatically.')

H('Key SDK Concepts', 2)
T(
    ['Concept', 'Description', 'Relevance to Frontera'],
    [
        ['Custom Tools', 'Python/TS functions exposed as tools for Claude to use', 'Database queries, external API calls, file processing'],
        ['In-Process MCP', 'MCP servers running within the agent process', 'No separate server process needed; direct function calls'],
        ['Subagents', 'Isolated context windows for parallel work', 'Transcript processing, research tasks'],
        ['Context Compaction', 'Automatic summarisation when approaching context limits', 'Long-running agent sessions for complex research'],
        ['Verification', 'Rules-based, visual, or LLM-as-judge output validation', 'AI eval suite integration; prompt quality checks'],
    ]
)

link_ref('Claude Agent SDK Overview', 'https://platform.claude.com/docs/en/agent-sdk/overview')
link_ref('Agent SDK Python Repo', 'https://github.com/anthropics/claude-agent-sdk-python')
link_ref('Building Agents with Agent SDK (Blog)', 'https://claude.com/blog/building-agents-with-the-claude-agent-sdk')
link_ref('Building Effective Agents (Research)', 'https://www.anthropic.com/research/building-effective-agents')

doc.add_page_break()

# ============================================================================
# 8. FRONTERA BUILD PLAN: AGENT STRATEGY BY PHASE
# ============================================================================
H('8. Frontera Build Plan: Agent Strategy by Phase', 1)

para(
    'This section maps each planned Frontera development phase to a specific agent strategy, '
    'with detailed task breakdowns and recommended approaches.'
)

# Phase: PSC Enhancements
H('Phase A: Product Strategy Coach Enhancements', 2)
P('Timeline: ', 'Sprints 1-5 (per Critical Analysis Enhancement Roadmap)')
P('Agent Approach: ', 'Layer 1 (single agent) with Layer 2 (research subagents)')

T(
    ['Sprint', 'Feature', 'Agent Pattern', 'Task Breakdown'],
    [
        ['1', 'Reduce questions to 27', 'Single agent', 'Edit 3 DeepDive components; update synthesis helpers; run tests'],
        ['1', 'Add confidence ratings', 'Prompt Chaining', 'DB type update \u2192 API route \u2192 UI component \u2192 synthesis integration'],
        ['1', 'Post-phase reflection prompts', 'Single agent', 'Update system prompt; add coaching behaviours'],
        ['1', 'Estimated time per phase', 'Single agent', 'UI component addition; session time tracking'],
        ['2-3', 'Micro-synthesis after territory', 'Prompt Chaining', 'New API route \u2192 synthesis helper \u2192 UI component \u2192 coach integration'],
        ['2-3', 'Strategy on a Page generator', 'Prompt Chaining', 'Data aggregation \u2192 template design \u2192 PDF generation \u2192 UI trigger'],
        ['2-3', 'Strategic Guardrails generator', 'Single agent', 'New synthesis output type; prompt engineering; UI display'],
        ['4-5', 'Assumption tracker dashboard', 'Prompt Chaining', 'DB schema \u2192 API routes \u2192 UI dashboard \u2192 coach integration'],
        ['4-5', 'Blind spot detection', 'Evaluator-Optimiser', 'Write detection logic; eval against test cases; iterate'],
    ]
)

P('Key Skill to Use: ', 'agent-extend for coaching behaviour changes; api-route for new endpoints; '
  'test-gen after each feature')

# Phase 5
H('Phase B: Strategic Activation (Phase 5)', 2)
P('Timeline: ', 'Sprints 6-8')
P('Agent Approach: ', 'Layer 3 (Agent Team) \u2014 this is the first feature that warrants full multi-agent orchestration')

para('Recommended Agent Team Structure:')
T(
    ['Agent', 'Role', 'File Ownership', 'Tasks'],
    [
        ['Lead', 'Orchestrator', 'None (coordination only)', 'Decompose Phase 5 into subtasks; assign; synthesise'],
        ['Agent A', 'Backend Developer', 'src/app/api/product-strategy-agent/activation/*', 'API routes: team briefs, OKRs, guardrails, communication pack generation'],
        ['Agent B', 'Frontend Developer', 'src/components/product-strategy-agent/CanvasPanel/Activation*', 'UI components: activation dashboard, team brief cards, OKR display, guardrails view'],
        ['Agent C', 'AI/Prompt Engineer', 'src/lib/agents/strategy-coach/* (prompt updates only)', 'System prompt updates for Phase 5 coaching; generation prompts for each output type'],
        ['Agent D', 'Test & Quality', 'tests/unit/*, tests/integration/*, tests/evals/*', 'Unit tests, integration tests, and AI eval tests for all new features'],
    ]
)

P('Pre-requisite: ', 'Database migration for activation phase tables (run BEFORE launching agent team)')

doc.add_page_break()

# Expert Perspectives
H('Phase C: Expert Perspectives (UC1)', 2)
P('Timeline: ', 'Sprint 9+')
P('Agent Approach: ', 'Layer 2 (Parallelised subagents for transcript ingestion)')

para('The transcript ingestion pipeline is highly parallelisable:')
B([
    'Step 1: Single agent creates database schema and ingestion API routes',
    'Step 2: Parallelise transcript processing across subagents (batch of 30 per agent, 10 agents = 300 transcripts)',
    'Step 3: Single agent builds Expert Sources UI panel and coach citation integration',
    'Step 4: Run eval suite to verify citation quality and relevance',
])

P('Agent SDK Opportunity: ', 'This is the best candidate for a custom Agent SDK pipeline. '
  'Build a Python script using the Agent SDK that processes transcripts with automatic retry, '
  'progress tracking, and quality validation.')

# UC2-UC5
H('Phase D: Sparring Partners, Case Studies, Tension Simulator, Leadership Playbook', 2)
P('Timeline: ', 'Sprint 10+')

T(
    ['Module', 'Agent Approach', 'Rationale'],
    [
        ['UC2: Sparring Partners', 'Single agent (Layer 1)', 'Contained: 3 persona files + registry update + UI grid update. Sequential, interdependent'],
        ['UC3: Case Study Engine', 'Parallelised subagents (Layer 2)', 'Extraction from 301 transcripts is independent; similar pipeline to UC1'],
        ['UC4: Tension Simulator', 'Prompt Chaining (Layer 1+2)', 'Data model \u2192 tension map \u2192 UI debate component \u2192 synthesis integration \u2192 tests'],
        ['UC5: Leadership Playbook', 'Evaluator-Optimiser (Layer 1)', 'Prompt engineering with eval loop; iterative quality improvement'],
    ]
)

# Integrations
H('Phase E: Integration Layer', 2)
P('Timeline: ', 'Sprint 11+')
P('Agent Approach: ', 'Layer 3 (Agent Team) \u2014 each integration is independent')

T(
    ['Agent', 'Integration', 'Scope'],
    [
        ['Agent A', 'Slack', 'OAuth setup, strategy notification bot, channel context pushes'],
        ['Agent B', 'Jira/Linear', 'Bidirectional sync: strategic bets \u2194 backlog items'],
        ['Agent C', 'Notion/Confluence', 'Strategy on a Page export to wiki pages'],
        ['Agent D', 'Calendar', 'Kill-date reminders, strategy review scheduling'],
    ]
)

doc.add_page_break()

# ============================================================================
# 9. RECOMMENDED AGENT ARCHITECTURE
# ============================================================================
H('9. Recommended Agent Architecture for Frontera', 1)

H('The 3-Layer Development Model', 2)

T(
    ['Layer', 'What', 'When', 'Cost', 'Context Efficiency'],
    [
        ['Layer 1: Interactive', 'Single Claude Code session, conversational development', 'Daily development, bug fixes, small features, code reviews', 'Base cost', '80-90% context usage'],
        ['Layer 2: Subagents', 'Main agent delegates research/exploration to Task subagents', 'Before modifying code, parallel research, batch processing', 'Base + subagent costs', '~60% (main) + isolated subagent contexts'],
        ['Layer 3: Agent Teams', 'Multiple Claude Code sessions with lead coordination', 'Large features spanning 3+ layers, batch processing, integration development', '3-5x base cost', '~40% per agent (clean, focused context)'],
    ]
)

H('CLAUDE.md Architecture Recommendations', 2)
B([
    'Split root CLAUDE.md to under 80 lines with @imports for detailed sections',
    'Add subdirectory CLAUDE.md files: src/app/api/, src/lib/agents/, src/components/, tests/',
    'Create .claude/AGENTS.md for compounding agent learnings across sessions',
    'Add "What NOT to Do" and "Common Gotchas" sections based on development history',
    'Keep design system details in src/components/CLAUDE.md, not root',
])

H('Skill Development Roadmap', 2)
B([
    'Create prompt-engineer skill for AI prompt iteration with eval loop integration',
    'Create canvas-section skill for scaffolding new CanvasPanel sections',
    'Create export-pdf skill for new PDFKit-based exports following the established pattern',
    'Create eval-test skill for Phase 8 AI evaluation test generation',
    'Create module-plan skill that uses Plan subagent for architecture design',
])

H('Hook Strategy', 2)
B([
    'Implement test-before-commit hook to enforce green tests before any commit',
    'Implement auth-check hook to verify all API routes include authentication',
    'Implement eval-after-prompt-change hook to catch AI quality regressions',
    'Implement design-system-check hook to enforce colour palette and component patterns',
])

H('The Compound Engineering Principle', 2)
Q('"80% planning and review, 20% execution. Detailed upfront specifications make agent delegation work well."')

para(
    'For Frontera, this means investing heavily in CLAUDE.md, skills, and hooks BEFORE '
    'launching large agent teams. Each investment compounds: a well-written CLAUDE.md benefits '
    'every future agent session. A well-designed skill accelerates every similar feature. '
    'A well-configured hook prevents every similar bug.'
)

B([
    'Write comprehensive task specifications before spawning agent teams',
    'Create SKILL.md files for repeating patterns (you\'ve already started this)',
    'Document learnings in AGENTS.md after each agent team session',
    'Refine CLAUDE.md based on what agents get wrong \u2014 prevention over correction',
    'Build hooks for the mistakes you see agents making repeatedly',
])

doc.add_page_break()

# ============================================================================
# 10. REFERENCE DIRECTORY
# ============================================================================
H('10. Reference Directory', 1)

H('Official Anthropic Resources', 2)
T(
    ['Resource', 'URL', 'Key Content'],
    [
        ['Building Effective Agents', 'anthropic.com/research/building-effective-agents', '5 workflow patterns, when to use agents vs workflows'],
        ['Claude Agent SDK Overview', 'platform.claude.com/docs/en/agent-sdk/overview', 'SDK architecture, tool use, context management'],
        ['Agent SDK Quickstart', 'platform.claude.com/docs/en/agent-sdk/quickstart', 'Getting started with custom agents'],
        ['Agent SDK Python Repo', 'github.com/anthropics/claude-agent-sdk-python', 'Source code, examples, documentation'],
        ['Claude Code Docs', 'code.claude.com/docs/en/overview', 'Full Claude Code documentation'],
        ['Claude Code Skills', 'code.claude.com/docs/en/skills', 'Creating and using custom skills'],
        ['Claude Code Agent Teams', 'code.claude.com/docs/en/agent-teams', 'Official agent teams documentation'],
    ]
)

H('Community Resources', 2)
T(
    ['Resource', 'URL', 'Key Content'],
    [
        ['Claude Code Agent Teams (Addy Osmani)', 'addyosmani.com/blog/claude-code-agent-teams', 'Comprehensive guide to swarm orchestration'],
        ['Hidden Multi-Agent System (Paddo)', 'paddo.dev/blog/claude-code-hidden-swarm', 'TeammateTool architecture analysis'],
        ['Complete CLAUDE.md Guide', 'builder.io/blog/claude-md-guide', 'Best practices for CLAUDE.md structure'],
        ['Claude Code Customisation Guide', 'alexop.dev/posts/claude-code-customization-guide', 'Skills, subagents, hooks comparison'],
        ['How I Use Every Feature (Shrivu)', 'blog.sshh.io/p/how-i-use-every-claude-code-feature', 'Practitioner guide to all features'],
        ['Awesome Claude Code (GitHub)', 'github.com/hesreallyhim/awesome-claude-code', 'Curated list of skills, hooks, resources'],
        ['Everything Claude Code (GitHub)', 'github.com/affaan-m/everything-claude-code', 'Battle-tested configs from hackathon winner'],
        ['Claude Code Showcase (GitHub)', 'github.com/ChrisWiles/claude-code-showcase', 'Comprehensive config examples'],
        ['Shipyard CLI Cheatsheet', 'shipyard.build/blog/claude-code-cheat-sheet', 'Commands, config, best practices'],
    ]
)

H('Multi-Agent Frameworks', 2)
T(
    ['Framework', 'URL', 'Use Case'],
    [
        ['Claude Flow', 'github.com/ruvnet/claude-flow', 'Enterprise agent orchestration with MCP'],
        ['ccswarm', 'github.com/nwiizo/ccswarm', 'Rust-native multi-agent with Git worktree isolation'],
        ['Claude Swarm', 'github.com/affaan-m/claude-swarm', 'Multi-agent orchestration with terminal UI'],
        ['Claude Code Agentrooms', 'claudecode.run', 'Multi-agent development workspace'],
    ]
)

doc.add_page_break()

# Document control
H('Document Control', 1)
T(
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ['1.0', 'February 2026', 'Frontera Product Team', 'Initial agent orchestration guide for Frontera development'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('\u2014 End of Document \u2014').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_path = r'c:\Users\deeks\frontera-platform\Background\Claude_Code_Agent_Orchestration_Guide.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
