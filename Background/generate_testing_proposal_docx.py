"""
Generate a Word document for the Product Strategy Coach User Testing Proposal
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Create document
doc = Document()

# Define colors
NAVY = RGBColor(0x1a, 0x1f, 0x3a)
GOLD = RGBColor(0xfb, 0xbf, 0x24)
CYAN = RGBColor(0x08, 0x91, 0xb2)

# Helper functions
def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(text, level=1):
    """Add a heading with consistent styling"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = NAVY
    return heading

def add_table_from_data(headers, rows, widths=None):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '1a1f3a')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)

    doc.add_paragraph()
    return table

# Title Page
title = doc.add_heading('Product Strategy Coach', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(36)
    run.font.color.rgb = NAVY

subtitle = doc.add_paragraph('User Testing Proposal for Pilot Validation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(24)
subtitle.runs[0].font.color.rgb = CYAN

doc.add_paragraph()
doc.add_paragraph()

# Document info
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Version: 1.0\n').bold = True
info.add_run('Date: February 2026\n')
info.add_run('Author: Frontera Product Team\n')
info.add_run('Status: Draft for Review')

doc.add_page_break()

# Table of Contents
add_heading('Table of Contents', 1)
toc_items = [
    ('1.', 'Executive Summary'),
    ('2.', 'Testing Objectives'),
    ('3.', 'Product Overview'),
    ('4.', 'Testing Framework'),
    ('5.', 'Maze Test Specifications'),
    ('6.', 'Moderated Session Protocol'),
    ('7.', 'Question Bank'),
    ('8.', 'Expert Review Framework'),
    ('9.', 'Participant Recruitment'),
    ('10.', 'Test Plan Timeline'),
    ('11.', 'Success Metrics'),
    ('12.', 'Budget Estimate'),
    ('13.', 'Appendices'),
]
for num, item in toc_items:
    p = doc.add_paragraph()
    p.add_run(f'{num} ').bold = True
    p.add_run(item)

doc.add_page_break()

# 1. Executive Summary
add_heading('1. Executive Summary', 1)

add_heading('Purpose', 2)
doc.add_paragraph(
    'This proposal outlines a comprehensive user testing strategy for the Product Strategy Coach pilot. '
    'The testing programme is designed to validate three critical dimensions:'
)

bullet_points = [
    'Experience & Usability - Can users navigate the interface effectively?',
    'Journey Completion - Do users progress through all methodology phases?',
    'Coaching Quality & Outputs - Is the AI coaching valuable? Are strategic outputs actionable?',
]
for point in bullet_points:
    p = doc.add_paragraph(point, style='List Bullet')

add_heading('Recommended Approach', 2)
doc.add_paragraph('A hybrid testing methodology combining:')
approaches = [
    'Maze Platform for modular, unmoderated usability testing',
    'Moderated Think-Aloud Sessions for full journey evaluation',
    'Expert Panel Review for coaching and output quality assessment',
]
for approach in approaches:
    doc.add_paragraph(approach, style='List Bullet')

add_heading('Key Outcomes', 2)
outcomes = [
    'Identify and prioritise usability friction points',
    'Validate the coaching methodology resonates with target users',
    'Assess the perceived value of AI-generated strategic outputs',
    'Gather actionable insights for pre-launch iteration',
]
for outcome in outcomes:
    doc.add_paragraph(outcome, style='List Bullet')

doc.add_page_break()

# 2. Testing Objectives
add_heading('2. Testing Objectives', 1)

add_heading('Primary Objectives', 2)
add_table_from_data(
    ['Objective', 'Success Criteria', 'Measurement Method'],
    [
        ['Validate navigation clarity', '>85% task success rate', 'Maze mission completion'],
        ['Confirm methodology comprehension', '>80% understand phase progression', 'Post-test survey'],
        ['Assess coach interaction value', 'Trust score >4.0/5.0', 'Session interviews'],
        ['Evaluate synthesis output quality', 'Actionability score >4.0/5.0', 'Expert panel review'],
        ['Identify critical friction points', '<3 blocking issues', 'Session observation'],
    ]
)

add_heading('Research Questions', 2)

add_heading('Experience & Usability', 3)
questions = [
    'How intuitive is the two-panel layout (coach + canvas)?',
    'Do users understand the relationship between chat coaching and canvas activities?',
    'Can users locate key actions (upload materials, explore territories, generate synthesis)?',
    'Is the progress stepper effective at communicating journey state?',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

add_heading('Journey & Methodology', 3)
questions = [
    'Do users understand the 4-phase methodology without explanation?',
    'What triggers users to progress from one phase to the next?',
    'Where do users abandon or express frustration?',
    'Is the "terrain mapping" metaphor meaningful?',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

add_heading('Coaching Quality', 3)
questions = [
    'Do users perceive the coach as helpful vs. intrusive?',
    'Is the coaching personalised enough to feel relevant?',
    'Do users trust AI-generated strategic guidance?',
    'What coaching behaviours increase or decrease trust?',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

add_heading('Output Value', 3)
questions = [
    'Are synthesis outputs immediately understandable?',
    'Would users share these outputs with stakeholders?',
    'Do outputs feel generic or specifically tailored?',
    'What additional outputs would users want?',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

doc.add_page_break()

# 3. Product Overview
add_heading('3. Product Overview', 1)

add_heading('Product Description', 2)
doc.add_paragraph(
    'The Product Strategy Coach is an AI-powered strategic coaching platform that guides enterprise '
    'leaders through a structured methodology to develop product transformation strategies. The experience combines:'
)
features = [
    'Persistent AI Coach Sidebar for contextual guidance and Q&A',
    'Scrollable Canvas with phased, card-based workflow',
    'Progressive Disclosure through "terrain mapping" that unlocks as users complete phases',
    'Document Upload to ground strategy in real source materials',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

add_heading('Methodology Phases', 2)
add_table_from_data(
    ['Phase', 'Name', 'Purpose', 'Key Activities'],
    [
        ['1', 'Discovery', 'Context Setting', 'Upload materials, establish strategic baseline'],
        ['2', 'Research', 'Terrain Mapping', 'Explore Company, Customer, Competitor territories'],
        ['3', 'Synthesis', 'Strategy Formation', 'Review AI-generated opportunities and insights'],
        ['4', 'Strategic Bets', 'Route Planning', 'Define hypothesis-driven strategic bets'],
    ]
)

add_heading('Target Users', 2)
doc.add_paragraph('Primary Persona: Strategic Product Leader', style='Intense Quote')
user_attrs = [
    'VP Product, Head of Strategy, Product Director',
    'Enterprise (500+ employees) or funded scale-up',
    'Currently leading or planning product transformation',
    '5+ years experience in product/strategy roles',
]
for attr in user_attrs:
    doc.add_paragraph(attr, style='List Bullet')

doc.add_page_break()

# 4. Testing Framework
add_heading('4. Testing Framework', 1)

add_heading('Testing Layers', 2)

doc.add_paragraph('LAYER 1: MODULAR USABILITY (Maze)', style='Intense Quote')
layer1 = [
    '6 focused tests, 5-10 minutes each',
    'Unmoderated, scalable recruitment',
    'Click tracking, heatmaps, task success metrics',
    'Quantitative data with qualitative follow-up',
]
for item in layer1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('LAYER 2: FULL JOURNEY (Moderated Sessions)', style='Intense Quote')
layer2 = [
    '45-60 minute think-aloud sessions',
    'Real context (user\'s actual company)',
    'Deep observation of coaching interaction',
    'Rich qualitative insights',
]
for item in layer2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('LAYER 3: OUTPUT QUALITY (Expert Review)', style='Intense Quote')
layer3 = [
    'Panel of 3-5 strategy experts',
    'Blind review of synthesis outputs',
    'Structured scoring rubric',
    'Comparative assessment',
]
for item in layer3:
    doc.add_paragraph(item, style='List Bullet')

add_heading('Why This Hybrid Approach?', 2)
add_table_from_data(
    ['Testing Need', 'Maze Fit', 'Moderated Fit', 'Expert Review'],
    [
        ['First-click accuracy', 'Excellent', 'Overkill', 'N/A'],
        ['Task completion rates', 'Excellent', 'Good', 'N/A'],
        ['Navigation patterns', 'Excellent', 'Good', 'N/A'],
        ['Emotional response', 'Limited', 'Excellent', 'N/A'],
        ['Trust assessment', 'Limited', 'Excellent', 'Limited'],
        ['Coaching quality', 'Poor', 'Excellent', 'Excellent'],
        ['Output actionability', 'Limited', 'Good', 'Excellent'],
        ['Long-form journey', 'Poor', 'Excellent', 'N/A'],
        ['Scale (many users)', 'Excellent', 'Poor', 'N/A'],
    ]
)

doc.add_page_break()

# 5. Maze Test Specifications
add_heading('5. Maze Test Specifications', 1)

add_heading('Test Overview', 2)
add_table_from_data(
    ['Test #', 'Name', 'Focus Area', 'Duration', 'Participants'],
    [
        ['1', 'First Impressions', 'Overall comprehension', '5 min', '15-20'],
        ['2', 'Discovery Phase', 'Context setting UX', '8 min', '15-20'],
        ['3', 'Coach Interaction', 'AI coaching perception', '6 min', '15-20'],
        ['4', 'Territory Navigation', 'Research phase UX', '8 min', '15-20'],
        ['5', 'Synthesis Review', 'Output comprehension', '7 min', '15-20'],
        ['6', 'End-to-End Prototype', 'Full flow (abbreviated)', '12 min', '10-15'],
    ]
)

# Test 1
add_heading('Test 1: First Impressions', 2)
doc.add_paragraph('Objective: Assess initial comprehension and emotional response to the interface')
doc.add_paragraph('Test Type: Opinion Test + First-Click Test')

add_heading('Block 1: Initial Reaction (30 seconds)', 3)

doc.add_paragraph('Question 1.1 - 5-Second Test', style='Intense Quote')
doc.add_paragraph('"Look at this screen for 5 seconds. What do you think this product does?"')
doc.add_paragraph('Type: Open text')
doc.add_paragraph('Purpose: Capture first impressions before cognitive processing')

doc.add_paragraph('Question 1.2 - Emotional Response', style='Intense Quote')
doc.add_paragraph('"What\'s your immediate reaction to this interface?"')
doc.add_paragraph('Type: Multiple choice (single select)')
options = [
    'Excited - I want to explore this',
    'Curious - I\'d like to learn more',
    'Neutral - No strong feeling',
    'Confused - I\'m not sure what I\'m looking at',
    'Overwhelmed - This looks complex',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Question 1.3 - Comprehension Check', style='Intense Quote')
doc.add_paragraph('"Based on what you see, what is the main purpose of this tool?"')
doc.add_paragraph('Type: Multiple choice (single select)')
options = [
    'A project management dashboard',
    'An AI-powered strategy coaching tool',
    'A document collaboration platform',
    'A customer analytics system',
    'I\'m not sure',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

add_heading('Block 2: Layout Understanding (60 seconds)', 3)

doc.add_paragraph('Question 1.4 - Component Identification', style='Intense Quote')
doc.add_paragraph('"What do you think the left sidebar is for?"')
doc.add_paragraph('Type: Open text')
doc.add_paragraph('Purpose: Validate coach sidebar mental model')

doc.add_paragraph('Question 1.5 - First-Click Test', style='Intense Quote')
doc.add_paragraph('"If you wanted to start working on your strategy, where would you click first?"')
doc.add_paragraph('Type: First-click (heatmap capture)')
doc.add_paragraph('Success zone: Discovery section, coach input, or upload materials button')

doc.add_paragraph('Question 1.6 - Progress Understanding', style='Intense Quote')
doc.add_paragraph('"Looking at the top of the main panel, what do the circles represent?"')
doc.add_paragraph('Type: Multiple choice (single select)')
options = [
    'Steps in a process I need to complete',
    'Different features I can access',
    'My team members',
    'Notification indicators',
    'Not sure',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

add_heading('Block 3: Value Proposition (45 seconds)', 3)

doc.add_paragraph('Question 1.7 - Value Clarity', style='Intense Quote')
doc.add_paragraph('"How clear is it what value this tool would provide to you?"')
doc.add_paragraph('Type: Linear scale 1-5 (1 = Not at all clear, 5 = Extremely clear)')

doc.add_paragraph('Question 1.8 - Relevance', style='Intense Quote')
doc.add_paragraph('"How relevant does this tool seem for your work?"')
doc.add_paragraph('Type: Linear scale 1-5 (1 = Not relevant at all, 5 = Highly relevant)')

doc.add_paragraph('Question 1.9 - Open Feedback', style='Intense Quote')
doc.add_paragraph('"What questions do you have about this interface?"')
doc.add_paragraph('Type: Open text')

doc.add_page_break()

# Test 2
add_heading('Test 2: Discovery Phase Deep Dive', 2)
doc.add_paragraph('Objective: Validate Discovery phase usability and task completion')
doc.add_paragraph('Test Type: Mission + Survey')

add_heading('Scenario Introduction', 3)
scenario = doc.add_paragraph()
scenario.add_run('Scenario: ').bold = True
scenario.add_run(
    'You\'re the VP of Product at a mid-sized fintech company. Your CEO has asked you to develop '
    'a product strategy for the next 18 months. You\'ve decided to use this AI coaching tool to '
    'help structure your thinking. You\'ve just logged in for the first time.'
)

add_heading('Mission 1: Locate Upload Function', 3)
doc.add_paragraph('Task: Find where you would upload your company\'s strategic documents.')
doc.add_paragraph('Type: Mission with click tracking')
doc.add_paragraph('Success criteria: Click on "Upload Files" button or drag-drop zone')
doc.add_paragraph('Time limit: 30 seconds')

doc.add_paragraph('Follow-up Question 2.1', style='Intense Quote')
doc.add_paragraph('"How easy was it to find the upload function?"')
doc.add_paragraph('Type: Linear scale 1-5 (1 = Very difficult, 5 = Very easy)')

add_heading('Mission 2: Understand Context Requirements', 3)
doc.add_paragraph('Task: What information does the tool need from you to get started?')
doc.add_paragraph('Type: Open text response after viewing Discovery section')

doc.add_paragraph('Follow-up Question 2.2', style='Intense Quote')
doc.add_paragraph('"Is it clear what materials you should upload?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_paragraph('Follow-up Question 2.3', style='Intense Quote')
doc.add_paragraph('"What documents would you upload if using this for your company?"')
doc.add_paragraph('Type: Multiple choice (multi-select)')
options = [
    'Annual report / financial statements',
    'Strategic plan / roadmap',
    'Market research reports',
    'Competitor analysis',
    'Customer feedback / NPS data',
    'Product specs / PRDs',
    'Board presentations',
    'I wouldn\'t know what to upload',
    'Other (please specify)',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

add_heading('Mission 3: Interact with Coach', 3)
doc.add_paragraph('Task: The AI coach has asked you a question. Respond to it.')
doc.add_paragraph('Type: Mission with interaction tracking')
doc.add_paragraph('Show: Coach message asking "What competitive dynamics are making product transformation urgent for your company?"')
doc.add_paragraph('Success criteria: Click on input field or type response')

doc.add_paragraph('Follow-up Question 2.4', style='Intense Quote')
doc.add_paragraph('"How would you describe the tone of the AI coach?"')
doc.add_paragraph('Type: Multiple choice (single select)')
options = [
    'Professional and authoritative',
    'Friendly and supportive',
    'Neutral and factual',
    'Pushy or demanding',
    'Robotic and impersonal',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Follow-up Question 2.5', style='Intense Quote')
doc.add_paragraph('"Is this question relevant to your situation?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_page_break()

# Test 3
add_heading('Test 3: Coach Interaction Evaluation', 2)
doc.add_paragraph('Objective: Assess perceptions of AI coaching quality and trust')
doc.add_paragraph('Test Type: Video Walkthrough + Opinion Survey')

add_heading('Block 1: Video Observation', 3)
doc.add_paragraph('Show: Screen recording of coach interaction including:')
items = [
    'User asking a question about competitive positioning',
    'Coach providing contextual response with follow-up question',
    'User responding',
    'Coach synthesising and providing insight',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Question 3.1 - Helpfulness', style='Intense Quote')
doc.add_paragraph('"How helpful was the coach\'s response?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_paragraph('Question 3.2 - Relevance', style='Intense Quote')
doc.add_paragraph('"How relevant was the coach\'s guidance to the user\'s question?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_paragraph('Question 3.3 - Interaction Style', style='Intense Quote')
doc.add_paragraph('"How would you describe the interaction between the user and the coach?"')
doc.add_paragraph('Type: Multiple choice (single select)')
options = [
    'Like talking to a knowledgeable consultant',
    'Like using a sophisticated chatbot',
    'Like following a guided tutorial',
    'Like having a conversation with a colleague',
    'Like receiving automated suggestions',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

add_heading('Block 2: Coaching Scenarios', 3)
doc.add_paragraph('Present 3 different coach response examples:')

doc.add_paragraph('Scenario A: Proactive Insight', style='Intense Quote')
doc.add_paragraph(
    'Coach message: "I notice you mentioned \'integration depth\' as a customer priority, but your '
    'competitor analysis suggests most alternatives struggle with this. This could be a significant '
    '\'How to Win\' opportunity. Would you like to explore what capabilities you\'d need to excel here?"'
)

doc.add_paragraph('Question 3.4', style='Intense Quote')
doc.add_paragraph('"How valuable is it for the coach to proactively identify connections like this?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_paragraph('Question 3.5', style='Intense Quote')
doc.add_paragraph('"How much would you trust this type of insight from an AI coach?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_paragraph('Scenario B: Challenge Response', style='Intense Quote')
doc.add_paragraph(
    'Coach message: "You mentioned that \'price is not a key differentiator\' for your customers. '
    'However, your earlier notes indicated several lost deals due to pricing. Can you help me '
    'understand this apparent tension?"'
)

doc.add_paragraph('Question 3.6', style='Intense Quote')
doc.add_paragraph('"How comfortable are you with the coach challenging your assumptions?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_paragraph('Question 3.7', style='Intense Quote')
doc.add_paragraph('"This type of challenge feels:"')
doc.add_paragraph('Type: Multiple choice (single select)')
options = [
    'Helpful - it makes me think more deeply',
    'Neutral - neither good nor bad',
    'Uncomfortable - I don\'t like being challenged by AI',
    'Annoying - the AI doesn\'t understand my context',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

add_heading('Block 3: Overall Coach Assessment', 3)

doc.add_paragraph('Question 3.9 - Trust Comparison', style='Intense Quote')
doc.add_paragraph('"Compared to a human strategy consultant, how much would you trust this AI coach?"')
doc.add_paragraph('Type: Linear scale 1-7 (1 = Much less, 4 = About the same, 7 = Much more)')

doc.add_paragraph('Question 3.10 - Ideal Coach Behaviour', style='Intense Quote')
doc.add_paragraph('"What behaviours would make you trust the AI coach more?"')
doc.add_paragraph('Type: Multiple choice (multi-select)')
options = [
    'Showing its reasoning/sources',
    'Asking clarifying questions before advising',
    'Admitting when it doesn\'t know something',
    'Providing options rather than single recommendations',
    'Referencing industry benchmarks or data',
    'Remembering context from earlier in the session',
    'Other (please specify)',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_page_break()

# Test 4
add_heading('Test 4: Territory Navigation (Research Phase)', 2)
doc.add_paragraph('Objective: Validate Research phase navigation and territory exploration UX')
doc.add_paragraph('Test Type: Mission + Heatmap + Survey')

add_heading('Scenario Introduction', 3)
doc.add_paragraph(
    'You\'ve completed the Discovery phase and uploaded your strategic documents. The tool is now '
    'asking you to "Map Your Strategic Terrain" by exploring three territories: Company, Customer, '
    'and Competitor.'
)

add_heading('Mission 1: Territory Selection', 3)
doc.add_paragraph('Task: You want to start by understanding your customers better. Where would you click?')
doc.add_paragraph('Type: First-click test')
doc.add_paragraph('Success zone: Customer territory card')
doc.add_paragraph('Capture: Heatmap of all clicks')

doc.add_paragraph('Question 4.1', style='Intense Quote')
doc.add_paragraph('"How clear was it which card to select?"')
doc.add_paragraph('Type: Linear scale 1-5')

add_heading('Mission 2: Navigate Territory Deep Dive', 3)
doc.add_paragraph('Task: You\'ve clicked on Customer territory. Now find where you would answer questions about customer segments.')
doc.add_paragraph('Type: Mission in prototype')
doc.add_paragraph('Success criteria: Navigate to "Segments & Needs" research area')

doc.add_paragraph('Question 4.2', style='Intense Quote')
doc.add_paragraph('"How easy was it to find the right section?"')
doc.add_paragraph('Type: Linear scale 1-5')

add_heading('Mission 3: Complete Research Input', 3)
doc.add_paragraph('Task: Answer the research questions you see for this area.')
doc.add_paragraph('Type: Observe interaction with research question inputs')

doc.add_paragraph('Question 4.3', style='Intense Quote')
doc.add_paragraph('"How clear were the research questions?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_paragraph('Question 4.4', style='Intense Quote')
doc.add_paragraph('"How relevant were these questions to developing a product strategy?"')
doc.add_paragraph('Type: Linear scale 1-5')

add_heading('Block: Territory Card Comprehension', 3)

doc.add_paragraph('Question 4.7', style='Intense Quote')
doc.add_paragraph('"The tool uses a \'terrain mapping\' metaphor with \'territories\' to explore. How intuitive is this metaphor?"')
doc.add_paragraph('Type: Linear scale 1-5 (1 = Very confusing, 5 = Very intuitive)')

doc.add_paragraph('Question 4.8', style='Intense Quote')
doc.add_paragraph('"The three territories are Company, Customer, and Competitor (the \'3Cs\'). Are these categories:"')
doc.add_paragraph('Type: Multiple choice (single select)')
options = [
    'Exactly right - these are the key areas to explore',
    'Mostly right - but I\'d add or change something',
    'Somewhat relevant - but missing important areas',
    'Not how I think about strategy',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Question 4.9', style='Intense Quote')
doc.add_paragraph('"What other \'territories\' or areas would you want to explore for product strategy?"')
doc.add_paragraph('Type: Open text')

doc.add_page_break()

# Test 5
add_heading('Test 5: Synthesis Output Review', 2)
doc.add_paragraph('Objective: Assess comprehension and perceived value of synthesis outputs')
doc.add_paragraph('Test Type: Comprehension Test + Opinion Survey')

add_heading('Block 1: Strategic Opportunity Map', 3)
doc.add_paragraph('Show: Full 2x2 matrix with 4 plotted opportunities')

doc.add_paragraph('Question 5.1 - 5-Second Comprehension', style='Intense Quote')
doc.add_paragraph('"Look at this visualisation for 5 seconds. What does it show?"')
doc.add_paragraph('Type: Open text')

doc.add_paragraph('Question 5.2 - Quadrant Understanding', style='Intense Quote')
doc.add_paragraph('"What does an opportunity in the top-right quadrant (INVEST) mean?"')
doc.add_paragraph('Type: Multiple choice (single select)')
options = [
    'High priority - execute immediately',
    'Interesting but needs more research',
    'Low priority - deprioritise',
    'I\'m not sure',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Question 5.3', style='Intense Quote')
doc.add_paragraph('"How useful is this visualisation for strategic decision-making?"')
doc.add_paragraph('Type: Linear scale 1-5')

add_heading('Block 2: Opportunity Card Review', 3)
doc.add_paragraph('Show: Detailed opportunity card including:')
items = [
    'Title: "Enterprise Integration Platform"',
    'Scores: Market Attractiveness 8/10, Capability Fit 7/10',
    'Where to Play: "Mid-market consolidators (£5-50bn AUM)"',
    'How to Win: "Deep platform integration + compliance automation"',
    'Evidence: 3 linked sources from research',
    'Assumptions: "What Would Have to Be True"',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Question 5.5', style='Intense Quote')
doc.add_paragraph('"Based on this card, what strategic action is being recommended?"')
doc.add_paragraph('Type: Open text')

doc.add_paragraph('Question 5.6', style='Intense Quote')
doc.add_paragraph('"The card shows \'Evidence\' linking back to your research. How valuable is this feature?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_paragraph('Question 5.8 - Actionability', style='Intense Quote')
doc.add_paragraph('"Could you take action on this strategic opportunity based on the information provided?"')
doc.add_paragraph('Type: Multiple choice (single select)')
options = [
    'Yes, I have enough information to act',
    'Partially, but I need more detail on some areas',
    'No, I need significantly more information',
    'I\'m not sure what action to take',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

add_heading('Block 3: Output Quality Assessment', 3)

doc.add_paragraph('Question 5.11 - Stakeholder Shareability', style='Intense Quote')
doc.add_paragraph('"Would you share this synthesis output with your executive team?"')
doc.add_paragraph('Type: Linear scale 1-5 (1 = Definitely not, 5 = Definitely yes)')

doc.add_paragraph('Question 5.12', style='Intense Quote')
doc.add_paragraph('"Why or why not would you share this with stakeholders?"')
doc.add_paragraph('Type: Open text')

doc.add_paragraph('Question 5.16', style='Intense Quote')
doc.add_paragraph('"What other outputs would you want from the synthesis phase?"')
doc.add_paragraph('Type: Multiple choice (multi-select)')
options = [
    'SWOT analysis',
    'Competitive positioning map',
    'Risk assessment',
    'Implementation timeline',
    'Resource requirements',
    'Financial projections',
    'Stakeholder communication plan',
    'Other (please specify)',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_page_break()

# Test 6
add_heading('Test 6: End-to-End Prototype Flow', 2)
doc.add_paragraph('Objective: Validate full journey comprehension with abbreviated prototype')
doc.add_paragraph('Test Type: Usability Mission + Post-Test Survey')
doc.add_paragraph('Duration: 10-12 minutes total')

add_heading('Introduction', 3)
doc.add_paragraph(
    'You\'re about to experience the full Product Strategy Coach journey, from Discovery to Strategic Bets. '
    'This is an abbreviated version - in the real product, you\'d spend more time in each phase. '
    'Please think aloud as you navigate, sharing what you\'re thinking and feeling.'
)

add_heading('Phase Missions', 3)
phases = [
    ('Phase 1: Discovery (2 min)', 'Complete the Discovery phase by uploading a document and answering the coach\'s first question.'),
    ('Phase 2: Research (3 min)', 'Complete at least one territory (Customer) by answering the research questions.'),
    ('Phase 3: Synthesis (2 min)', 'Review the generated synthesis and identify the top strategic opportunity.'),
    ('Phase 4: Strategic Bets (2 min)', 'Create a strategic bet based on the synthesis.'),
]
for phase, mission in phases:
    p = doc.add_paragraph()
    p.add_run(phase + ': ').bold = True
    p.add_run(mission)

add_heading('Post-Flow Survey', 3)

doc.add_paragraph('Question 6.1', style='Intense Quote')
doc.add_paragraph('"How clear was the journey from start to finish?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_paragraph('Question 6.2', style='Intense Quote')
doc.add_paragraph('"Were the transitions between phases smooth and logical?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_paragraph('Question 6.3', style='Intense Quote')
doc.add_paragraph('"For the value provided, the time investment in this process is:"')
doc.add_paragraph('Type: Multiple choice (single select)')
options = [
    'Worth much more than the time required',
    'Worth the time required',
    'Worth somewhat less than the time required',
    'Not worth the time required',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Question 6.5', style='Intense Quote')
doc.add_paragraph('"How satisfied are you with the overall Product Strategy Coach experience?"')
doc.add_paragraph('Type: Linear scale 1-5')

doc.add_paragraph('Question 6.6 - Net Promoter Score', style='Intense Quote')
doc.add_paragraph('"How likely are you to recommend this tool to a colleague?"')
doc.add_paragraph('Type: Linear scale 0-10 (NPS format)')

doc.add_paragraph('Question 6.7', style='Intense Quote')
doc.add_paragraph('"What is the most valuable aspect of this tool?"')
doc.add_paragraph('Type: Open text')

doc.add_paragraph('Question 6.8', style='Intense Quote')
doc.add_paragraph('"What is your biggest concern or hesitation about using this tool?"')
doc.add_paragraph('Type: Open text')

doc.add_paragraph('Question 6.10', style='Intense Quote')
doc.add_paragraph('"Would you or your company pay for a tool like this?"')
doc.add_paragraph('Type: Multiple choice (single select)')
options = [
    'Yes, definitely',
    'Yes, probably',
    'Not sure',
    'Probably not',
    'Definitely not',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_page_break()

# 6. Moderated Session Protocol
add_heading('6. Moderated Session Protocol', 1)

add_heading('Session Overview', 2)
add_table_from_data(
    ['Element', 'Details'],
    [
        ['Duration', '60 minutes'],
        ['Format', 'Remote video call with screen sharing'],
        ['Recording', 'Video + audio (with consent)'],
        ['Participants', '6-8 strategic product leaders'],
    ]
)

add_heading('Session Structure', 2)
add_table_from_data(
    ['Section', 'Duration', 'Purpose'],
    [
        ['1. Introduction', '5 minutes', 'Set expectations, obtain consent'],
        ['2. Context Setting', '5 minutes', 'Understand participant background'],
        ['3. First Impressions', '5 minutes', 'Initial reaction to interface'],
        ['4. Discovery Phase', '10-15 minutes', 'Walkthrough with real context'],
        ['5. Research Phase', '15-20 minutes', 'Territory exploration'],
        ['6. Synthesis Phase', '10 minutes', 'Output review'],
        ['7. Debrief Interview', '10-15 minutes', 'Structured feedback'],
    ]
)

add_heading('Facilitator Script: Introduction', 3)
script = doc.add_paragraph()
script.add_run(
    '"Thank you for joining today. I\'m [name], and I\'ll be facilitating this session.\n\n'
    'We\'re developing an AI-powered tool to help product leaders develop strategy. Today, I\'d like '
    'you to use the tool and share your thoughts as you go.\n\n'
    'A few things before we start:\n'
    '- Please think aloud - share what you\'re looking at, thinking, and feeling\n'
    '- There are no wrong answers. If something is confusing, that\'s valuable feedback\n'
    '- We\'re testing the product, not you\n'
    '- Feel free to be candid - criticism helps us improve\n\n'
    'Do you have any questions before we begin?"'
)
script.italic = True

add_heading('Context Setting Questions', 3)
questions = [
    '"Tell me briefly about your role and company."',
    '"What strategic challenges are you currently working on?"',
    '"Have you used any tools or consultants to help with product strategy before?"',
    '"What did you like or dislike about those experiences?"',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

add_heading('During-Task Observation Prompts', 3)
add_table_from_data(
    ['Situation', 'Prompt'],
    [
        ['Participant pauses', '"What are you looking at right now?"'],
        ['Participant seems confused', '"What are you expecting to see here?"'],
        ['Participant clicks unexpectedly', '"What made you click there?"'],
        ['Participant skips content', '"I noticed you scrolled past that - was it not relevant?"'],
        ['Participant sighs/frustrated', '"What\'s causing that reaction?"'],
        ['Participant smiles/satisfied', '"What do you like about this?"'],
    ]
)

add_heading('Debrief Interview Questions', 3)

doc.add_paragraph('Overall Experience', style='Intense Quote')
debrief_overall = [
    '"How would you describe this experience in a few words?"',
    '"What was the most valuable part?"',
    '"What was the most frustrating part?"',
]
for q in debrief_overall:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph('Coaching Quality', style='Intense Quote')
debrief_coach = [
    '"How did the AI coach compare to working with a human consultant?"',
    '"Did you trust the coach\'s guidance? Why or why not?"',
    '"What would make you trust it more?"',
]
for q in debrief_coach:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph('Output Value', style='Intense Quote')
debrief_output = [
    '"How actionable were the strategic outputs?"',
    '"What would you do next with this synthesis?"',
    '"What outputs are missing that you\'d want?"',
]
for q in debrief_output:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph('Adoption Likelihood', style='Intense Quote')
debrief_adoption = [
    '"Would you use this tool for a real strategic project?"',
    '"What would need to change for you to use it?"',
    '"What would you be willing to pay for this?"',
]
for q in debrief_adoption:
    doc.add_paragraph(q, style='List Bullet')

doc.add_page_break()

# 7. Question Bank (abbreviated)
add_heading('7. Question Bank', 1)

add_heading('Post-Session Survey (Quantitative)', 2)
add_table_from_data(
    ['Question', 'Scale', 'Purpose'],
    [
        ['Overall satisfaction', '1-5', 'CSAT'],
        ['Likelihood to recommend', '0-10', 'NPS'],
        ['Ease of use', '1-5', 'SUS component'],
        ['Feature completeness', '1-5', 'Gap analysis'],
        ['Trust in AI guidance', '1-5', 'Trust metric'],
        ['Likelihood to use for real project', '1-5', 'Adoption intent'],
        ['Value vs. time invested', '1-5', 'ROI perception'],
    ]
)

add_heading('System Usability Scale (SUS) - Standard 10 Questions', 2)
sus_questions = [
    'I think that I would like to use this system frequently.',
    'I found the system unnecessarily complex.',
    'I thought the system was easy to use.',
    'I think that I would need the support of a technical person to use this system.',
    'I found the various functions in this system were well integrated.',
    'I thought there was too much inconsistency in this system.',
    'I would imagine that most people would learn to use this system very quickly.',
    'I found the system very cumbersome to use.',
    'I felt very confident using the system.',
    'I needed to learn a lot of things before I could get going with this system.',
]
for i, q in enumerate(sus_questions, 1):
    doc.add_paragraph(f'{i}. {q}')
doc.add_paragraph('Scale: Strongly Disagree (1) to Strongly Agree (5)')

doc.add_page_break()

# 8. Expert Review Framework
add_heading('8. Expert Review Framework', 1)

add_heading('Panel Composition', 2)
doc.add_paragraph('Recommended: 3-5 strategy experts with:')
experts = [
    'Mix of backgrounds: management consulting, corporate strategy, product leadership',
    '10+ years experience in strategy development',
    'No prior exposure to the tool',
]
for exp in experts:
    doc.add_paragraph(exp, style='List Bullet')

add_heading('Scoring Rubric', 2)

doc.add_paragraph('Dimension 1: Coaching Relevance (1-5)', style='Intense Quote')
add_table_from_data(
    ['Score', 'Description'],
    [
        ['1 - Poor', 'Coaching completely misses the point; generic or off-topic'],
        ['2 - Below Average', 'Some relevant points but largely unhelpful'],
        ['3 - Adequate', 'Reasonably relevant but not insightful'],
        ['4 - Good', 'Relevant and occasionally insightful'],
        ['5 - Excellent', 'Highly relevant, insightful, and well-timed'],
    ]
)

doc.add_paragraph('Dimension 2: Strategic Actionability (1-5)', style='Intense Quote')
add_table_from_data(
    ['Score', 'Description'],
    [
        ['1 - Poor', 'Outputs are too vague to act upon'],
        ['2 - Below Average', 'Some actionable elements but mostly abstract'],
        ['3 - Adequate', 'Actionable but requires significant interpretation'],
        ['4 - Good', 'Clear actions with reasonable specificity'],
        ['5 - Excellent', 'Highly actionable with clear next steps'],
    ]
)

doc.add_paragraph('Dimension 3: Insight Quality (1-5)', style='Intense Quote')
add_table_from_data(
    ['Score', 'Description'],
    [
        ['1 - Poor', 'No insights; purely descriptive'],
        ['2 - Below Average', 'Obvious insights only'],
        ['3 - Adequate', 'Some non-obvious connections'],
        ['4 - Good', 'Multiple meaningful insights'],
        ['5 - Excellent', 'Breakthrough insights that reframe the problem'],
    ]
)

doc.add_paragraph('Dimension 4: Stakeholder Readiness (1-5)', style='Intense Quote')
add_table_from_data(
    ['Score', 'Description'],
    [
        ['1 - Poor', 'Would embarrass the user if shared'],
        ['2 - Below Average', 'Needs significant editing before sharing'],
        ['3 - Adequate', 'Shareable with caveats'],
        ['4 - Good', 'Ready to share with minor polish'],
        ['5 - Excellent', 'Could be presented to the board immediately'],
    ]
)

doc.add_page_break()

# 9. Participant Recruitment
add_heading('9. Participant Recruitment', 1)

add_heading('Target Profile', 2)
add_table_from_data(
    ['Attribute', 'Requirement'],
    [
        ['Title', 'VP Product, Head of Strategy, Product Director, CPO'],
        ['Company size', '200+ employees (enterprise focus)'],
        ['Experience', '5+ years in product or strategy roles'],
        ['Current context', 'Actively involved in strategic initiatives'],
        ['Industry', 'Technology, Financial Services, Healthcare, B2B SaaS'],
    ]
)

add_heading('Screening Questions', 2)

doc.add_paragraph('1. Role Confirmation', style='Intense Quote')
doc.add_paragraph('What is your current job title?')
options = [
    'VP/Head/Director of Product or Strategy (qualify)',
    'Product Manager (qualify if senior)',
    'Other (screen out)',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('2. Strategic Involvement', style='Intense Quote')
doc.add_paragraph('In the past 12 months, have you been responsible for developing or significantly contributing to your company\'s product strategy?')
options = [
    'Yes, I led the process (ideal)',
    'Yes, I was a key contributor (qualify)',
    'No, I was not involved (screen out)',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('3. Company Size', style='Intense Quote')
doc.add_paragraph('Approximately how many employees does your company have?')
options = [
    '1-49 (screen out unless funded startup)',
    '50-199 (qualify with caution)',
    '200-999 (qualify)',
    '1000+ (ideal)',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

add_heading('Incentives', 2)
add_table_from_data(
    ['Test Type', 'Duration', 'Suggested Incentive'],
    [
        ['Maze Test (unmoderated)', '5-10 min', '$20-30 or prize draw'],
        ['Maze Test (longer)', '10-15 min', '$40-50'],
        ['Moderated Session', '60 min', '$150-200'],
        ['Expert Panel Review', '30 min per review', '$100-150 per review'],
    ]
)

add_heading('Sample Size Recommendations', 2)
add_table_from_data(
    ['Test Type', 'Minimum', 'Recommended', 'Notes'],
    [
        ['Maze usability tests', '10', '15-20', '85% confidence at n=15'],
        ['Moderated sessions', '5', '6-8', 'Diminishing returns after 8'],
        ['Expert reviews', '3', '5', 'For inter-rater reliability'],
    ]
)

doc.add_page_break()

# 10. Test Plan Timeline
add_heading('10. Test Plan Timeline', 1)

add_heading('6-Week Programme Overview', 2)
add_table_from_data(
    ['Week', 'Activity', 'Deliverable'],
    [
        ['1', 'Preparation: Recruit, build tests', 'Approved test plan, ready tests'],
        ['2', 'Maze Tests 1-3', 'Live tests, early insights'],
        ['3', 'Maze Tests 4-6', 'Complete Maze data'],
        ['4', 'Moderated Sessions (6-8)', 'Recordings, theme summary'],
        ['5', 'Expert Review + Analysis', 'Integrated findings'],
        ['6', 'Reporting + Planning', 'Final report, sprint plan'],
    ]
)

doc.add_page_break()

# 11. Success Metrics
add_heading('11. Success Metrics', 1)

add_heading('Quantitative Targets', 2)
add_table_from_data(
    ['Metric', 'Target', 'Source', 'Priority'],
    [
        ['Task success rate (first-click)', '≥85%', 'Maze', 'P0'],
        ['Task completion rate', '≥80%', 'Maze', 'P0'],
        ['Time-on-task (Discovery)', '<3 min', 'Maze', 'P1'],
        ['System Usability Scale (SUS)', '≥75', 'Post-session survey', 'P0'],
        ['Net Promoter Score (NPS)', '≥30', 'Post-session survey', 'P0'],
        ['Coach trust score', '≥4.0/5.0', 'Session interviews', 'P0'],
        ['Synthesis actionability', '≥4.0/5.0', 'Expert review', 'P0'],
        ['Journey completion rate', '≥70%', 'Session observation', 'P1'],
        ['Willingness to pay', '≥60% "yes"', 'Post-session survey', 'P1'],
    ]
)

add_heading('Risk Thresholds', 2)
add_table_from_data(
    ['Metric', 'Acceptable', 'Warning', 'Critical'],
    [
        ['Task success', '≥85%', '70-84%', '<70%'],
        ['SUS Score', '≥75', '60-74', '<60'],
        ['NPS', '≥30', '0-29', '<0'],
        ['Coach trust', '≥4.0', '3.0-3.9', '<3.0'],
        ['Abandonment rate', '<20%', '20-40%', '>40%'],
    ]
)

doc.add_page_break()

# 12. Budget Estimate
add_heading('12. Budget Estimate', 1)

add_heading('External Costs', 2)
add_table_from_data(
    ['Category', 'Item', 'Quantity', 'Unit Cost', 'Total'],
    [
        ['Maze Platform', 'Subscription (Pro)', '1 month', '$99', '$99'],
        ['Maze Platform', 'Panel recruitment', '100 participants', '$3-5', '$300-500'],
        ['Moderated Sessions', 'Participant incentives', '8', '$150', '$1,200'],
        ['Moderated Sessions', 'Recruitment platform', '1', '$300', '$300'],
        ['Moderated Sessions', 'Video transcription', '8 hours', '$50', '$400'],
        ['Expert Review', 'Reviewer fees', '5 x 3 reviews', '$100', '$1,500'],
    ]
)

add_heading('Total Budget Summary', 2)
add_table_from_data(
    ['Category', 'Low Estimate', 'High Estimate'],
    [
        ['External costs', '$3,800', '$5,000'],
        ['Internal time (est.)', '£7,175', '£9,000'],
        ['Contingency (15%)', '$570 + £1,076', '$750 + £1,350'],
        ['TOTAL', '$4,370 + £8,251', '$5,750 + £10,350'],
    ]
)

doc.add_page_break()

# 13. Appendices
add_heading('13. Appendices', 1)

add_heading('Appendix A: Consent Form Template', 2)
consent = doc.add_paragraph()
consent.add_run('PARTICIPANT CONSENT FORM\n\n').bold = True
consent.add_run(
    'Study Title: Product Strategy Coach User Testing\n'
    'Organisation: Frontera\n\n'
    'Purpose:\n'
    'We are conducting research to evaluate and improve our product strategy coaching tool. '
    'Your participation will help us understand how users interact with the tool.\n\n'
    'What to Expect:\n'
    '• Duration: Approximately [X] minutes\n'
    '• Activities: Using a prototype, answering questions, sharing feedback\n'
    '• Recording: This session will be recorded for research purposes\n\n'
    'Your Rights:\n'
    '• Participation is voluntary\n'
    '• You may withdraw at any time without penalty\n'
    '• You may skip any question you prefer not to answer\n'
    '• Your responses will be kept confidential\n\n'
    'By signing below, you confirm that you have read and understood this consent form '
    'and agree to participate in this research study.\n\n'
    'Signature: _______________________\n'
    'Date: ___________________________\n'
    'Name (printed): __________________'
)

add_heading('Appendix B: Session Checklist', 2)

doc.add_paragraph('Pre-Session Checklist', style='Intense Quote')
pre_session = [
    'Prototype URL confirmed working',
    'Recording software tested',
    'Consent form sent and signed',
    'Participant background reviewed',
    'Observation template ready',
    'Backup scenarios prepared',
    'Calendar invite sent with joining link',
    'Incentive payment prepared',
]
for item in pre_session:
    doc.add_paragraph(f'☐ {item}')

doc.add_paragraph('During-Session Checklist', style='Intense Quote')
during_session = [
    'Welcome and introductions',
    'Confirm consent to record',
    'Start recording',
    'Context-setting questions',
    'First impressions task',
    'Discovery phase walkthrough',
    'Research phase walkthrough',
    'Synthesis review',
    'Debrief interview',
    'Thank participant',
    'Stop recording',
]
for item in during_session:
    doc.add_paragraph(f'☐ {item}')

doc.add_paragraph('Post-Session Checklist', style='Intense Quote')
post_session = [
    'Complete observation notes within 24 hours',
    'Tag key moments in recording',
    'Update participant tracker',
    'Send thank-you email',
    'Process incentive payment',
    'Share preliminary insights with team',
]
for item in post_session:
    doc.add_paragraph(f'☐ {item}')

# Document control
doc.add_page_break()
add_heading('Document Control', 1)
add_table_from_data(
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ['1.0', 'February 2026', 'Frontera Product Team', 'Initial draft'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('— End of Document —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
output_path = r'c:\Users\deeks\frontera-platform\Background\Product_Strategy_Coach_User_Testing_Proposal.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
