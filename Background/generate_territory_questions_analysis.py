"""
Generate a Word document analysing the Territory Research Questions
in the Product Strategy Coach.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

NAVY = RGBColor(0x1a, 0x1f, 0x3a)
GOLD = RGBColor(0xfb, 0xbf, 0x24)
CYAN = RGBColor(0x08, 0x91, 0xb2)
EMERALD = RGBColor(0x10, 0xb9, 0x81)
RED = RGBColor(0xef, 0x44, 0x44)
AMBER = RGBColor(0xd9, 0x77, 0x06)


def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_heading(text, level=1):
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = NAVY
    return heading


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '1a1f3a')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
    doc.add_paragraph()
    return table


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def quote(text):
    doc.add_paragraph(text, style='Intense Quote')


def bold_para(label, text):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)
    return p


# ============================================================================
# TITLE PAGE
# ============================================================================
title = doc.add_heading('Territory Research Questions', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY

subtitle = doc.add_paragraph('Critical Analysis & Optimisation Review')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(22)
subtitle.runs[0].font.color.rgb = CYAN

doc.add_paragraph()

tagline = doc.add_paragraph('Product Strategy Coach — Terrain Mapping Phase')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.runs[0].font.size = Pt(14)
tagline.runs[0].italic = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Date: February 2026\n')
info.add_run('Purpose: Question inventory, strategic rationale, and reduction impact analysis\n')
info.add_run('Current: 36 questions (4 per area × 9 areas × 3 territories)\n')
info.add_run('Proposed: 27 questions (3 per area × 9 areas × 3 territories)')

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
add_heading('Table of Contents', 1)
toc = [
    ('1.', 'Executive Summary'),
    ('2.', 'Question Inventory Overview'),
    ('3.', 'Company Territory — Full Analysis'),
    ('4.', 'Customer Territory — Full Analysis'),
    ('5.', 'Competitor Territory — Full Analysis'),
    ('6.', 'How Questions Feed Synthesis'),
    ('7.', 'Reduction Impact Analysis'),
    ('8.', 'Recommended 3-Question Sets'),
    ('9.', 'Risk Assessment & Mitigations'),
]
for num, item in toc:
    p = doc.add_paragraph()
    p.add_run(f'{num} ').bold = True
    p.add_run(item)

doc.add_page_break()

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================
add_heading('1. Executive Summary', 1)

add_heading('Current State', 2)
add_table(
    ['Dimension', 'Value'],
    [
        ['Total Territories', '3 (Company, Customer, Competitor)'],
        ['Research Areas per Territory', '3'],
        ['Questions per Area', '4'],
        ['Total Questions', '36'],
        ['Question Type', 'All open-ended text (textarea)'],
        ['Input Format', 'Free-text, no character limit'],
        ['Completion Requirement', 'All 4 questions to mark area as "Mapped"'],
        ['Minimum for Synthesis', '4 of 9 areas with at least 1 written response'],
    ]
)

add_heading('Key Findings', 2)
bullets([
    'All 36 questions are open-ended text inputs — no variation in question type (rating, multiple choice, ranking)',
    'Each area has exactly 4 questions — a uniform structure that does not reflect varying complexity across areas',
    'Questions follow a consistent pattern: (1) current state, (2) differentiation/depth, (3) constraints/threats, (4) trajectory/evolution',
    'The 4th question in each area consistently asks about trends or forward-looking dynamics — often the hardest to answer without preparation',
    'Several questions overlap across territories (e.g., market position appears in Company and Competitor)',
    'Synthesis requires minimum 4 areas, but quality improves dramatically with balanced coverage across all 3 territories',
    'Reducing to 3 questions per area removes 9 questions (25% reduction) and cuts estimated completion time by ~20-30 minutes',
])

add_heading('Recommendation', 2)
doc.add_paragraph(
    'Reduce to 3 questions per area by merging or removing the lowest-signal question in each set. '
    'This reduces total questions from 36 to 27, cuts user effort by ~25%, and preserves 85-90% of synthesis '
    'input quality. The 4th question in most areas either overlaps with other areas or asks for '
    'forward-looking analysis that is better surfaced through coaching conversation than structured input.'
)

doc.add_page_break()

# ============================================================================
# 2. QUESTION INVENTORY OVERVIEW
# ============================================================================
add_heading('2. Question Inventory Overview', 1)

add_heading('Complete Question Count', 2)
add_table(
    ['Territory', 'Area 1', 'Area 2', 'Area 3', 'Total'],
    [
        ['Company', '4 questions', '4 questions', '4 questions', '12'],
        ['Customer', '4 questions', '4 questions', '4 questions', '12'],
        ['Competitor', '4 questions', '4 questions', '4 questions', '12'],
        ['TOTAL', '', '', '', '36'],
    ]
)

add_heading('Question Type Analysis', 2)
add_table(
    ['Question Type', 'Count', '% of Total', 'Assessment'],
    [
        ['Open-ended text (textarea)', '36', '100%', 'No variation — cognitive fatigue risk'],
        ['Rating/scale', '0', '0%', 'Missing — could reduce effort for some questions'],
        ['Multiple choice', '0', '0%', 'Missing — could accelerate known-answer questions'],
        ['Ranking/prioritisation', '0', '0%', 'Missing — could surface priorities faster'],
    ]
)

add_heading('Question Pattern Analysis', 2)
doc.add_paragraph('Each research area follows a consistent 4-question pattern:')
add_table(
    ['Position', 'Typical Function', 'Cognitive Load', 'Synthesis Value'],
    [
        ['Q1', 'Landscape scan — "What/Who are your...?"', 'Low-Medium', 'High — establishes foundation'],
        ['Q2', 'Differentiation — "What makes them/you different?"', 'Medium', 'High — reveals unique positioning'],
        ['Q3', 'Constraints/gaps — "What limits or threatens?"', 'Medium-High', 'High — identifies vulnerabilities'],
        ['Q4', 'Trajectory — "How is this evolving/trending?"', 'High', 'Medium — forward-looking, often speculative'],
    ]
)

doc.add_paragraph(
    'The consistent pattern helps users develop a rhythm, but the uniform 4-question structure '
    'does not account for the reality that some areas need deeper exploration than others.'
)

add_heading('Estimated Completion Time', 2)
add_table(
    ['Scenario', 'Per Question', 'Per Area', 'Per Territory', 'Full Completion'],
    [
        ['Minimal (brief answers)', '2 min', '8 min', '24 min', '72 min (~1.2 hrs)'],
        ['Moderate (thoughtful)', '4 min', '16 min', '48 min', '144 min (~2.4 hrs)'],
        ['Deep (comprehensive)', '6 min', '24 min', '72 min', '216 min (~3.6 hrs)'],
    ]
)

doc.add_paragraph(
    'At 4 questions per area, even moderate effort requires ~2.5 hours of focused strategic writing. '
    'This is a significant cognitive investment that risks abandonment in the Research phase.'
)

doc.add_page_break()

# ============================================================================
# 3. COMPANY TERRITORY
# ============================================================================
add_heading('3. Company Territory — Full Analysis', 1)

# Area 1
add_heading('Area 1: Core Capabilities & Constraints', 2)
quote('What are your organization\'s unique strengths, and what fundamental constraints shape your strategic options?')

add_table(
    ['#', 'Question', 'Type', 'Synthesis Role'],
    [
        ['Q1', 'What are your organization\'s core competencies and differentiating capabilities?', 'Landscape', 'Identifies "How to Win" candidates — what the org does uniquely well'],
        ['Q2', 'What key resources (technical, human, IP) do you control that competitors don\'t?', 'Differentiation', 'Feeds Capability Fit scoring — proprietary advantages'],
        ['Q3', 'What structural constraints limit your strategic freedom (legacy systems, contracts, regulations)?', 'Constraints', 'Identifies execution barriers for Assumptions (WWHBT)'],
        ['Q4', 'Which capabilities are table stakes vs. truly differentiated in your market?', 'Evaluation', 'Refines competitive advantage scoring — separates signal from noise'],
    ]
)

bold_para('Strategic Purpose: ', 'Establishes the internal capability foundation for "How to Win" strategies. '
          'Without this area, synthesis cannot assess Capability Fit (35% of opportunity scoring).')
bold_para('Synthesis Connection: ', 'Directly feeds Capability Fit scores (1-10) and "How to Win" in Playing to Win framework. '
          'Evidence from Q1-Q2 frequently cited in opportunity cards.')
bold_para('Overlap Risk: ', 'Q4 (table stakes vs. differentiated) overlaps conceptually with Q1 (core competencies). '
          'Both ask "what makes us special?" from slightly different angles.')
bold_para('Reduction Candidate: ', 'Q4. Merge into Q1 by reframing as "What are your truly differentiated capabilities '
          '(beyond table stakes)?" The table-stakes/differentiation distinction can emerge from coaching conversation.')

# Area 2
add_heading('Area 2: Resource Reality', 2)
quote('What team, technology, and funding realities will enable or constrain your strategy execution?')

add_table(
    ['#', 'Question', 'Type', 'Synthesis Role'],
    [
        ['Q1', 'What is the current composition and skill distribution of your team?', 'Landscape', 'Identifies capability gaps for "Capabilities Required" in PTW'],
        ['Q2', 'What technology stack and infrastructure do you have in place?', 'Asset Inventory', 'Feeds Capability Fit — technical enablement or constraint'],
        ['Q3', 'What funding runway and burn rate define your strategic timeline?', 'Constraints', 'Sets time horizon for strategic bets — urgency signal'],
        ['Q4', 'What hiring constraints or talent gaps could impact execution?', 'Risk', 'Feeds Assumptions — "We can hire X in Y timeframe"'],
    ]
)

bold_para('Strategic Purpose: ', 'Grounds strategy in execution reality. Prevents synthesis from generating '
          'aspirational strategies the company cannot resource.')
bold_para('Synthesis Connection: ', 'Feeds "Capabilities Required" and "Management Systems" fields in Playing to Win. '
          'Also directly influences Capability Fit scoring via resource constraints.')
bold_para('Overlap Risk: ', 'Q1 (team composition) and Q4 (hiring/talent gaps) are closely related — '
          'knowing the team composition naturally surfaces gaps.')
bold_para('Reduction Candidate: ', 'Q4. Merge talent gap identification into Q1 by reframing as '
          '"What is your team composition, and what critical skill gaps exist?"')

# Area 3
add_heading('Area 3: Product Portfolio & Market Position', 2)
quote('How do your current products perform in the market, and what does your portfolio reveal about strategic direction?')

add_table(
    ['#', 'Question', 'Type', 'Synthesis Role'],
    [
        ['Q1', 'What products/services comprise your current portfolio and how do they perform?', 'Landscape', 'Establishes portfolio baseline for "Where to Play" decisions'],
        ['Q2', 'Which products are growth drivers vs. legacy offerings?', 'Evaluation', 'Identifies invest/harvest decisions — directly maps to 2x2 quadrants'],
        ['Q3', 'What is your current market position and competitive standing?', 'Positioning', 'Feeds Market Attractiveness and Competitive Advantage scores'],
        ['Q4', 'What gaps exist between your current portfolio and market opportunities?', 'Trajectory', 'Identifies "Where to Play" expansion candidates'],
    ]
)

bold_para('Strategic Purpose: ', 'Maps existing portfolio health and positioning. Essential for determining '
          'which quadrant (INVEST/EXPLORE/HARVEST/DIVEST) applies to each opportunity.')
bold_para('Synthesis Connection: ', 'Q2 directly informs the 2×2 Strategic Opportunity Map. '
          'Q3 overlaps heavily with Competitor territory — competitive standing is better assessed with competitor context.')
bold_para('Overlap Risk: ', 'Q3 (market position/competitive standing) duplicates analysis done in Competitor Territory. '
          'The user is asked about competitive position here AND about direct competitors there.')
bold_para('Reduction Candidate: ', 'Q3. Competitive standing is better assessed in Competitor Territory where '
          'the user has competitor-specific context. Reframe Q4 to include positioning: '
          '"What gaps exist between your portfolio and market opportunities, and how does your current position address them?"')

doc.add_page_break()

# ============================================================================
# 4. CUSTOMER TERRITORY
# ============================================================================
add_heading('4. Customer Territory — Full Analysis', 1)

# Area 1
add_heading('Area 1: Customer Segmentation & Behaviours', 2)
quote('Who are your customers, how do they behave, and what drives their decisions?')

add_table(
    ['#', 'Question', 'Type', 'Synthesis Role'],
    [
        ['Q1', 'What are your primary customer segments, and how do they differ in needs, behaviors, and value?', 'Landscape', 'Defines "Where to Play" segments — most critical input for synthesis'],
        ['Q2', 'How do customers currently discover, evaluate, and purchase solutions in your category?', 'Journey', 'Feeds "How to Win" via channel/GTM strategy'],
        ['Q3', 'What decision-making criteria matter most to each segment (price, features, trust, speed)?', 'Drivers', 'Directly feeds "How to Win" — what competitive levers matter'],
        ['Q4', 'Which customer segments are growing, declining, or emerging in your market?', 'Trajectory', 'Feeds Market Attractiveness scoring — growth potential'],
    ]
)

bold_para('Strategic Purpose: ', 'The most important area for synthesis. Customer segments define "Where to Play" '
          'and decision criteria define "How to Win." Without this, synthesis lacks its primary input.')
bold_para('Synthesis Connection: ', 'Q1 directly drives segment selection in every strategic opportunity. '
          'Q3 defines the competitive battleground. Q4 informs Market Attractiveness (40% of scoring).')
bold_para('Overlap Risk: ', 'Q4 (segment growth trends) overlaps with Customer Area 3 (Market Dynamics) '
          'which also asks about evolving expectations and trends.')
bold_para('Reduction Candidate: ', 'Q2 (purchase journey). While useful, the discovery/evaluation/purchase '
          'journey is tactical rather than strategic. It is less relevant to Playing to Win synthesis '
          'and better explored through coaching conversation. However, this is a difficult cut — '
          'Q2 informs GTM strategy. Alternative: cut Q4 and cover segment trajectory in Area 3.')

# Area 2
add_heading('Area 2: Unmet Needs & Pain Points', 2)
quote('What problems do customers face that current solutions fail to address adequately?')

add_table(
    ['#', 'Question', 'Type', 'Synthesis Role'],
    [
        ['Q1', 'What are the most significant pain points customers experience with existing solutions (including yours)?', 'Landscape', 'Identifies opportunity spaces for "Where to Play"'],
        ['Q2', 'What jobs-to-be-done are customers hiring products for, and where do current solutions fall short?', 'JTBD Framework', 'Core input for "How to Win" — solving underserved jobs'],
        ['Q3', 'What workarounds, hacks, or compromises do customers make to get their jobs done?', 'Signal Detection', 'Reveals latent demand — highest-signal indicator of market opportunity'],
        ['Q4', 'What emerging needs or latent desires are customers beginning to express?', 'Trajectory', 'Forward-looking opportunity identification'],
    ]
)

bold_para('Strategic Purpose: ', 'Identifies the opportunity space. Pain points and workarounds are the '
          'strongest signals of market opportunity and defensible positioning.')
bold_para('Synthesis Connection: ', 'Q1-Q3 are the primary evidence sources for strategic opportunities. '
          'Claude cites these responses most frequently in evidence trails. '
          'Q3 (workarounds) is particularly high-signal — it reveals where demand exists but supply doesn\'t.')
bold_para('Overlap Risk: ', 'Q1 (pain points) and Q2 (JTBD shortfalls) are closely related frameworks '
          'for the same underlying question: "Where do solutions fail?" However, they elicit different '
          'types of responses — Q1 surfaces problems, Q2 surfaces unmet functional needs.')
bold_para('Reduction Candidate: ', 'Q4 (emerging needs). This is speculative and overlaps with '
          'Customer Area 3 (Market Dynamics Q1: evolving expectations). The remaining Q1-Q3 form a '
          'strong problem-discovery triad: problems → jobs → workarounds.')

# Area 3
add_heading('Area 3: Market Dynamics & Customer Evolution', 2)
quote('How are customer expectations, behaviors, and the competitive landscape changing over time?')

add_table(
    ['#', 'Question', 'Type', 'Synthesis Role'],
    [
        ['Q1', 'How have customer expectations evolved in the past 2-3 years, and what trends are accelerating?', 'Trajectory', 'Feeds Market Attractiveness — trend alignment'],
        ['Q2', 'What new alternatives or substitutes are customers considering that didn\'t exist before?', 'Threats', 'Overlaps with Competitor Territory (substitutes & adjacent threats)'],
        ['Q3', 'How are customer acquisition costs, retention rates, and lifetime value trending?', 'Economics', 'Feeds Assumptions — unit economics viability'],
        ['Q4', 'What external forces (technology, regulation, economics, culture) are reshaping customer needs?', 'Macro Forces', 'Overlaps with Competitor Area 3 (market forces)'],
    ]
)

bold_para('Strategic Purpose: ', 'Captures the dynamic dimension — how the landscape is shifting. '
          'Important for validating whether today\'s opportunities will persist.')
bold_para('Synthesis Connection: ', 'Q1 informs Market Attractiveness trends. Q3 feeds Assumptions '
          '("What Would Have to Be True" about unit economics). Q2 and Q4 provide competitive context.')
bold_para('Overlap Risk: ', 'This area has the HIGHEST overlap with other territories. '
          'Q2 (new alternatives) is nearly identical to Competitor Area 2 Q1 (non-traditional solutions). '
          'Q4 (external forces) duplicates Competitor Area 3 Q1 (macro trends). Users answering both '
          'areas will experience significant repetition.')
bold_para('Reduction Candidate: ', 'Q2 (new alternatives). This is fully covered in Competitor Territory Area 2. '
          'Removing it eliminates the most obvious cross-territory duplication. '
          'Alternative: remove Q4 (external forces) which duplicates Competitor Area 3.')

doc.add_page_break()

# ============================================================================
# 5. COMPETITOR TERRITORY
# ============================================================================
add_heading('5. Competitor Territory — Full Analysis', 1)

# Area 1
add_heading('Area 1: Direct Competitor Landscape', 2)
quote('Who are your primary competitors, and how do they compete for the same customers?')

add_table(
    ['#', 'Question', 'Type', 'Synthesis Role'],
    [
        ['Q1', 'Who are your top 3-5 direct competitors, and what makes them your primary competition?', 'Landscape', 'Establishes competitive set for Competitive Advantage scoring'],
        ['Q2', 'What are each competitor\'s core value propositions, and how do they differentiate?', 'Differentiation', 'Feeds "How to Win" — understanding what you compete against'],
        ['Q3', 'Where do competitors have clear advantages over your current offering?', 'Gap Analysis', 'Identifies vulnerability — feeds Strategic Tensions'],
        ['Q4', 'What competitive moves or announcements have you observed in the past 12 months?', 'Intelligence', 'Forward-looking competitor trajectory — feeds urgency signals'],
    ]
)

bold_para('Strategic Purpose: ', 'Maps the competitive battleground. Without knowing who you compete against '
          'and their strengths, "How to Win" strategies are blind.')
bold_para('Synthesis Connection: ', 'Q1-Q3 directly feed Competitive Advantage scoring (25% of overall score). '
          'Q3 (competitor advantages) is a primary source for Strategic Tensions. '
          'Q2 (competitor value props) helps Claude assess whether a proposed "How to Win" is differentiated.')
bold_para('Overlap Risk: ', 'Q4 (recent moves) is time-sensitive intelligence that may not be available '
          'or may be speculative. It also overlaps with the forward-looking nature of Area 3 (market forces).')
bold_para('Reduction Candidate: ', 'Q4 (recent moves). While useful colour, it is the most perishable data '
          'and the most likely to elicit "I\'m not sure" responses. Q1-Q3 provide the structural competitive '
          'understanding needed for synthesis.')

# Area 2
add_heading('Area 2: Substitute & Adjacent Threats', 2)
quote('What alternative solutions or emerging players could capture your customers\' attention?')

add_table(
    ['#', 'Question', 'Type', 'Synthesis Role'],
    [
        ['Q1', 'What non-traditional solutions do customers use to solve the same problems you address?', 'Substitutes', 'Broadens competitive frame — "Where NOT to Play" signals'],
        ['Q2', 'What adjacent products or services are expanding into your market space?', 'Adjacent Threats', 'Identifies emerging competitive pressure'],
        ['Q3', 'What emerging startups or disruptors are gaining traction with your target customers?', 'Disruptors', 'Forward threat identification — urgency for strategic bets'],
        ['Q4', 'How might technology shifts (AI, automation, platforms) create new competitive threats?', 'Technology Risks', 'Feeds Assumptions — technology disruption risk'],
    ]
)

bold_para('Strategic Purpose: ', 'Expands the competitive lens beyond direct competitors to include '
          'substitutes, adjacency moves, and disruptors. Critical for avoiding strategic blindspots.')
bold_para('Synthesis Connection: ', 'Q1 (substitutes) broadens "Where to Play" by revealing what customers '
          'actually use. Q2-Q3 feed Strategic Tensions by identifying threats to proposed strategies. '
          'Q4 feeds Assumptions with technology disruption scenarios.')
bold_para('Overlap Risk: ', 'Q1 (non-traditional solutions) overlaps with Customer Area 3 Q2 (new alternatives). '
          'Q4 (technology shifts) is broad and speculative, overlapping with Area 3 (market forces).')
bold_para('Reduction Candidate: ', 'Q4 (technology shifts). This is the most speculative question and overlaps with '
          'Area 3 Q1 (macro trends). The remaining Q1-Q3 cover the substitute/adjacent/disruptor triad '
          'which is more specific and actionable.')

# Area 3
add_heading('Area 3: Market Forces & Dynamics', 2)
quote('What broader market trends and forces are reshaping the competitive landscape?')

add_table(
    ['#', 'Question', 'Type', 'Synthesis Role'],
    [
        ['Q1', 'What macroeconomic, regulatory, or industry trends are most impacting your market?', 'Macro Trends', 'Feeds Market Attractiveness — external forces'],
        ['Q2', 'How is the overall market size and growth trajectory evolving?', 'Market Sizing', 'Direct input to Market Attractiveness scoring (40% weight)'],
        ['Q3', 'What barriers to entry exist, and are they strengthening or weakening?', 'Moat Analysis', 'Feeds Competitive Advantage — defensibility assessment'],
        ['Q4', 'What emerging customer expectations or behaviors are changing competitive dynamics?', 'Trajectory', 'Overlaps with Customer Area 3 Q1 (evolving expectations)'],
    ]
)

bold_para('Strategic Purpose: ', 'Provides the macro context for strategy. Market size, growth, and barriers '
          'to entry are fundamental inputs for investment decisions.')
bold_para('Synthesis Connection: ', 'Q2 (market size/growth) is the most direct input to Market Attractiveness. '
          'Q3 (barriers) feeds Competitive Advantage durability. Both are critical for accurate '
          'scoring in the 2×2 Strategic Opportunity Map.')
bold_para('Overlap Risk: ', 'Q4 (changing customer expectations) directly duplicates Customer Area 3 Q1 '
          '(evolved customer expectations). This is the clearest cross-territory overlap.')
bold_para('Reduction Candidate: ', 'Q4 (changing customer expectations). This belongs in Customer Territory, '
          'not Competitor Territory. Its removal eliminates the most obvious duplication and '
          'keeps the area focused on market forces rather than customer behaviour.')

doc.add_page_break()

# ============================================================================
# 6. HOW QUESTIONS FEED SYNTHESIS
# ============================================================================
add_heading('6. How Questions Feed Synthesis', 1)

add_heading('Data Flow', 2)
doc.add_paragraph(
    'All territory responses are passed to Claude as a single context block. Claude generates '
    '3-5 strategic opportunities using the Playing to Win framework. Each opportunity must cite '
    'at least 2 evidence quotes from the research.'
)

add_heading('Playing to Win Mapping', 2)
add_table(
    ['PTW Element', 'Primary Question Sources', 'Territory'],
    [
        ['Where to Play', 'Customer Area 1 Q1 (segments), Customer Area 2 Q1 (pain points), Competitor Area 2 Q1 (substitutes)', 'Customer + Competitor'],
        ['How to Win', 'Company Area 1 Q1-Q2 (capabilities), Customer Area 1 Q3 (decision criteria), Competitor Area 1 Q2 (competitor value props)', 'All three'],
        ['Capabilities Required', 'Company Area 2 Q1-Q2 (team + tech), Company Area 1 Q3 (constraints)', 'Company'],
        ['Winning Aspiration', 'Company Area 3 Q2 (growth drivers), Customer Area 2 Q2 (JTBD)', 'Company + Customer'],
        ['WWHBT Assumptions', 'Company Area 2 Q3 (funding), Customer Area 3 Q3 (unit economics), Competitor Area 3 Q3 (barriers)', 'All three'],
    ]
)

add_heading('Scoring Input Mapping', 2)
add_table(
    ['Score Dimension', 'Weight', 'Primary Question Sources'],
    [
        ['Market Attractiveness', '40%', 'Customer Area 1 Q4 (segment growth), Customer Area 3 Q1 (trends), Competitor Area 3 Q2 (market size)'],
        ['Capability Fit', '35%', 'Company Area 1 Q1-Q2 (capabilities), Company Area 2 Q1-Q2 (team/tech), Company Area 3 Q1 (portfolio)'],
        ['Competitive Advantage', '25%', 'Competitor Area 1 Q2-Q3 (competitor strengths/weaknesses), Competitor Area 3 Q3 (barriers)'],
    ]
)

add_heading('Evidence Citation Frequency', 2)
doc.add_paragraph('Based on the synthesis prompt structure, Claude is most likely to cite:')
add_table(
    ['Source Area', 'Expected Citation Frequency', 'Reason'],
    [
        ['Customer: Unmet Needs (Area 2)', 'Very High', 'Pain points and workarounds are the richest opportunity signal'],
        ['Company: Core Capabilities (Area 1)', 'High', 'Capability evidence is needed for every "How to Win"'],
        ['Competitor: Direct Landscape (Area 1)', 'High', 'Competitive positioning evidence validates differentiation'],
        ['Customer: Segmentation (Area 1)', 'High', '"Where to Play" requires segment definition'],
        ['Company: Portfolio (Area 3)', 'Medium', 'Portfolio health informs invest/harvest decisions'],
        ['Competitor: Market Forces (Area 3)', 'Medium', 'Market sizing feeds attractiveness but is less quotable'],
        ['Company: Resource Reality (Area 2)', 'Medium-Low', 'Execution constraints cited in assumptions, not opportunities'],
        ['Customer: Market Dynamics (Area 3)', 'Low', 'Trend data is contextual, rarely cited as specific evidence'],
        ['Competitor: Substitutes (Area 2)', 'Low', 'Substitutes broaden frame but are less directly cited'],
    ]
)

doc.add_page_break()

# ============================================================================
# 7. REDUCTION IMPACT ANALYSIS
# ============================================================================
add_heading('7. Reduction Impact Analysis', 1)

add_heading('Impact Summary: 4 Questions → 3 Questions Per Area', 2)
add_table(
    ['Metric', 'Current (4 per area)', 'Proposed (3 per area)', 'Change'],
    [
        ['Total Questions', '36', '27', '-9 (25% reduction)'],
        ['Est. Completion Time (moderate)', '~2.5 hours', '~1.75 hours', '-45 minutes'],
        ['Est. Completion Time (minimal)', '~1.2 hours', '~50 minutes', '-22 minutes'],
        ['Synthesis Input Coverage', '100%', '~85-90%', 'Minor signal loss'],
        ['Cross-Territory Overlap', 'Significant (4-5 pairs)', 'Minimal (1-2 pairs)', 'Major improvement'],
        ['User Cognitive Load', 'High', 'Moderate', 'Significant improvement'],
        ['Abandonment Risk', 'Medium-High', 'Low-Medium', 'Significant improvement'],
    ]
)

add_heading('What You Gain', 2)
bullets([
    'Faster time-to-synthesis: Users reach the high-value synthesis phase sooner',
    'Reduced abandonment: 9 fewer open-ended questions reduces fatigue and dropout risk',
    'Less repetition: Removing cross-territory overlaps eliminates "didn\'t I already answer this?" frustration',
    'Sharper focus: Each remaining question carries more weight, encouraging deeper answers',
    'Preserved methodology: The 3 territories × 3 areas × 3 questions structure remains clean and balanced',
])

add_heading('What You Lose', 2)
bullets([
    'Forward-looking depth: Most removed questions are trajectory/trend questions (Q4 pattern)',
    'Some nuance in capability differentiation (Company Area 1 Q4: table stakes vs. differentiated)',
    'Explicit talent gap identification (Company Area 2 Q4) — though absorbed into team composition',
    'Purchase journey detail (Customer Area 1 Q2 option) — tactical GTM insight',
    'Technology disruption prompting (Competitor Area 2 Q4) — speculative but occasionally valuable',
])

add_heading('Net Assessment', 2)
doc.add_paragraph(
    'The reduction is strongly positive. The 9 removed questions are consistently the lowest-signal '
    'items in each set: either speculative forward-looking questions, cross-territory duplicates, or '
    'questions whose answers naturally emerge from the remaining three. The synthesis engine does not '
    'require 4 responses per area — it performs fuzzy matching on response content, so 3 substantive '
    'responses provide nearly equivalent evidence for opportunity generation.'
)

add_heading('Impact on Synthesis Quality', 2)
add_table(
    ['Synthesis Element', 'Impact of Reduction', 'Severity'],
    [
        ['Strategic Opportunities (3-5)', 'Negligible — primary inputs (segments, capabilities, competitors) preserved', 'Low'],
        ['Evidence Trail Quality', 'Minor — slightly fewer quotable responses, but higher quality per response', 'Low'],
        ['Market Attractiveness Scores', 'Minor — trend data reduced but market size/growth questions retained', 'Low-Medium'],
        ['Capability Fit Scores', 'Negligible — core capability and resource questions fully preserved', 'Low'],
        ['Competitive Advantage Scores', 'Negligible — direct competitor analysis fully preserved', 'Low'],
        ['Strategic Tensions', 'Minor — fewer cross-referencing data points, but key tensions still emerge', 'Low-Medium'],
        ['WWHBT Assumptions', 'Minor — some speculative trajectory inputs removed, but testable assumptions still generated', 'Low'],
        ['Overall Synthesis Confidence', 'Marginal reduction from "High" threshold (fewer evidence points per opportunity)', 'Low'],
    ]
)

doc.add_page_break()

# ============================================================================
# 8. RECOMMENDED 3-QUESTION SETS
# ============================================================================
add_heading('8. Recommended 3-Question Sets', 1)

doc.add_paragraph(
    'The following recommendations remove the lowest-signal question from each area. '
    'Where a question is removed, its key intent is absorbed into a reframed remaining question.'
)

# Company
add_heading('Company Territory (Revised)', 2)

quote('Area 1: Core Capabilities & Constraints — Remove Q4')
add_table(
    ['#', 'Revised Question', 'Change'],
    [
        ['Q1', 'What are your organization\'s truly differentiated capabilities — beyond table stakes?', 'Merged original Q1 + Q4. Adds "beyond table stakes" filter.'],
        ['Q2', 'What key resources (technical, human, IP) do you control that competitors don\'t?', 'Unchanged'],
        ['Q3', 'What structural constraints limit your strategic freedom (legacy systems, contracts, regulations)?', 'Unchanged'],
    ]
)
bold_para('Removed: ', '"Which capabilities are table stakes vs. truly differentiated?" — absorbed into Q1 reframe.')

quote('Area 2: Resource Reality — Remove Q4')
add_table(
    ['#', 'Revised Question', 'Change'],
    [
        ['Q1', 'What is the current composition of your team, and what critical skill gaps exist?', 'Merged original Q1 + Q4. Adds gap identification.'],
        ['Q2', 'What technology stack and infrastructure do you have in place?', 'Unchanged'],
        ['Q3', 'What funding runway and burn rate define your strategic timeline?', 'Unchanged'],
    ]
)
bold_para('Removed: ', '"What hiring constraints or talent gaps could impact execution?" — absorbed into Q1 reframe.')

quote('Area 3: Product Portfolio & Market Position — Remove Q3')
add_table(
    ['#', 'Revised Question', 'Change'],
    [
        ['Q1', 'What products/services comprise your current portfolio and how do they perform?', 'Unchanged'],
        ['Q2', 'Which products are growth drivers vs. legacy offerings?', 'Unchanged'],
        ['Q3', 'What gaps exist between your current portfolio and market opportunities?', 'Moved from original Q4 position'],
    ]
)
bold_para('Removed: ', '"What is your current market position and competitive standing?" — better covered in Competitor Territory.')

doc.add_page_break()

# Customer
add_heading('Customer Territory (Revised)', 2)

quote('Area 1: Customer Segmentation & Behaviours — Remove Q4')
add_table(
    ['#', 'Revised Question', 'Change'],
    [
        ['Q1', 'What are your primary customer segments, and how do they differ in needs, behaviors, and value?', 'Unchanged'],
        ['Q2', 'How do customers currently discover, evaluate, and purchase solutions in your category?', 'Unchanged'],
        ['Q3', 'What decision-making criteria matter most to each segment (price, features, trust, speed)?', 'Unchanged'],
    ]
)
bold_para('Removed: ', '"Which customer segments are growing, declining, or emerging?" — covered in Area 3 (market dynamics & trends).')

quote('Area 2: Unmet Needs & Pain Points — Remove Q4')
add_table(
    ['#', 'Revised Question', 'Change'],
    [
        ['Q1', 'What are the most significant pain points customers experience with existing solutions (including yours)?', 'Unchanged'],
        ['Q2', 'What jobs-to-be-done are customers hiring products for, and where do current solutions fall short?', 'Unchanged'],
        ['Q3', 'What workarounds, hacks, or compromises do customers make to get their jobs done?', 'Unchanged'],
    ]
)
bold_para('Removed: ', '"What emerging needs or latent desires are customers beginning to express?" — speculative; covered by Area 3.')

quote('Area 3: Market Dynamics & Customer Evolution — Remove Q2')
add_table(
    ['#', 'Revised Question', 'Change'],
    [
        ['Q1', 'How have customer expectations evolved in the past 2-3 years, and what trends are accelerating?', 'Unchanged'],
        ['Q2', 'How are customer acquisition costs, retention rates, and lifetime value trending?', 'Moved from original Q3 position'],
        ['Q3', 'What external forces (technology, regulation, economics, culture) are reshaping customer needs?', 'Moved from original Q4 position'],
    ]
)
bold_para('Removed: ', '"What new alternatives or substitutes are customers considering?" — fully covered in Competitor Territory Area 2.')

doc.add_page_break()

# Competitor
add_heading('Competitor Territory (Revised)', 2)

quote('Area 1: Direct Competitor Landscape — Remove Q4')
add_table(
    ['#', 'Revised Question', 'Change'],
    [
        ['Q1', 'Who are your top 3-5 direct competitors, and what makes them your primary competition?', 'Unchanged'],
        ['Q2', 'What are each competitor\'s core value propositions, and how do they differentiate?', 'Unchanged'],
        ['Q3', 'Where do competitors have clear advantages over your current offering?', 'Unchanged'],
    ]
)
bold_para('Removed: ', '"What competitive moves or announcements have you observed in the past 12 months?" — perishable intelligence; lower synthesis value.')

quote('Area 2: Substitute & Adjacent Threats — Remove Q4')
add_table(
    ['#', 'Revised Question', 'Change'],
    [
        ['Q1', 'What non-traditional solutions do customers use to solve the same problems you address?', 'Unchanged'],
        ['Q2', 'What adjacent products or services are expanding into your market space?', 'Unchanged'],
        ['Q3', 'What emerging startups or disruptors are gaining traction with your target customers?', 'Unchanged'],
    ]
)
bold_para('Removed: ', '"How might technology shifts create new competitive threats?" — speculative; overlaps with Area 3.')

quote('Area 3: Market Forces & Dynamics — Remove Q4')
add_table(
    ['#', 'Revised Question', 'Change'],
    [
        ['Q1', 'What macroeconomic, regulatory, or industry trends are most impacting your market?', 'Unchanged'],
        ['Q2', 'How is the overall market size and growth trajectory evolving?', 'Unchanged'],
        ['Q3', 'What barriers to entry exist, and are they strengthening or weakening?', 'Unchanged'],
    ]
)
bold_para('Removed: ', '"What emerging customer expectations or behaviors are changing competitive dynamics?" — this belongs in Customer Territory, not here.')

doc.add_page_break()

# ============================================================================
# 9. RISK ASSESSMENT & MITIGATIONS
# ============================================================================
add_heading('9. Risk Assessment & Mitigations', 1)

add_heading('Risk 1: Loss of Forward-Looking Signal', 2)
bold_para('Risk: ', 'Removing trajectory questions (Q4 pattern) reduces the forward-looking dimension of research.')
bold_para('Severity: ', 'Low-Medium')
bold_para('Mitigation: ', 'The AI coach naturally probes for trends and trajectory during conversation. '
          'Forward-looking questions can be surfaced as coaching prompts rather than mandatory structured inputs. '
          'Additionally, Customer Area 3 and Competitor Area 3 retain trend questions.')

add_heading('Risk 2: Synthesis Evidence Density', 2)
bold_para('Risk: ', 'Fewer responses = fewer potential evidence quotes for Claude to cite.')
bold_para('Severity: ', 'Low')
bold_para('Mitigation: ', 'The synthesis requires minimum 2 evidence quotes per opportunity. With 27 responses '
          '(3 per area × 9 areas), there is more than sufficient evidence. The confidence score calculation '
          'awards 2 points per evidence item — 3 high-quality responses per area will achieve "High" confidence '
          'more reliably than 4 thin responses.')

add_heading('Risk 3: User Expects More Structure', 2)
bold_para('Risk: ', 'Some users may want more guided questioning, not less.')
bold_para('Severity: ', 'Low')
bold_para('Mitigation: ', 'The coaching sidebar provides supplementary prompting. Users who want depth '
          'can engage the coach for follow-up questions. The structured questions establish the baseline; '
          'the coach provides the depth. This also supports differentiated user experiences — quick users '
          'answer 3 and move on; thorough users engage the coach.')

add_heading('Risk 4: Question Reframing Breaks Existing Data', 2)
bold_para('Risk: ', 'Changing Q1 wording in Company Areas 1 and 2 may affect existing responses if '
          'the change is deployed mid-session for active users.')
bold_para('Severity: ', 'Medium')
bold_para('Mitigation: ', 'Implement as a schema change that only affects NEW conversations. Existing '
          'conversations retain their original question set and response data. The territory_insights '
          'table stores responses by index, not by question text, so the backend is unaffected.')

add_heading('Summary: Question Change Register', 2)
add_table(
    ['Area', 'Question Removed', 'Absorbed Into', 'Synthesis Impact'],
    [
        ['Company: Capabilities', 'Q4 (table stakes vs. differentiated)', 'Q1 reframe', 'Negligible'],
        ['Company: Resources', 'Q4 (talent gaps)', 'Q1 reframe', 'Negligible'],
        ['Company: Portfolio', 'Q3 (market position)', 'Competitor Territory', 'None — better placement'],
        ['Customer: Segmentation', 'Q4 (segment growth)', 'Customer Area 3', 'None — covered elsewhere'],
        ['Customer: Unmet Needs', 'Q4 (emerging needs)', 'Customer Area 3 + Coach', 'Minor'],
        ['Customer: Dynamics', 'Q2 (new alternatives)', 'Competitor Area 2', 'None — duplicate removed'],
        ['Competitor: Direct', 'Q4 (recent moves)', 'Coach conversation', 'Minor'],
        ['Competitor: Substitutes', 'Q4 (tech shifts)', 'Competitor Area 3', 'Negligible'],
        ['Competitor: Forces', 'Q4 (customer expectations)', 'Customer Territory', 'None — misplaced question'],
    ]
)

# Document control
doc.add_page_break()
add_heading('Document Control', 1)
add_table(
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ['1.0', 'February 2026', 'Frontera Product Team', 'Initial question analysis and reduction recommendations'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('— End of Document —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_path = r'c:\Users\deeks\frontera-platform\Background\Territory_Research_Questions_Analysis.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
