"""
Generate UX Journey & Platform Inspiration Research Document
for Frontera Platform Development.

Deep research into leading user experience journeys that can stimulate
ideas for the Frontera platform build, including PSC enhancements,
next strategy modules, and the connected user journey.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

NAVY = RGBColor(0x1a, 0x1f, 0x3a)
GOLD = RGBColor(0xfb, 0xbf, 0x24)
CYAN = RGBColor(0x08, 0x91, 0xb2)
WHITE = RGBColor(0xff, 0xff, 0xff)
SLATE = RGBColor(0x47, 0x55, 0x69)


def set_cell_shading(cell, color_hex):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(s)


def H(text, level=1):
    h = doc.add_heading(text, level)
    for r in h.runs:
        r.font.color.rgb = NAVY
    return h


def T(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hc = t.rows[0].cells
    for i, header in enumerate(headers):
        hc[i].text = header
        hc[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hc[i], '1a1f3a')
        hc[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        hc[i].paragraphs[0].runs[0].font.size = Pt(9)
    for rd in rows:
        r = t.add_row()
        for i, ct in enumerate(rd):
            r.cells[i].text = str(ct)
            for p in r.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return t


def B(items):
    for i in items:
        p = doc.add_paragraph(i, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)


def Q(text):
    p = doc.add_paragraph(text, style='Intense Quote')
    for run in p.runs:
        run.font.size = Pt(10)


def P(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p


def para(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)


def link_ref(label, url):
    """Add a reference link as formatted text."""
    p = doc.add_paragraph()
    r = p.add_run('\u2192 ')
    r.font.size = Pt(9)
    r.font.color.rgb = CYAN
    r2 = p.add_run(f'{label}')
    r2.font.size = Pt(9)
    r2.bold = True
    r2.font.color.rgb = CYAN
    r3 = p.add_run(f'\n   {url}')
    r3.font.size = Pt(8)
    r3.font.color.rgb = SLATE


def section_divider():
    p = doc.add_paragraph()
    p.add_run('\u2500' * 60).font.color.rgb = RGBColor(0xa5, 0xf3, 0xfc)


# ============================================================================
# TITLE PAGE
# ============================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('UX Journey & Platform Inspiration', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    r.font.size = Pt(32)
    r.font.color.rgb = NAVY

sub = doc.add_paragraph('Research for the Frontera Platform Build')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(20)
sub.runs[0].font.color.rgb = CYAN

doc.add_paragraph()
tag = doc.add_paragraph('Leading UX Journeys, Connected Experiences & Module Inspiration')
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.runs[0].font.size = Pt(14)
tag.runs[0].italic = True

doc.add_paragraph()
tag2 = doc.add_paragraph('Strategy Coach Enhancements \u2022 Next Modules \u2022 Connected User Journey')
tag2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag2.runs[0].font.size = Pt(12)
tag2.runs[0].font.color.rgb = SLATE

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Date: February 2026\n').font.size = Pt(10)
info.add_run('Research Scope: 25+ platforms across strategy, coaching, analytics & collaboration\n').font.size = Pt(10)
info.add_run('Cross-Referenced: Strategy Challenges Research, PSC Critical Analysis, Frontera PRD v2.1\n').font.size = Pt(10)
info.add_run('Perspective: Global Product Strategy & Product Design Expert').font.size = Pt(10)

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
H('Table of Contents', 1)
toc = [
    ('1.', 'Executive Context & Purpose'),
    ('2.', 'Platform Category Analysis: 6 UX Journey Models'),
    ('3.', 'The 12 UX Patterns That Should Inspire Frontera'),
    ('4.', 'Product Strategy Coach Enhancement Inspiration'),
    ('5.', 'Next Module Inspiration: Beyond the Strategy Coach'),
    ('6.', 'The Connected Frontera Journey: Full Vision'),
    ('7.', 'Reference Platform Directory'),
    ('8.', 'Cross-Reference: Challenges \u2192 Inspiration \u2192 Frontera Solution'),
]
for n, i in toc:
    p = doc.add_paragraph()
    r = p.add_run(f'{n} ')
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(i)
    r2.font.size = Pt(11)

doc.add_page_break()

# ============================================================================
# 1. EXECUTIVE CONTEXT & PURPOSE
# ============================================================================
H('1. Executive Context & Purpose', 1)

H('Why This Research Matters', 2)
para(
    'Frontera is building a category-defining platform at the intersection of three converging trends: '
    'AI-powered coaching, strategic planning software, and execution management. No existing platform '
    'occupies this precise intersection. This creates both an opportunity and a design challenge \u2014 '
    'there is no single platform to copy. Instead, Frontera must synthesise the best UX patterns from '
    'adjacent categories to create something genuinely new.'
)

para(
    'This research examines 25+ leading platforms across six categories to identify the UX patterns, '
    'journey architectures, and design innovations that should inform the next phase of Frontera\'s '
    'development. Each finding is cross-referenced back to the Strategy Challenges Research document '
    'and the Product Strategy Coach Critical Analysis to ensure every recommendation is grounded in '
    'real user pain points.'
)

H('The Frontera Opportunity Space', 2)
T(
    ['Category', 'Example Platforms', 'What They Do Well', 'What They Miss'],
    [
        ['Strategy Execution', 'Cascade, Perdoo, Quantive', 'OKR tracking, goal alignment, dashboards', 'No strategic thinking coaching \u2014 assumes strategy already exists'],
        ['AI Coaching', 'BetterUp, CoachHub, Valence', 'Personalised development, human+AI blend', 'No strategic methodology \u2014 coaching is generic leadership'],
        ['Product Management', 'Productboard, Linear, Amplitude', 'Roadmaps, analytics, team workflows', 'No strategy formulation \u2014 tools for executing, not thinking'],
        ['Research & Synthesis', 'Dovetail, Condens, EnjoyHQ', 'Evidence repository, tagging, insight extraction', 'No strategic framing \u2014 research without direction'],
        ['Collaboration', 'Miro, FigJam, Notion', 'Visual workshops, flexible canvases, AI assist', 'No guided methodology \u2014 blank canvas problem'],
        ['Performance Mgmt', 'Lattice, 15Five, Asana Goals', 'Goals cascade, 1:1s, review cycles', 'No strategy origination \u2014 goals without strategic grounding'],
    ]
)

P('Frontera\'s unique position: ', 'The only platform that helps leaders THINK strategically (coaching), '
  'FORMULATE strategy (methodology), AND bridge to EXECUTION (activation). Every other platform '
  'assumes at least one of these three already exists.')

doc.add_page_break()

H('Cross-Reference: Current Gaps Being Addressed', 2)
para(
    'This research directly addresses the gaps identified in two prior Frontera analyses:'
)

T(
    ['Source Document', 'Key Finding', 'What This Research Provides'],
    [
        ['Strategy Challenges Research', '84% worry products won\'t succeed; strategy disconnected from execution', 'Execution bridging patterns from Cascade, Perdoo, Asana Goals'],
        ['Strategy Challenges Research', '49% lack time for strategic planning', 'Time-compression UX from Linear, progressive gratification patterns'],
        ['Strategy Challenges Research', '80% don\'t involve engineers early', 'Collaboration patterns from Miro, FigJam, Notion'],
        ['PSC Critical Analysis', 'No Phase 5: Strategic Activation', 'Strategy-to-execution journeys from Lattice, Perdoo, Quantive'],
        ['PSC Critical Analysis', 'Energy problem \u2014 wow moment comes too late', 'Micro-celebration and progressive gratification patterns'],
        ['PSC Critical Analysis', 'Strategy is static, not adaptive', 'Living strategy patterns from Cascade, Notion, version control UX'],
        ['PSC Critical Analysis', 'Single-user tool, no collaboration', 'Multi-stakeholder patterns from Miro, Gong, FigJam'],
        ['PSC Critical Analysis', 'Missing artefacts (Strategy on a Page, Guardrails)', 'Artefact generation patterns from strategy dashboard tools'],
    ]
)

doc.add_page_break()

# ============================================================================
# 2. PLATFORM CATEGORY ANALYSIS
# ============================================================================
H('2. Platform Category Analysis: 6 UX Journey Models', 1)

para(
    'Each category below represents a distinct approach to solving parts of the problem Frontera '
    'addresses. By studying the UX journey of each, we can extract patterns that inform Frontera\'s '
    'connected experience.'
)

# --- Category 1: Strategy Execution ---
H('Category 1: Strategy Execution Platforms', 2)
Q('How platforms like Cascade, Perdoo, and Quantive connect strategy to measurable outcomes')

H('Cascade Strategy (cascade.app)', 3)
para(
    'Cascade is the world\'s #1 strategy execution platform, built on a 3-step journey model: '
    'Connect \u2192 Accelerate \u2192 Reap. The platform uses five core capability tabs (Planner, '
    'Alignment & Relationships, Metrics & Measures, Dashboards & Reports, Integrations) with '
    'animated progress ring visualisations that auto-cycle between views.'
)

P('Key UX Patterns for Frontera:', '')
B([
    'Strategy Map Visualisation: Cascade shows relationship graphs between goals, teams, and initiatives. Frontera should adopt similar visual mapping from strategic bets to team objectives',
    'Real-time Health Scores: Every goal has a calculated alignment and health score. Frontera\'s strategic bets should show similar "strategic health" indicators',
    'Template-Driven Onboarding: 1000+ templates reduce blank-page anxiety. Frontera\'s Research phase should offer pre-populated templates by industry',
    'AI Layer (Tapestry AI): "AI-powered insights that cut through the noise" \u2014 filtering and prioritisation built into the interface, not bolted on',
    'Tabbed Feature Navigation: Animated tab system creating a carousel-like discovery pattern that reveals platform capabilities progressively',
])

link_ref('Cascade Strategy Platform', 'https://www.cascade.app/strategy-platform')
link_ref('Cascade Solutions', 'https://www.cascade.app/solutions/strategy')

H('Perdoo (perdoo.com)', 3)
para(
    'Perdoo is the #1 OKR and strategy execution platform, distinguished by connecting goals for '
    'change (OKRs) with goals for stability (KPIs) in the same dashboard. Its Strategy Map feature '
    'visualises how company-level strategy cascades to team-level objectives.'
)

P('Key UX Patterns for Frontera:', '')
B([
    'Dual Goal Tracking: OKRs (change) + KPIs (stability) in one view. Frontera\'s Phase 5 should distinguish between strategic bets (change) and business-as-usual metrics (stability)',
    'Strategy Map: Visual cascade from company strategy \u2192 team OKRs \u2192 individual initiatives. Direct inspiration for Frontera\'s strategy-to-execution bridge',
    'Regular Check-ins: Built-in check-in cadence removes recency bias and keeps strategy alive. Frontera should embed review cadences into the post-strategy experience',
    'Integration Breadth: Jira, Asana, Salesforce, HubSpot, Slack, Teams. Frontera should plan integration architecture early',
    'Simplicity: "Overall, the tool is very simple to use" \u2014 Frontera must resist complexity creep as modules expand',
])

link_ref('Perdoo Platform', 'https://www.perdoo.com')
link_ref('Perdoo Strategy Alignment', 'https://www.perdoo.com/solutions/strategy')

doc.add_page_break()

# --- Category 2: AI Coaching ---
H('Category 2: AI-Powered Coaching Platforms', 2)
Q('How BetterUp, CoachHub, and Valence blend human expertise with AI to scale personalised development')

H('BetterUp (betterup.com)', 3)
para(
    'BetterUp is the market leader in enterprise coaching, having launched its AI Coaching product '
    'in January 2025. Their approach is instructive: AI handles daily skill-building while human '
    'coaches tackle high-stakes transformational breakthroughs. Every interaction builds on the last, '
    'creating what they call "synchronized Human + AI coaching."'
)

P('Key UX Patterns for Frontera:', '')
B([
    'AI as Complement, Not Replacement: "Our AI features enhance the capabilities of our expert coaches, not replace or mimic them." Frontera\'s AI coach should feel like a complement to the leader\'s own judgment, not a substitute',
    'Intelligent Matching (89% accuracy): Matching algorithms connect employees with the right coach. Frontera\'s persona selection could use similar matching based on the Personal Profile intake',
    'Multi-Persona Interfaces: BetterUp serves three distinct user flows (Partners/Intelligence, Employees/Learning, Coaches/Tools). Frontera should design for Admin, Leader, and Team Member personas',
    'In-Flow Integration: Slack and Teams integration means coaching happens where work happens. Frontera should consider embedding strategic nudges into daily workflow tools',
    'Workday Integration: Performance data and role expectations feed personalised coaching. Frontera should consider integrating HR/performance context into coaching recommendations',
    'Real-time Celebration: "Feedback and insights that celebrate their journey toward transformation" \u2014 progressive gratification embedded throughout the experience',
])

link_ref('BetterUp AI Coaching', 'https://www.betterup.com/products/betterup-ai-coaching')
link_ref('BetterUp Powered by AI', 'https://www.betterup.com/powered-by-ai')
link_ref('BetterUp Fall 2025 Release', 'https://www.betterup.com/platform-releases/fall-2025')

H('CoachHub (coachhub.com)', 3)
para(
    'CoachHub is Europe\'s leading digital coaching platform with Microsoft Teams integration '
    'bringing coaching directly into the workplace collaboration environment. Their UX design '
    'process is notably user-centred, with usability testing guiding feature prioritisation.'
)

P('Key UX Patterns for Frontera:', '')
B([
    'Dual Calendar Views: Distinct views and interactions for coaches vs. clients. Frontera should design different views for strategy leaders vs. team members consuming outputs',
    'CoachHub Insights\u2122: Real-time analysis of how coaching impacts company-wide performance. Frontera should track and visualise how strategic coaching translates to organisational outcomes',
    'Mobile-First Design: iOS and Android apps with 100+ language support. Frontera should plan mobile experience for strategy review and check-ins',
    'Accessibility as Foundation: "Accessible HTML5, unobtrusive Javascript, and mobile-responsive CSS" built into core technology. Frontera should embed accessibility from the start, not retrofit',
])

link_ref('CoachHub Platform', 'https://www.coachhub.com/coaching-platform')

doc.add_page_break()

# --- Category 3: Product Management ---
H('Category 3: Product Management Platforms', 2)
Q('How Productboard, Linear, and Amplitude create opinionated, efficient workflows for product teams')

H('Linear (linear.app) \u2014 THE DESIGN BENCHMARK', 3)
para(
    'Linear is the most design-forward B2B SaaS product in market. At $1.25B valuation with only '
    '100 employees, their opinionated UX philosophy is a masterclass in what enterprise software '
    'can be. Linear doesn\'t try to serve every workflow \u2014 it decides one excellent way to do things '
    'and optimises relentlessly for that path.'
)

P('Core Design Principles (directly applicable to Frontera):', '')
B([
    'Opinionated Defaults: "There may not be one right way to do something, but there is a default way." Frontera\'s methodology phases should have strong defaults that guide without constraining',
    'Speed as Feature: Updates sync in milliseconds. Keyboard-driven interface. Frontera should invest in perceived performance \u2014 every interaction should feel instant',
    'Directional Clarity: Linear design adds linearity \u2014 "being direct and offering minimal choices." Frontera\'s multi-phase journey should feel like a clear forward path, not a choose-your-own-adventure',
    'Visual Minimalism with Identity: Linear\'s 2025 redesign moved to monochrome black/white with very few bold colours. Frontera\'s Navy + Gold + Cyan is already distinctive \u2014 maintain restraint',
    'Craft over Features: "Championing craft, speed, and clarity over complexity." Each Frontera phase should feel crafted, not just functional',
    'No Manual Required: "You shouldn\'t need a handbook to start using Linear." Frontera\'s coaching methodology should be self-evident in the interface, not explained in documentation',
])

Q('"Linear designs it so that there\'s one really good way of doing things. Rather than flexible software that lets everyone invent their own workflows, which eventually creates chaos as teams scale."')

link_ref('The Linear Method (Figma Blog)', 'https://www.figma.com/blog/the-linear-method-opinionated-software/')
link_ref('Linear Design: The SaaS UX Trend (LogRocket)', 'https://blog.logrocket.com/ux-design/linear-design/')
link_ref('Linear Case Study (Eleken)', 'https://www.eleken.co/blog-posts/linear-app-case-study')
link_ref('How Linear Grew to $1.25B (Aakash Gupta)', 'https://aakashgupta.medium.com/linear-hit-1-25b-with-100-employees-heres-how-they-did-it-54e168a5145f')

H('Productboard (productboard.com)', 3)
para(
    'Productboard excels at connecting customer feedback to product decisions through a structured '
    'insights-to-roadmap pipeline. Their portal feature creates shared roadmap visibility that '
    'aligns stakeholders without endless meetings.'
)

P('Key UX Patterns for Frontera:', '')
B([
    'Insights-to-Roadmap Pipeline: Customer feedback flows through structured analysis to roadmap decisions. Frontera\'s territory research should similarly flow visibly into synthesis and bets',
    'Role-Based Views: Different views for PMs, executives, and stakeholders. Frontera should consider tailored views of strategy outputs for different organisational levels',
    'Interactive Roadmaps: Collaborative, shared-source-of-truth roadmaps. Frontera\'s strategic bets could be presented as an interactive strategic roadmap',
    'Milestone Tracking: Visual dependency and milestone tracking. Strategy execution should show dependencies between bets',
])

link_ref('Productboard Roadmapping', 'https://www.productboard.com/use-cases/product-roadmapping/')

H('Amplitude (amplitude.com)', 3)
para(
    'Amplitude represents the cutting edge of AI-integrated analytics. Their recent "Automated '
    'Insights" feature uses AI to run analyst-level workflows automatically, surfacing patterns '
    'that humans might miss. They connect usage analytics, segmentation, A/B testing, and audience '
    'activation in one unified interface.'
)

P('Key UX Patterns for Frontera:', '')
B([
    'AI-Powered Insight Generation: AI surfaces patterns automatically rather than requiring manual analysis. Frontera\'s synthesis should adopt similar proactive insight surfacing',
    'Natural Language Querying: Users can prompt insights conversationally. Frontera\'s coach should be able to answer analytical questions about the user\'s strategy data',
    'Unified Analytics: 45,000+ digital products connected. Frontera should consider connecting to product analytics tools to ground strategy in actual usage data',
    'GenAI Analytics: Automated insight summaries in natural language. Strategy review summaries should be generated and updated automatically',
])

link_ref('Amplitude AI Analytics', 'https://amplitude.com/')

doc.add_page_break()

# --- Category 4: Research & Synthesis ---
H('Category 4: Research & Synthesis Platforms', 2)
Q('How Dovetail, Miro, and FigJam handle evidence collection, synthesis, and collaborative insight generation')

H('Dovetail (dovetail.com)', 3)
para(
    'Dovetail is a research repository that aggregates video, audio, and text data to surface '
    'insights through AI-assisted tagging, auto-clustering themes, and rich reporting. Their '
    'approach to turning raw research into structured insights is directly relevant to Frontera\'s '
    'territory research and synthesis pipeline.'
)

P('Key UX Patterns for Frontera:', '')
B([
    'Auto-Clustering Themes: AI automatically groups related research into themes. Frontera\'s synthesis could use similar clustering to group territory responses before generating opportunities',
    'Rich Reporting: Insights presented through "mini blogs" containing quotes, videos, tags, and themes. Frontera\'s synthesis outputs should be similarly rich \u2014 not just cards, but narrative-quality reports',
    'Customisable Templates: Structured research output templates for different methodologies. Frontera\'s territories should offer industry-specific research templates',
    'Evidence Trail: Every insight is traceable back to source data. Frontera already does this well \u2014 maintain and extend the evidence linking pattern',
])

link_ref('Dovetail AI Research', 'https://www.perpetualny.com/blog/ai-enhanced-user-research-with-dovetail')
link_ref('Research Repository Guide', 'https://greatquestion.co/blog/ux-research-repository-guide')

H('Miro & FigJam: Collaborative Strategy Workshops', 3)
para(
    'Miro and FigJam represent the visual collaboration approach to strategy. Key innovations '
    'include AI-powered template generation, real-time multi-user interaction, and the ability '
    'to transition from brainstorming to structured output seamlessly.'
)

P('Key UX Patterns for Frontera:', '')
B([
    'AI Template Generation: FigJam generates entire workshop templates from natural language prompts. Frontera could generate pre-populated research frameworks from company context',
    'Infinite Canvas with Structure: Freedom within frameworks \u2014 users explore within guided structures. Frontera\'s canvas should balance structure (methodology) with flexibility (exploration)',
    'Real-Time Collaboration: Multiple users contributing simultaneously. Frontera\'s stakeholder input feature should support async and real-time contribution',
    'Workshop-to-Output Flow: Brainstorming seamlessly transitions into structured documents. Frontera\'s coaching conversations should similarly flow into structured strategic outputs',
    'Strategy Canvas Templates: Miro offers Blue Ocean Strategy Canvas, Value Curve, Playing to Win templates. Frontera should offer interactive versions of these within its methodology',
])

link_ref('FigJam Collaborative Whiteboard', 'https://www.figma.com/figjam/')
link_ref('Miro Strategy Canvas Templates', 'https://miro.com/strategic-planning/strategy-canvas-examples/')
link_ref('Miro UX Strategy Workshop', 'https://miro.com/miroverse/ux-strategy-workshop/')
link_ref('AI Workshop Facilitation (Medium)', 'https://medium.com/@uxraspberry/redesigning-the-workshop-ai-powered-facilitation-in-miro-figjam-and-mural-84d8b3deab62')

doc.add_page_break()

# --- Category 5: Performance Management ---
H('Category 5: Performance & Goal Management Platforms', 2)
Q('How Lattice, 15Five, and Asana Goals cascade strategy to team-level execution')

H('Lattice (lattice.com)', 3)
para(
    'Lattice connects OKRs with performance reviews, 1:1s, and career development. Their key '
    'insight: goals without performance context are disconnected from reality. Their 2025 launch '
    'of an AI Agent for HR provides always-on assistance for routine tasks.'
)

P('Key UX Patterns for Frontera:', '')
B([
    'Goals + Performance Integration: OKRs are integrated into feedback conversations and reviews. Frontera\'s strategic bets should be reviewable in check-in sessions',
    'Cascading Alignment: Visual progress tracking shows how individual work connects to company objectives. Frontera needs this cascade from strategy to team to individual',
    'AI Agent (2025): Always-on assistant for HR questions. Frontera\'s coach should be similarly available beyond the strategy creation session \u2014 answering strategic questions throughout the quarter',
    'Integration Depth: Jira, Salesforce, Slack, Teams. Strategy should be visible where teams already work',
])

P('UX Warning from Lattice: ', '"Reviewers find navigating Lattice\'s platform challenging due to its confusing setup '
  'and non-intuitive design." Frontera must avoid this pitfall as modules multiply. Clarity over comprehensiveness.')

link_ref('Lattice OKR Software', 'https://lattice.com/platform/goals/okrs')
link_ref('Lattice Goals', 'https://lattice.com/goals')

H('Asana Goals (asana.com)', 3)
para(
    'Asana\'s 2025 strategy is built on "intelligent orchestration" \u2014 streamlining work, '
    'surfacing risks, and aligning outcomes through automation and interconnectivity. Their '
    'Goals feature connects tasks and projects to OKRs and portfolios.'
)

P('Key UX Patterns for Frontera:', '')
B([
    'Laddering Visibility: "Teams can see how their efforts ladder up across departments and programs." Frontera should visualise how strategic bets cascade to team-level initiatives',
    'AI Studio Rule Builder: Intuitive rule builder for automated workflows. Frontera could use similar rule-based triggers for strategy review prompts',
    'Portfolio Views: Company-wide goal visibility with clear owners, timelines, and progress. Frontera\'s bet portfolio should offer similar executive-level visibility',
    'Connected Work: Every task connects to bigger goals. Strategic context should permeate daily work',
])

link_ref('Asana Goals Vision (Case Study)', 'https://cattsmall.com/design/asana-goals-vision/')

doc.add_page_break()

# --- Category 6: Connected Workspace ---
H('Category 6: Connected Workspace Platforms', 2)
Q('How Notion, Gong, and Salesforce create unified experiences across multiple modules')

H('Notion (notion.so) \u2014 THE CONNECTED JOURNEY BENCHMARK', 3)
para(
    'Notion represents the most ambitious attempt to create a unified workspace that connects '
    'documents, databases, wikis, and AI into a single platform. Their September 2025 AI Agent 3.0 '
    'launch introduced autonomous agents that can work across hundreds of pages simultaneously, '
    'performing up to 20 minutes of autonomous work.'
)

P('Key UX Patterns for Frontera:', '')
B([
    'Deep Workspace Awareness: Notion AI understands "page relationships, database patterns, team workflows, and historical context." Frontera\'s coach should have similar awareness of all user data across phases',
    'AI Agent Autonomy: Notion agents perform 20 minutes of autonomous work across hundreds of pages. Frontera\'s synthesis and bet generation should feel similarly autonomous \u2014 the AI doing the heavy lifting',
    'Unified Information Architecture: One platform replaces note-taking, project management, wikis, databases, and documentation. Frontera should aspire to be the single platform for strategy (not a collection of disconnected modules)',
    'Auto-Model Selection: Notion 3.2 automatically selects the best AI model (GPT-5.2, Claude, Gemini) for each task. Frontera could similarly route different coaching needs to different AI capabilities',
    'Relational Data: Databases connect to documents which connect to projects. Frontera\'s territory research should visibly connect to synthesis which connects to bets which connects to team outputs',
])

link_ref('Notion Releases', 'https://www.notion.com/releases')
link_ref('Notion AI Agent Use Cases', 'https://thecrunch.io/notion-ai-agent/')
link_ref('Notion AI Features', 'https://kipwise.com/blog/notion-ai-features-capabilities')

H('Gong (gong.io) \u2014 THE INSIGHT-TO-ACTION BENCHMARK', 3)
para(
    'Gong transforms unstructured conversation data into structured, actionable intelligence. '
    'Their platform captures, transcribes, and analyses business conversations \u2014 then surfaces '
    'patterns, objections, and buying signals that would otherwise be missed. Gong is a 2025 '
    'Gartner Magic Quadrant Leader for Revenue Action Orchestration.'
)

P('Key UX Patterns for Frontera:', '')
B([
    'Behaviour Design through Data: "Call scorecards and talk ratio visuals subtly shape rep performance without requiring formal training." Frontera\'s coaching should similarly shape strategic thinking through embedded patterns, not explicit teaching',
    'Automatic Capture + AI Summaries: "Connect your Zoom or CRM once, and Gong does the rest." Frontera should minimise manual data entry \u2014 capture context automatically from uploaded documents, AI research, and coaching conversations',
    'From Analysis to Action: Gong\'s 2025 shift toward orchestration \u2014 auto-populating CRM fields from conversation data. Frontera\'s synthesis should auto-populate strategic frameworks from research data',
    'Pattern Recognition at Scale: AI identifies patterns across thousands of conversations. Frontera should identify patterns across multiple users\' strategy sessions (aggregate insights for platform improvement)',
])

link_ref('Gong Conversation Intelligence', 'https://www.gong.io/conversation-intelligence')
link_ref('Gong Analytics Features', 'https://www.oliv.ai/blog/gong-analytics')

doc.add_page_break()

# ============================================================================
# 3. THE 12 UX PATTERNS
# ============================================================================
H('3. The 12 UX Patterns That Should Inspire Frontera', 1)

para(
    'Synthesising across all 25+ platforms researched, twelve UX patterns emerge as critical '
    'inspiration for Frontera\'s next development phase. Each pattern is grounded in the Strategy '
    'Challenges Research and PSC Critical Analysis.'
)

# Pattern 1
H('Pattern 1: Progressive Gratification (\u2018Micro-Wow\u2019 Moments)', 2)
P('Challenge Addressed: ', 'PSC Critical Analysis \u2014 "Energy problem: wow moment comes too late"')
P('Inspired By: ', 'BetterUp celebrations, Kontentino milestone modals, HubSpot onboarding')

para(
    'The most engaged enterprise users experience regular moments of reward throughout their journey, '
    'not just at the end. The PSC currently requires ~2.5 hours of research input before users see any '
    'strategic output. This is a known abandonment risk.'
)

P('Frontera Application:', '')
B([
    'After each territory completion: generate a "Territory Snapshot" \u2014 a 3-bullet summary of key themes',
    'After 3 of 9 research areas: show a "Research Progress Milestone" with emerging patterns',
    'During research: real-time word cloud or tag cloud visualising themes as they emerge',
    'On phase transition: celebratory micro-animation with summary of what was accomplished',
    'Progress indicators that show content richness, not just completion (e.g., "87% of companies in your industry provide 4+ territory responses \u2014 you\'re at 3")',
])

# Pattern 2
H('Pattern 2: Opinionated Defaults with Escape Hatches', 2)
P('Challenge Addressed: ', 'Strategy Challenges \u2014 49% lack time for strategic thinking')
P('Inspired By: ', 'Linear (one good way), Cascade (templates), Perdoo (simplicity)')

para(
    'Linear\'s core philosophy \u2014 "There may not be one right way, but there is a default way" \u2014 '
    'is the antidote to blank-page paralysis. Frontera\'s methodology already provides structure, '
    'but it should go further by offering intelligent defaults at every step.'
)

P('Frontera Application:', '')
B([
    'Pre-populate research questions with AI-generated draft answers from uploaded documents',
    'Offer "Express Mode" that converts the 36-question research into a 15-minute coach-led interview',
    'Provide industry-specific territory templates (e.g., FinTech Customer Territory vs. Healthcare Customer Territory)',
    'Default scoring in synthesis with "adjust and rationale" option \u2014 don\'t ask users to score from scratch',
    'Smart defaults for kill criteria and success metrics based on industry benchmarks',
])

# Pattern 3
H('Pattern 3: Living Strategy, Not Static Documents', 2)
P('Challenge Addressed: ', 'PSC Critical Analysis \u2014 "Strategy is static, not adaptive" (HIGH severity)')
P('Inspired By: ', 'Notion (living workspace), Cascade (real-time health), Perdoo (check-ins)')

para(
    'Every leading platform has moved from documents to dynamic, continuously-updated views. '
    'Frontera\'s strategy currently terminates as a PDF export. This fundamentally contradicts '
    'Cagan\'s principle that "good strategies are constantly updated based on learnings."'
)

P('Frontera Application:', '')
B([
    'Replace PDF-as-endpoint with a "Strategy Dashboard" \u2014 a living view of the current strategy',
    'Strategy versioning with visual diff: "What changed since last quarter?"',
    'Signal log where users record market events that may affect strategy',
    'Assumption tracker: validated/invalidated/untested status for each WWHBT assumption',
    'Kill-date notifications that trigger automatic strategy review prompts',
    'Quarterly strategy health check prompted by the coach',
])

doc.add_page_break()

# Pattern 4
H('Pattern 4: Multi-Stakeholder Collaboration', 2)
P('Challenge Addressed: ', 'Strategy Challenges \u2014 80% don\'t involve engineers early; 35% cite stakeholder influence')
P('Inspired By: ', 'Miro/FigJam (real-time collaboration), Gong (automatic capture), Notion (shared workspace)')

P('Frontera Application:', '')
B([
    'Shareable research area links: "Ask your Head of Sales to complete the Customer Segments research area"',
    'Async contribution mode: stakeholders receive focused questions, not the full interface',
    'Comment/annotation layer on synthesis outputs for team discussion',
    'Role-based views: strategist sees full canvas, team member sees relevant bet context only',
    'Alignment Gap Detector: flag contradictions between different contributors\' research responses',
])

# Pattern 5
H('Pattern 5: AI as Ambient Intelligence', 2)
P('Challenge Addressed: ', 'Strategy Challenges \u2014 AI not yet helping with strategic work (only saving 10-60 min on docs)')
P('Inspired By: ', 'BetterUp (synchronized coaching), Notion 3.2 (autonomous agents), Amplitude (automated insights)')

P('Frontera Application:', '')
B([
    'Proactive blind-spot detection: "I notice you haven\'t mentioned retention in your Customer Territory"',
    'Auto-generated research summaries from uploaded documents before users start writing',
    'Coach that remembers patterns across sessions: "You tend to focus on features rather than markets"',
    'Ambient strategy monitoring: weekly email digest of news relevant to the user\'s competitive landscape',
    'AI-generated connection maps showing relationships between territory insights and strategic opportunities',
])

# Pattern 6
H('Pattern 6: The Strategy-to-Execution Bridge', 2)
P('Challenge Addressed: ', 'Strategy Challenges \u2014 84% worry products won\'t succeed; PSC Critical Analysis \u2014 Missing Phase 5')
P('Inspired By: ', 'Cascade (strategy map), Perdoo (OKR cascade), Lattice (goals + performance), Asana (laddering)')

P('Frontera Application (Phase 5: Strategic Activation):', '')
B([
    'Team Brief Generator: translate each bet into problem-to-solve, guardrails, success metrics, and decision framework for team leads',
    'OKR Cascade: auto-generate draft OKRs from strategic bets with coach review',
    'Strategy on a Page: one-view summary suitable for Slack, all-hands, or wall display',
    'Strategic Guardrails Document: "We will / We will not" statements derived from synthesis tensions',
    'Decision Framework: "When choosing between X and Y, prioritise..." rules derived from strategy',
    'Stakeholder Communication Pack: tailored messaging for CEO (2-min narrative), CTO (capability needs), Sales (positioning)',
])

doc.add_page_break()

# Pattern 7
H('Pattern 7: Evidence Quality Visualisation', 2)
P('Challenge Addressed: ', 'Strategy Challenges \u2014 49% can\'t fit in data analysis; PSC Critical Analysis \u2014 no confidence ratings')
P('Inspired By: ', 'Dovetail (evidence trails), Gong (pattern detection), Amplitude (data visualisation)')

P('Frontera Application:', '')
B([
    'Confidence rating per research response (High/Medium/Low/Guess) \u2014 visualised as colour-coded evidence bars',
    'Evidence strength indicator on each synthesis opportunity showing how well-supported it is',
    'Gap Analysis Radar: visual representation of research coverage vs. gaps across all 9 areas',
    'Source quality badges: distinguish between primary research, AI-generated, and anecdotal evidence',
    'Interactive evidence map: click any synthesis finding to see all supporting research responses',
])

# Pattern 8
H('Pattern 8: Gamified Progress with Substance', 2)
P('Challenge Addressed: ', 'Strategy Challenges \u2014 49% lack time (need momentum); PSC \u2014 research abandonment risk')
P('Inspired By: ', 'Enterprise onboarding (checklists, Zeigarnik effect), BetterUp (celebrations), SaaS trends 2025')

para(
    'The Zeigarnik effect \u2014 people remember uncompleted tasks \u2014 makes checklists psychologically '
    'sticky. 92% of top SaaS apps implement some form of in-app guidance. But gamification must '
    'feel substantive in enterprise contexts, not trivial.'
)

P('Frontera Application:', '')
B([
    'Strategic Readiness Score (0-100): calculated from research depth, evidence quality, and coverage breadth',
    'Research completeness benchmarking: "Companies in your sector typically provide 350+ words per research area \u2014 your average is 180"',
    'Phase completion badges with strategic insight summaries (not empty trophies)',
    'Coach acknowledgment of effort: "You\'ve invested 90 minutes mapping your competitive terrain. That\'s more rigorous than most strategy consulting engagements I\'ve seen"',
    'Time-invested tracking: show users their cumulative strategic thinking time vs. industry averages',
])

# Pattern 9
H('Pattern 9: Personalised Adaptive Coaching', 2)
P('Challenge Addressed: ', 'Strategy Challenges \u2014 59% need strategy skills; PSC \u2014 no progressive difficulty')
P('Inspired By: ', 'BetterUp (intelligent matching), CoachHub (personalised insights), Frontera Personal Profile PRD')

P('Frontera Application:', '')
B([
    'Personal Profile shapes coaching depth (first-time founder vs. 20-year CPO)',
    'Strategic maturity assessment at onboarding adapts research question complexity',
    'Progressive challenge escalation: gentle probing in Discovery, demanding rigour in Bets',
    'Learning patterns tracked: "You consistently provide strong customer insights but skip competitive analysis"',
    'Post-phase reflection prompts: "What surprised you? What assumption changed?"',
    'Coach persona recommendation based on profile analysis (not just self-selection)',
])

doc.add_page_break()

# Pattern 10
H('Pattern 10: Visual Strategy Communication', 2)
P('Challenge Addressed: ', 'PSC Critical Analysis \u2014 missing Strategy on a Page, no portfolio view')
P('Inspired By: ', 'Cascade (strategy maps), Blue Ocean (strategy canvas), Miro (visual frameworks)')

P('Frontera Application:', '')
B([
    'Interactive 2x2 Strategy Map (already exists) \u2014 add click-through to bet details and evidence',
    'Strategy Canvas / Value Curve: plot company positioning against competitors on key factors',
    'Bet Portfolio Timeline: visual sequencing showing bet dependencies and critical path',
    'Effort vs. Impact Scatter Plot: interactive view of all bets with drag-to-adjust capability',
    'Competitive Response Scenarios: visual wargaming of "If we execute Bet A, Competitor X likely responds with..."',
    'Strategy on a Page Generator: auto-generated single-view summary of entire strategy',
])

# Pattern 11
H('Pattern 11: Workflow Integration Architecture', 2)
P('Challenge Addressed: ', 'Strategy Challenges \u2014 strategy disconnected from execution')
P('Inspired By: ', 'Perdoo (200+ integrations), Lattice (Jira/Slack/Teams), Cascade (bidirectional sync)')

P('Frontera Application:', '')
B([
    'Slack/Teams integration: push strategy context updates to channels where teams work',
    'Jira/Linear/Asana integration: link strategic bets to product backlog items',
    'Notion/Confluence export: push Strategy on a Page to company wikis',
    'Calendar integration: auto-schedule strategy review sessions based on kill dates',
    'API for custom integrations: allow enterprises to connect Frontera to their tech stack',
])

# Pattern 12
H('Pattern 12: Continuous Strategy Evolution', 2)
P('Challenge Addressed: ', 'PSC Critical Analysis \u2014 "Strategy is static" (CRITICAL gap)')
P('Inspired By: ', 'Notion (living workspace), Version control UX, Cascade (real-time health scores)')

P('Frontera Application:', '')
B([
    'Strategy Review Mode: quarterly prompted re-assessment with guided review template',
    'Signal Log: structured log of market events, competitive moves, and internal changes',
    'Assumption Tracker Dashboard: validated/invalidated/untested with visual progress',
    'Strategy Diff View: "Since your last review, 2 assumptions have been invalidated and 1 new competitor has entered"',
    'Feedback Loop: bet outcomes update territory research confidence ratings',
    'Coach-prompted reviews: "Your kill date for Bet 3 is next week. Let\'s review the evidence"',
])

doc.add_page_break()

# ============================================================================
# 4. PSC ENHANCEMENT INSPIRATION
# ============================================================================
H('4. Product Strategy Coach Enhancement Inspiration', 1)

para(
    'The following enhancements to the existing Product Strategy Coach are inspired by the platform '
    'research above, prioritised by the gaps identified in the Critical Analysis.'
)

H('Discovery Phase Enhancements', 2)
T(
    ['Enhancement', 'Inspired By', 'Impact', 'Challenge Ref'],
    [
        ['AI auto-extract from uploaded docs \u2192 pre-fill research areas', 'Gong (automatic capture), Notion (workspace awareness)', 'Saves 30-45 min; reduces blank-page anxiety', '#1 Time, #4 Research'],
        ['Personal Profile coaching intake', 'BetterUp (intelligent matching), CoachHub (personalisation)', 'Enables adaptive coaching from first interaction', '#7 Competency'],
        ['Industry-specific Discovery templates', 'Cascade (1000+ templates), Miro (strategy templates)', 'Contextualises experience immediately', '#1 Time'],
        ['Proactive coach questioning in Discovery', 'BetterUp (real-time guidance), Gong (behaviour design)', 'Under-coached Discovery becomes actively guided', 'PSC Critical Analysis'],
    ]
)

H('Research Phase Enhancements', 2)
T(
    ['Enhancement', 'Inspired By', 'Impact', 'Challenge Ref'],
    [
        ['Micro-synthesis after each territory', 'Dovetail (auto-clustering), progressive gratification pattern', 'Breaks energy problem; provides wow moments during research', 'PSC Critical Analysis'],
        ['Confidence rating per response', 'Dovetail (evidence quality), Amplitude (data quality indicators)', 'Surfaces evidence quality; feeds synthesis confidence', '#4 Research'],
        ['Express Mode (coach-led interview)', 'Linear (opinionated defaults), BetterUp (guided sessions)', 'Reduces 2.5 hours to 45 min for time-constrained leaders', '#1 Time'],
        ['Research method hints per question', 'CoachHub (coaching methodology), BetterUp (skill building)', 'Builds capability while guiding; addresses competency gap', '#7 Competency'],
        ['Blind spot detection', 'Gong (pattern recognition), BetterUp (proactive coaching)', '"I notice you haven\'t mentioned retention" \u2014 proactive coaching', 'PSC Critical Analysis'],
    ]
)

H('Synthesis Phase Enhancements', 2)
T(
    ['Enhancement', 'Inspired By', 'Impact', 'Challenge Ref'],
    [
        ['Interactive evidence map', 'Dovetail (evidence trails), Amplitude (data visualisation)', 'Click any finding to see all supporting research', '#4 Research'],
        ['Leadership score calibration', 'Perdoo (check-ins), Lattice (performance calibration)', 'Users validate/adjust AI-generated scores with rationale', '#3 Prioritisation'],
        ['Strategy Canvas visualisation', 'Blue Ocean (strategy canvas), Miro (visual frameworks)', 'Competitive positioning visualisation alongside synthesis', '#4 Research'],
        ['Tension Simulator integration', 'Frontera PRD-UC4, Miro (debate templates)', 'Expert debate mode for strategic tensions', 'PSC planned module'],
    ]
)

H('Strategic Bets Phase Enhancements', 2)
T(
    ['Enhancement', 'Inspired By', 'Impact', 'Challenge Ref'],
    [
        ['Interactive portfolio view', 'Cascade (strategy map), Productboard (roadmap)', 'Effort vs. impact scatter plot for all bets', '#3 Prioritisation'],
        ['Bet sequencing timeline', 'Asana Goals (milestone tracking), Linear (project timelines)', 'Shows dependencies and critical path', '#3 Prioritisation'],
        ['Strategy on a Page generator', 'Cascade (one-page plan), strategy dashboard tools', 'Shareable single-view strategy summary', '#2 Execution gap'],
        ['Challenge escalation in coaching', 'BetterUp (adaptive coaching), PSC Critical Analysis', 'Coach demands more rigour in Bets than Discovery', 'PSC Critical Analysis'],
    ]
)

doc.add_page_break()

# ============================================================================
# 5. NEXT MODULE INSPIRATION
# ============================================================================
H('5. Next Module Inspiration: Beyond the Strategy Coach', 1)

para(
    'The Frontera platform has five planned modules beyond the current Product Strategy Coach. '
    'The following analysis maps each to platform inspiration and recommended UX patterns.'
)

# Module: Phase 5
H('Module: Phase 5 \u2014 Strategic Activation (NEW \u2014 HIGHEST PRIORITY)', 2)
P('Gap Origin: ', 'PSC Critical Analysis \u2014 "No mechanism to cascade strategy to team-level objectives"')
P('Platform Inspiration: ', 'Cascade (strategy map), Perdoo (OKR cascade), Lattice (goals + reviews), Asana Goals (laddering)')

para(
    'This is the most critical missing module. It bridges the gap between "here are your bets" '
    'and "here\'s what each team should do." No existing strategy tool does this well because '
    'most assume strategy already exists. Frontera is uniquely positioned to build this bridge '
    'because it owns the full journey from research to strategy.'
)

P('Recommended UX Journey:', '')
B([
    'Step 1: Strategy on a Page \u2014 auto-generated single-view summary with editable sections (Inspired by: Cascade one-page plans, Perdoo strategy map)',
    'Step 2: Strategic Guardrails \u2014 "We will / We will not" statements derived from synthesis tensions (Inspired by: Cagan\'s guardrails concept)',
    'Step 3: Team Brief Generator \u2014 for each bet, generate: problem to solve, guardrails, success metrics, decision framework (Inspired by: BetterUp role-based views, Lattice cascading alignment)',
    'Step 4: OKR Draft Generator \u2014 translate bets into draft Objectives and Key Results (Inspired by: Perdoo OKR+KPI, Lattice OKR software)',
    'Step 5: Decision Framework \u2014 "When choosing between X and Y, prioritise Z because..." rules from strategy (Inspired by: Linear opinionated defaults)',
    'Step 6: Stakeholder Communication Pack \u2014 tailored views for CEO, CTO, Sales VP (Inspired by: BetterUp multi-persona interfaces)',
])

link_ref('Cascade Strategy Map', 'https://www.cascade.app/strategy-platform')
link_ref('Perdoo Strategy Alignment', 'https://www.perdoo.com/solutions/strategy')
link_ref('Asana Goals Vision', 'https://cattsmall.com/design/asana-goals-vision/')

doc.add_page_break()

# Module: Expert Perspectives (UC1)
H('Module: Expert Perspectives (UC1 \u2014 Planned)', 2)
P('PRD Reference: ', 'PRD-UC1-Expert-Perspectives.md')
P('Platform Inspiration: ', 'Gong (pattern detection from conversations), Dovetail (evidence linking), Amplitude (insight surfacing)')

para(
    'Expert Perspectives indexes 301 Lenny\'s Podcast transcripts to enrich coaching with '
    'real practitioner wisdom. This transforms coaching from methodology-driven to '
    'experience-grounded.'
)

P('UX Inspiration:', '')
B([
    'Gong\'s citation pattern: relevant insights surfaced automatically in context, not as a separate library',
    'Dovetail\'s evidence trail: every expert citation should be traceable to source transcript',
    'Amplitude\'s natural language querying: users should be able to ask "What do experts say about pricing strategy?" and get relevant citations',
    'Expert Sources panel should use Productboard-style filtering (by speaker, company, topic, phase)',
])

# Module: Strategy Sparring Partners (UC2)
H('Module: Strategy Sparring Partners (UC2 \u2014 Planned)', 2)
P('PRD Reference: ', 'PRD-UC2-Strategy-Sparring-Partners.md')
P('Platform Inspiration: ', 'BetterUp (intelligent matching), CoachHub (dual views), Notion (auto-model selection)')

para(
    'Three new expert personas (Growth Architect, Product Purist, Scale Navigator) extend the '
    'coaching system with domain-specific expertise. The UX challenge is making persona selection '
    'feel meaningful, not arbitrary.'
)

P('UX Inspiration:', '')
B([
    'BetterUp\'s 89% match accuracy: persona recommendation should be based on Personal Profile analysis, not random self-selection',
    'Notion 3.2\'s auto-model selection: the system should suggest the most relevant persona based on current strategic challenge',
    'CoachHub\'s dual views: different persona coaching should feel genuinely different \u2014 vocabulary, frameworks, challenge intensity',
    'Mid-session persona switching (PRD requirement) should include a "why switch?" prompt that feeds back into profile refinement',
])

# Module: Case Study Engine (UC3)
H('Module: Case Study Engine (UC3 \u2014 Planned)', 2)
P('PRD Reference: ', 'PRD-UC3-Case-Study-Engine.md')
P('Platform Inspiration: ', 'Dovetail (research repository), Amplitude (contextual insights), Gong (proactive suggestions)')

P('UX Inspiration:', '')
B([
    'Dovetail\'s mini-blog format: case studies should be narrative-quality, not database records',
    'Gong\'s proactive suggestions: coach should suggest cases once per 5 messages, not as a separate library (PRD already specifies this)',
    'Amplitude\'s contextual relevance: cases should be scored by relevance to user\'s current context and surfaced accordingly',
    'Productboard\'s filtering: Case Library should filter by industry, stage, challenge type, and phase relevance',
])

# Module: Tension Simulator (UC4)
H('Module: Tension Simulator (UC4 \u2014 Planned)', 2)
P('PRD Reference: ', 'PRD-UC4-Tension-Simulator.md')
P('Platform Inspiration: ', 'Miro (debate templates), FigJam (workshop facilitation), strategy canvas tools')

P('UX Inspiration:', '')
B([
    'Miro\'s debate facilitation: visual two-column layout with colour-coded positioning (blue vs. red)',
    'FigJam\'s AI workshop generation: tension simulations could be auto-generated from synthesis tensions',
    'Interactive voting mechanism: user\'s choice should carry forward into strategic bets with evidence trail',
    'Strategy canvas integration: tensions could be visualised as competing value curves on a strategy canvas',
    'The 20 initial tensions mapped to expert positions is a strong foundation \u2014 the UX should make the debate feel visceral, not academic',
])

# Module: Leadership Playbook (UC5)
H('Module: Leadership Playbook Generator (UC5 \u2014 Planned)', 2)
P('PRD Reference: ', 'PRD-UC5-Leadership-Playbook.md')
P('Platform Inspiration: ', 'BetterUp (personalised development), Lattice (performance development), CoachHub (coaching insights)')

P('UX Inspiration:', '')
B([
    'BetterUp\'s journey celebration: playbook generation should feel like a milestone achievement, not just another output',
    'Lattice\'s performance integration: playbook themes should connect to strategic challenges, not generic leadership content',
    'CoachHub Insights: track how leadership development connects to strategic execution outcomes',
    'Shareable sections: individual playbook themes should be shareable with team members for cascading development',
    'Regeneration trigger: playbook should update when strategy evolves (new bets, new tensions)',
])

doc.add_page_break()

# ============================================================================
# 6. THE CONNECTED FRONTERA JOURNEY
# ============================================================================
H('6. The Connected Frontera Journey: Full Vision', 1)

para(
    'Drawing from all platform research, this section presents the recommended end-to-end '
    'Frontera user journey \u2014 from first login to ongoing strategic evolution. Each step '
    'references the platform inspiration and design pattern that should inform its design.'
)

H('Journey Architecture', 2)

T(
    ['Stage', 'What Happens', 'Platform Inspiration', 'Key UX Pattern'],
    [
        ['1. Onboarding', 'Org setup, company context, industry selection', 'Cascade (template onboarding), HubSpot (guided setup)', 'Industry-specific templates'],
        ['2. Personal Profile', '5-min coaching intake \u2014 role, objectives, style, experience', 'BetterUp (intelligent matching), CoachHub (personalised coaching)', 'Conversational intake, not forms'],
        ['3. Persona Selection', 'Coach persona recommendation + selection', 'BetterUp (89% match), Notion (auto-model selection)', 'AI-recommended with override'],
        ['4. Discovery', 'Document upload, AI Research, strategic baseline', 'Gong (automatic capture), Notion (workspace awareness)', 'AI pre-fill, minimal manual entry'],
        ['5. Research', 'Territory mapping across Company, Customer, Competitor', 'Dovetail (evidence repository), Linear (opinionated defaults)', 'Express Mode option, micro-synthesis'],
        ['6. Synthesis', 'Cross-pillar insight generation, opportunity scoring', 'Amplitude (automated insights), Dovetail (auto-clustering)', 'Interactive evidence map, calibration'],
        ['7. Strategic Bets', 'Hypothesis-driven planning with kill criteria', 'Productboard (roadmapping), Linear (project management)', 'Portfolio view, sequencing timeline'],
        ['8. Activation (NEW)', 'Team briefs, OKRs, guardrails, communication packs', 'Cascade + Perdoo + Lattice (strategy-to-execution)', 'Strategy on a Page, cascading outputs'],
        ['9. Living Strategy', 'Ongoing review, signal monitoring, assumption tracking', 'Notion (living workspace), Cascade (real-time health)', 'Strategy dashboard, version history'],
        ['10. Evolution', 'Quarterly review cycles, learning feedback loop', 'Perdoo (check-ins), Lattice (performance reviews)', 'Coach-prompted reviews, diff view'],
    ]
)

H('The "One Platform" Principle', 2)
Q('"What sets Notion AI apart is its deep workspace awareness \u2014 page relationships, database patterns, team workflows, and historical context." \u2014 This is the aspiration for Frontera.')

para(
    'Frontera\'s connected journey must feel like ONE platform, not a collection of separate modules. '
    'The key to achieving this is data continuity \u2014 every piece of data entered in one phase should '
    'be visible and usable in every subsequent phase. Specifically:'
)

B([
    'Personal Profile \u2192 shapes coaching tone, depth, and persona recommendation throughout',
    'Uploaded documents \u2192 referenced in coaching conversations, pre-fill research areas, cited in synthesis',
    'Territory research \u2192 feeds synthesis, grounds strategic bets, informs leadership playbook',
    'Synthesis insights \u2192 become evidence for bets, tensions for Tension Simulator, context for Case Studies',
    'Strategic bets \u2192 cascade to team briefs, OKRs, guardrails, and communication packs',
    'Coaching conversations \u2192 inform persona adaptation, identify blind spots, track competency growth',
    'Expert citations \u2192 enrichen coaching, strengthen evidence trails, feed leadership playbook',
    'Bet outcomes \u2192 update territory research confidence, inform next strategy cycle',
])

H('Navigation Architecture', 2)
para(
    'The navigation should reflect the journey metaphor while allowing non-linear access. '
    'Inspired by Linear\'s sidebar navigation and Notion\'s relational page structure:'
)

B([
    'Primary Navigation (Left Sidebar): Strategy sessions list, Dashboard, Team, Settings',
    'Strategy Session Navigation (Top): Phase stepper showing Discovery \u2192 Research \u2192 Synthesis \u2192 Bets \u2192 Activation',
    'Canvas Navigation (Right Panel): Phase-specific content with tabbed sub-navigation',
    'Coach (Left Panel): Always-present coaching sidebar with phase-aware context',
    'Quick Actions Bar: Export, Share, Generate, Review \u2014 available at all times',
    'Breadcrumb Trail: shows current position in journey hierarchy',
])

doc.add_page_break()

H('The Daily Engagement Pattern', 2)
para(
    'Strategy is not a one-time event. Frontera must design for ongoing engagement beyond the '
    'initial strategy creation journey. Inspired by Perdoo\'s check-in cadence, Lattice\'s review '
    'cycles, and BetterUp\'s daily coaching nudges:'
)

T(
    ['Cadence', 'What Happens', 'UX Pattern', 'Platform Inspiration'],
    [
        ['Daily', 'Strategy context notification in Slack/Teams (if integrated)', 'Ambient intelligence', 'BetterUp in-flow coaching'],
        ['Weekly', 'Quick check-in: "Any new signals this week?" (2-min prompt)', 'Micro check-in', 'Perdoo regular check-ins'],
        ['Monthly', 'Strategy health review: assumption tracker, bet progress', 'Dashboard review', 'Cascade real-time health'],
        ['Quarterly', 'Full strategy review: what changed, what\'s validated, what\'s new', 'Guided review', 'Lattice performance review'],
        ['On Trigger', 'Kill date reached, market event logged, assumption invalidated', 'Event-driven review', 'Notion AI agents (autonomous work)'],
    ]
)

H('Module Interconnection Map', 2)
para(
    'Each module should connect to others, creating a web of strategic intelligence rather than '
    'isolated features:'
)

T(
    ['Module', 'Feeds Into', 'Receives From'],
    [
        ['Personal Profile', 'Coach personalisation, Persona recommendation, Playbook themes', 'Coaching patterns (progressive refinement)'],
        ['Strategy Coach', 'Territory research, Synthesis, Strategic bets', 'Personal Profile, Expert Perspectives, Case Studies'],
        ['Expert Perspectives (UC1)', 'Coach citations, Synthesis evidence, Playbook themes', 'Territory context (relevance matching)'],
        ['Sparring Partners (UC2)', 'Coaching conversations, Bet coaching', 'Personal Profile (persona match)'],
        ['Case Study Engine (UC3)', 'Coach suggestions, Synthesis evidence, Team briefs', 'Territory context, Synthesis tensions'],
        ['Tension Simulator (UC4)', 'Bet choices, Evidence trail, Team alignment', 'Synthesis tensions, Expert positions'],
        ['Leadership Playbook (UC5)', 'Personal development, Team sharing', 'Strategic challenges, Coaching patterns, Expert themes'],
        ['Strategic Activation (Phase 5)', 'Team briefs, OKRs, Guardrails, Communication', 'Strategic bets, Synthesis, Territory research'],
    ]
)

doc.add_page_break()

# ============================================================================
# 7. REFERENCE PLATFORM DIRECTORY
# ============================================================================
H('7. Reference Platform Directory', 1)

para(
    'Complete directory of platforms researched with URLs for UX review and screenshots. '
    'Organised by category for easy reference during design reviews.'
)

H('Strategy Execution Platforms', 2)
T(
    ['Platform', 'URL', 'Key Feature to Study'],
    [
        ['Cascade Strategy', 'cascade.app', 'Strategy map visualisation, AI insights (Tapestry)'],
        ['Perdoo', 'perdoo.com', 'OKR + KPI dual tracking, strategy map cascade'],
        ['Quantive (formerly Gtmhub)', 'quantive.com', 'OKR platform with strategic planning'],
        ['Strategy Software (Envisio)', 'strategysoftware.com', 'AI-powered analytics dashboard design'],
        ['Spider Strategies', 'spiderstrategies.com', 'Strategic plan dashboard design patterns'],
        ['AchieveIt', 'achieveit.com', 'Actionable strategic dashboards'],
    ]
)

H('AI Coaching Platforms', 2)
T(
    ['Platform', 'URL', 'Key Feature to Study'],
    [
        ['BetterUp', 'betterup.com', 'Human+AI coaching, intelligent matching, in-flow integration'],
        ['CoachHub', 'coachhub.com', 'Teams integration, coaching insights analytics'],
        ['Valence', 'valence.co', 'AI team coaching, group dynamics'],
        ['15Five', '15five.com', 'Performance + coaching integration'],
    ]
)

H('Product Management & Analytics', 2)
T(
    ['Platform', 'URL', 'Key Feature to Study'],
    [
        ['Linear', 'linear.app', 'Opinionated UX, speed, visual minimalism \u2014 THE design benchmark'],
        ['Productboard', 'productboard.com', 'Insights-to-roadmap pipeline, role-based views'],
        ['Amplitude', 'amplitude.com', 'AI analytics, automated insights, GenAI querying'],
        ['Dovetail', 'dovetail.com', 'Research repository, auto-clustering, evidence trails'],
    ]
)

H('Collaboration & Workspace', 2)
T(
    ['Platform', 'URL', 'Key Feature to Study'],
    [
        ['Notion', 'notion.so', 'Connected workspace, AI agents, relational data'],
        ['Miro', 'miro.com', 'Strategy templates, visual collaboration, workshops'],
        ['FigJam', 'figma.com/figjam', 'AI template generation, workshop facilitation'],
        ['Loom', 'loom.com', 'Async communication for strategy presentations'],
    ]
)

H('Performance & Goals', 2)
T(
    ['Platform', 'URL', 'Key Feature to Study'],
    [
        ['Lattice', 'lattice.com', 'Goals + performance, cascading alignment, AI agent'],
        ['Asana Goals', 'asana.com/goals', 'Laddering visibility, portfolio views, AI Studio'],
        ['Monday.com', 'monday.com', 'Multi-view work management, connected boards'],
    ]
)

H('Intelligence & Analytics', 2)
T(
    ['Platform', 'URL', 'Key Feature to Study'],
    [
        ['Gong', 'gong.io', 'Behaviour design through data, automatic capture, pattern recognition'],
        ['Clari', 'clari.com', 'Revenue intelligence, forecast accuracy'],
        ['Gainsight', 'gainsight.com', 'Customer success platform, health scores'],
    ]
)

H('Strategy Visualisation Tools', 2)
T(
    ['Platform', 'URL', 'Key Feature to Study'],
    [
        ['Blue Ocean Strategy Canvas', 'blueoceanstrategy.com/tools/strategy-canvas', 'Interactive value curve visualisation'],
        ['Creately', 'creately.com', 'Strategy canvas templates, visual collaboration'],
        ['StratNavApp', 'stratnavapp.com', 'Collaborative Blue Ocean Strategy Canvas'],
        ['Lucidchart', 'lucidchart.com', 'Strategy visualisation, flow diagrams'],
    ]
)

doc.add_page_break()

H('Key Articles & Design References', 2)
T(
    ['Title', 'Source', 'URL', 'Relevance'],
    [
        ['The Linear Method: Opinionated Software', 'Figma Blog', 'figma.com/blog/the-linear-method-opinionated-software', 'Core UX philosophy for Frontera'],
        ['Linear Design: The SaaS UX Trend', 'LogRocket', 'blog.logrocket.com/ux-design/linear-design', 'Minimalism and directional clarity'],
        ['How Linear Hit $1.25B with 100 Employees', 'Aakash Gupta', 'aakashgupta.medium.com', 'Growth through design quality'],
        ['6 Top UX Trends Transforming B2B SaaS', 'SuperUser Studio', 'superuserstudio.com', 'AI integration, guided workflows'],
        ['SaaS Design Trends 2025', 'Good Side', 'goodside.fi', 'Gamification, micro-interactions'],
        ['AI Workshop Facilitation in Miro/FigJam', 'UX Raspberry (Medium)', 'medium.com/@uxraspberry', 'AI-powered collaborative design'],
        ['Enterprise UX Design Trends 2025', 'AufaitUX', 'aufaitux.com/blog/enterprise-ux-design-trends', 'Enterprise micro-interactions'],
        ['Asana Goals 2-Year Vision', 'Catt Small', 'cattsmall.com/design/asana-goals-vision', 'Goals cascade UX design'],
        ['Best SaaS Onboarding Examples', 'InSaim Design', 'insaim.design', 'Progress, checklists, momentum'],
        ['BetterUp Launches AI Coaching', 'BusinessWire', 'businesswire.com', 'Human+AI coaching positioning'],
    ]
)

doc.add_page_break()

# ============================================================================
# 8. CROSS-REFERENCE MAP
# ============================================================================
H('8. Cross-Reference: Challenges \u2192 Inspiration \u2192 Frontera Solution', 1)

para(
    'This final section maps from the 8 core product strategy challenges (Strategy Challenges '
    'Research) through the platform inspiration (this document) to specific Frontera design '
    'decisions. Use this as a design checklist during sprint planning.'
)

T(
    ['Challenge', 'Platform Inspiration', 'Frontera Design Decision', 'Priority'],
    [
        ['#1 No Time', 'Linear (opinionated defaults), BetterUp (in-flow)', 'Express Mode, AI pre-fill, industry templates', 'P0'],
        ['#2 Execution Gap', 'Cascade + Perdoo + Lattice (strategy cascade)', 'Phase 5: Strategic Activation module', 'P0'],
        ['#3 Prioritisation', 'Cascade (strategy map), Productboard (roadmap)', 'Interactive portfolio view, effort/impact scatter', 'P1'],
        ['#4 Research Rigour', 'Dovetail (evidence trails), Amplitude (analytics)', 'Confidence ratings, evidence quality visualisation', 'P1'],
        ['#5 Misalignment', 'Miro/FigJam (collaboration), Gong (capture)', 'Stakeholder input mode, alignment gap detector', 'P1'],
        ['#6 Static Strategy', 'Notion (living workspace), Cascade (health scores)', 'Living Strategy Dashboard, assumption tracker', 'P0'],
        ['#7 Competency', 'BetterUp (personalised coaching), CoachHub (insights)', 'Personal Profile, adaptive coaching depth', 'P1'],
        ['#8 Measuring Impact', 'Perdoo (OKR+KPI), Lattice (goals+performance)', 'OKR cascade, outcome tracking, strategy health', 'P2'],
    ]
)

H('Recommended Build Sequence', 2)
T(
    ['Phase', 'Focus', 'Modules/Features', 'Platform Inspiration'],
    [
        ['Now', 'Strengthen Core', 'PSC enhancements (Express Mode, micro-synthesis, confidence ratings)', 'Linear, BetterUp, Dovetail'],
        ['Next', 'Bridge to Execution', 'Phase 5: Strategic Activation + Living Strategy Dashboard', 'Cascade, Perdoo, Lattice, Notion'],
        ['Then', 'Enrich Coaching', 'Expert Perspectives (UC1) + Sparring Partners (UC2)', 'Gong, BetterUp, Amplitude'],
        ['After', 'Deepen Methodology', 'Case Study Engine (UC3) + Tension Simulator (UC4)', 'Dovetail, Miro, FigJam'],
        ['Later', 'Develop Leaders', 'Leadership Playbook (UC5) + Personal Profile', 'BetterUp, CoachHub, Lattice'],
        ['Ongoing', 'Connect Ecosystem', 'Integrations (Slack, Jira, Notion) + API', 'Perdoo, Lattice, Cascade'],
    ]
)

doc.add_page_break()

# Document control
H('Document Control', 1)
T(
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ['1.0', 'February 2026', 'Frontera Product Team', 'Initial UX journey and platform inspiration research'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('\u2014 End of Document \u2014').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_path = r'c:\Users\deeks\frontera-platform\Background\UX_Journey_Platform_Inspiration_Research.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
