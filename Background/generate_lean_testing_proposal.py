"""
Generate a Word document for the Lean ($1k) Product Strategy Coach User Testing Proposal
Uses Maze (Free) + Ballpark HQ + Self-Recruited Moderated Sessions
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

NAVY = RGBColor(0x1a, 0x1f, 0x3a)
GOLD = RGBColor(0xfb, 0xbf, 0x24)
CYAN = RGBColor(0x08, 0x91, 0xb2)


def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_heading(text, level=1):
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = NAVY
    return heading


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '1a1f3a')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
    doc.add_paragraph()
    return table


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def quote(text):
    doc.add_paragraph(text, style='Intense Quote')


# ============================================================================
# TITLE PAGE
# ============================================================================
title = doc.add_heading('Product Strategy Coach', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(36)
    run.font.color.rgb = NAVY

subtitle = doc.add_paragraph('Lean User Testing Plan')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(24)
subtitle.runs[0].font.color.rgb = CYAN

doc.add_paragraph()

tagline = doc.add_paragraph('Maximum Signal on a $1,000 Budget')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.runs[0].font.size = Pt(16)
tagline.runs[0].italic = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Version: 1.0\n').bold = True
info.add_run('Date: February 2026\n')
info.add_run('Author: Frontera Product Team\n')
info.add_run('Tools: Maze (Free) + Ballpark HQ + Zoom\n')
info.add_run('Budget: ~$1,000\n')
info.add_run('Status: Draft for Review')

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
add_heading('Table of Contents', 1)
toc = [
    ('1.', 'Executive Summary'),
    ('2.', 'Research Objectives'),
    ('3.', 'Testing Strategy & Tool Selection'),
    ('4.', 'Maze A: Discovery & Navigation'),
    ('5.', 'Maze B: Synthesis & Value'),
    ('6.', 'Ballpark Study: Coach Perception & First Impressions'),
    ('7.', 'Moderated Sessions: Full Journey & Coaching Quality'),
    ('8.', 'Participant Recruitment'),
    ('9.', 'Three-Week Timeline'),
    ('10.', 'Success Metrics'),
    ('11.', 'Budget Breakdown'),
    ('12.', 'Appendices'),
]
for num, item in toc:
    p = doc.add_paragraph()
    p.add_run(f'{num} ').bold = True
    p.add_run(item)

doc.add_page_break()

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================
add_heading('1. Executive Summary', 1)

add_heading('Purpose', 2)
doc.add_paragraph(
    'This plan outlines a lean user testing programme for the Product Strategy Coach pilot, '
    'designed to deliver actionable UX insights within a ~$1,000 budget. The approach uses '
    'a deliberate blend of tools to match each research question to the cheapest method that '
    'can answer it reliably.'
)

add_heading('Core Principle: Right Tool for the Right Question', 2)
add_table(
    ['Research Question Type', 'Best Tool', 'Why'],
    [
        ['Can testers navigate the UI?', 'Maze (prototype missions)', 'Usability Score, heatmaps, click paths'],
        ['Do testers understand outputs?', 'Maze (5-Second Test + questions)', 'Comprehension metrics'],
        ['Do testers trust the AI coach?', 'Ballpark (video + ratings)', 'Cheaper for opinion/attitude data'],
        ['What are the first impressions?', 'Ballpark (5-Second Test)', 'Lower cost per tester for surveys'],
        ['Why do testers struggle?', 'Moderated sessions (Zoom)', 'Free tool, deep qualitative insight'],
        ['Is the coaching valuable?', 'Moderated sessions (Zoom)', 'Cannot be tested unmoderated'],
    ]
)

add_heading('Testing Layers', 2)
add_table(
    ['Layer', 'Tool', 'Testers', 'Cost'],
    [
        ['1. Prototype usability', 'Maze (Free plan)', '8-10 per maze', '$0 (tool)'],
        ['2. Attitudes & perception', 'Ballpark HQ', '10-12', '~$50-80 (credits)'],
        ['3. Deep qualitative', 'Zoom (free)', '5-6', '$0 (tool)'],
        ['4. Tester incentives', 'Amazon gift cards', '23-28 total', '~$860'],
        ['', 'TOTAL', '', '~$910-940'],
    ]
)

add_heading('What This Plan Covers', 2)
bullets([
    'Navigation usability across Discovery and Research phases (Maze heatmaps + usability scores)',
    'Synthesis output comprehension and perceived value (Maze 5-Second Tests + opinion scales)',
    'AI coach perception, trust, and interaction quality (Ballpark video review + ratings)',
    'First impressions and product comprehension (Ballpark 5-Second Test)',
    'Full journey experience and coaching quality (5 moderated think-aloud sessions)',
    'Willingness to adopt and pay (moderated session debrief)',
])

add_heading('What This Plan Defers', 2)
bullets([
    'End-to-end unmoderated prototype test (covered by moderated sessions instead)',
    'Expert panel review of synthesis quality (ask 1-2 advisors informally)',
    'Large-sample quantitative data (prioritise depth over breadth)',
    'NPS/willingness-to-pay survey (ask in moderated debrief)',
])

doc.add_page_break()

# ============================================================================
# 2. RESEARCH OBJECTIVES
# ============================================================================
add_heading('2. Research Objectives', 1)

add_heading('Primary Objectives', 2)
add_table(
    ['Objective', 'Success Criteria', 'Measured By'],
    [
        ['Validate navigation clarity', 'Usability Score ≥75 on Maze missions', 'Maze A'],
        ['Confirm methodology comprehension', '>70% direct success on key missions', 'Maze A + B'],
        ['Assess coach interaction value', 'Trust rating ≥4/5', 'Ballpark + Moderated'],
        ['Evaluate synthesis output value', 'Actionability rating ≥4/5', 'Maze B + Moderated'],
        ['Identify critical friction points', 'Misclick rate <15%', 'Maze heatmaps'],
    ]
)

add_heading('Hypotheses', 2)
bullets([
    'H1: Testers will find the document upload function within 30 seconds (>75% direct success)',
    'H2: The two-panel layout (coach + canvas) will be understood without instruction',
    'H3: Testers will trust AI coaching enough to engage voluntarily (trust ≥4/5)',
    'H4: Synthesis outputs will be perceived as actionable by product leaders',
    'H5: The terrain mapping metaphor will be intuitive (clarity ≥7/10)',
])

doc.add_page_break()

# ============================================================================
# 3. TESTING STRATEGY & TOOL SELECTION
# ============================================================================
add_heading('3. Testing Strategy & Tool Selection', 1)

add_heading('Why a Blended Approach?', 2)
doc.add_paragraph(
    'A single tool cannot cost-effectively answer all research questions. This plan deliberately '
    'separates behavioural testing (what testers do) from attitudinal testing (what testers think) '
    'and uses the cheapest appropriate tool for each.'
)

add_heading('Tool 1: Maze (Free Plan)', 2)
doc.add_paragraph('Used for: Prototype missions where usability metrics matter')

p = doc.add_paragraph()
p.add_run('Free plan includes:\n').bold = True
bullets([
    '1 active maze at a time (run sequentially)',
    'Up to 10 blocks per maze',
    'Prototype test blocks (goal-based and free explore)',
    '5-Second Test blocks',
    'Question blocks (Opinion Scale, Multiple Choice, Open Question)',
    'Usability Score calculation (0-100)',
    'Heatmaps and click path analysis',
    'Direct/Indirect Success and Bounce tracking',
    'Misclick rate measurement',
])

p = doc.add_paragraph()
p.add_run('Why Maze for prototype testing:\n').bold = True
bullets([
    'Automated Usability Score removes subjective bias',
    'Heatmaps show exactly where testers click (and mis-click)',
    'Expected Path vs. actual path reveals navigation confusion',
    'Time on Task benchmarks identify slow/frustrating flows',
])

p = doc.add_paragraph()
p.add_run('Limitation: ').bold = True
doc.add_paragraph(
    'Free plan restricts to 1 maze at a time. This plan runs two mazes sequentially '
    'across Week 1 (Maze A on Days 1-3, Maze B on Days 4-5).'
)

add_heading('Tool 2: Ballpark HQ', 2)
doc.add_paragraph('Used for: Opinion/attitude data where prototype interaction is not needed')

p = doc.add_paragraph()
p.add_run('Key features used:\n').bold = True
bullets([
    '5-Second Tests for first impressions',
    'Rating scale questions for trust and perception',
    'Video walkthrough review (show coach interaction recording)',
    'Open text responses for qualitative feedback',
    'Multiple choice for categorisation',
])

p = doc.add_paragraph()
p.add_run('Why Ballpark for attitudinal testing:\n').bold = True
bullets([
    'Per-credit pricing means short studies cost $50-80 (not $99+/month)',
    'Ideal for showing screenshots/videos and capturing reactions',
    'Simpler setup for non-prototype research',
    'Complements Maze behavioural data with attitudinal data',
])

add_heading('Tool 3: Zoom (Free)', 2)
doc.add_paragraph('Used for: Moderated think-aloud sessions')

p = doc.add_paragraph()
p.add_run('Why moderated sessions are non-negotiable:\n').bold = True
bullets([
    'AI coaching quality cannot be tested unmoderated (requires real interaction)',
    'Full journey evaluation needs real company context',
    'Trust and value perception require probing follow-up questions',
    'The "why" behind friction points only emerges through observation',
    'Nielsen\'s research: 5 users find 85% of usability issues',
])

add_heading('What Goes Where', 2)
add_table(
    ['Test Content', 'Tool', 'Rationale'],
    [
        ['Document upload navigation', 'Maze', 'Needs click tracking + success rate'],
        ['Territory card navigation', 'Maze', 'Needs heatmap + path analysis'],
        ['Research question interaction', 'Maze', 'Needs free explore observation'],
        ['Synthesis output comprehension', 'Maze', 'Needs 5-Second Test + mission'],
        ['First interface impression', 'Ballpark', 'Cheaper for screenshot + opinion'],
        ['Coach tone perception', 'Ballpark', 'Video review + rating (no prototype)'],
        ['Coach trust assessment', 'Ballpark', 'Scenario-based rating (no prototype)'],
        ['Trust-building behaviours', 'Ballpark', 'Multiple choice survey (cheaper)'],
        ['Full journey experience', 'Moderated', 'Requires real context + probing'],
        ['Coaching quality assessment', 'Moderated', 'Cannot test unmoderated'],
        ['Adoption likelihood', 'Moderated', 'Requires conversation, not survey'],
        ['Willingness to pay', 'Moderated', 'Debrief question (free)'],
    ]
)

doc.add_page_break()

# ============================================================================
# 4. MAZE A: DISCOVERY & NAVIGATION
# ============================================================================
add_heading('4. Maze A: Discovery & Navigation', 1)

doc.add_paragraph('Objective: Validate that testers can navigate the Discovery and Research phases')
p = doc.add_paragraph()
p.add_run('Duration: ').bold = True
p.add_run('~8 minutes')
p = doc.add_paragraph()
p.add_run('Testers: ').bold = True
p.add_run('8-10 (self-recruited)')
p = doc.add_paragraph()
p.add_run('Run: ').bold = True
p.add_run('Week 1, Days 1-3')

add_heading('Block Structure', 2)
add_table(
    ['Block #', 'Type', 'Content', 'Key Metric'],
    [
        ['1', 'Welcome Block', 'Scenario: VP Product at fintech, first login', '-'],
        ['2', 'Prototype Test (Goal)', 'Mission: Find document upload', 'Direct Success, Misclick Rate'],
        ['3', 'Opinion Scale (0-10)', 'How easy was it to find the upload?', 'Mean score'],
        ['4', 'Prototype Test (Goal)', 'Mission: Respond to coach question', 'Engagement rate'],
        ['5', 'Opinion Scale (0-10)', 'How relevant was the coach\'s question?', 'Mean score'],
        ['6', 'Prototype Test (Goal)', 'Mission: Navigate to Customer territory', 'Direct Success, Heatmap'],
        ['7', 'Prototype Test (Goal)', 'Mission: Find Segments & Needs area', 'Direct Success, Time on Task'],
        ['8', 'Prototype Test (Free Explore)', 'Explore research questions (60 seconds)', 'Click paths, Time distribution'],
        ['9', 'Opinion Scale (0-10)', 'How clear were the research questions?', 'Mean score'],
        ['10', 'Open Question', 'What was confusing or frustrating?', 'Qualitative themes'],
    ]
)

add_heading('Block Details', 2)

quote('Block 1: Welcome Block')
doc.add_paragraph(
    'Scenario text: "You\'re the VP of Product at a mid-sized fintech company. Your CEO has asked you '
    'to develop a product strategy for the next 18 months. You\'ve decided to use this AI coaching '
    'tool to help structure your thinking. You\'ve just logged in for the first time."'
)

quote('Block 2: Prototype Test (Goal-Based) — Find Upload')
p = doc.add_paragraph()
p.add_run('Mission: ').bold = True
p.add_run('"Find where you would upload your company\'s strategic documents."')
doc.add_paragraph()
add_table(
    ['Setting', 'Value'],
    [
        ['Expected Path', 'Landing → Discovery Section → Source Materials → Upload Button'],
        ['Goal Screen', 'Upload button focused or drag-drop zone visible'],
        ['Time Limit', '30 seconds'],
        ['Hotspots', 'Upload Files button, drag-drop zone, Add Link button'],
    ]
)

p = doc.add_paragraph()
p.add_run('What to look for in results:\n').bold = True
bullets([
    'Direct Success >75%: Testers find upload without wrong turns',
    'Indirect Success <15%: Few testers use unexpected routes',
    'Misclick heatmap: Where do testers click instead of upload?',
    'Time on Task <15s: Upload should be immediately discoverable',
])

quote('Block 4: Prototype Test (Goal-Based) — Coach Interaction')
p = doc.add_paragraph()
p.add_run('Mission: ').bold = True
p.add_run('"The AI coach has sent you a message. Find and respond to it."')
doc.add_paragraph()
add_table(
    ['Setting', 'Value'],
    [
        ['Context Screen', 'Coach message: "What competitive dynamics are making product transformation urgent?"'],
        ['Goal Screen', 'Input field focused or response submitted'],
        ['Hotspots', 'Coach input field, send button'],
    ]
)

p = doc.add_paragraph()
p.add_run('What to look for:\n').bold = True
bullets([
    'Do testers notice the coach sidebar?',
    'Do testers understand they can type a response?',
    'Heatmap: Do clicks concentrate on input area or wander?',
])

quote('Block 6: Prototype Test (Goal-Based) — Customer Territory')
p = doc.add_paragraph()
p.add_run('Mission: ').bold = True
p.add_run('"You want to understand your customers better. Navigate to the Customer territory."')
doc.add_paragraph()
add_table(
    ['Setting', 'Value'],
    [
        ['Expected Path', 'Research Phase → Customer Territory Card → Click'],
        ['Goal Screen', 'Customer territory deep-dive view'],
        ['Hotspots', 'Customer territory card'],
    ]
)

p = doc.add_paragraph()
p.add_run('What to look for:\n').bold = True
bullets([
    'Do testers select the correct territory card on first click?',
    'Heatmap: Even distribution across 3 cards or confusion between them?',
    'Misclicks: Do testers click territory labels vs. cards?',
])

quote('Block 8: Prototype Test (Free Explore) — Research Questions')
p = doc.add_paragraph()
p.add_run('Instruction: ').bold = True
p.add_run('"Take 60 seconds to explore the research questions for this area. Click around naturally."')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Observe natural navigation without a defined goal. Reveals what testers gravitate toward, '
          'what they skip, and whether the question structure is intuitive.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('What to look for:\n').bold = True
bullets([
    'Do testers scroll through all questions or stop early?',
    'Do testers click on input fields (intent to answer)?',
    'Time distribution: Which areas get the most attention?',
    'Do testers interact with Save/Mark as Mapped buttons?',
])

doc.add_page_break()

# ============================================================================
# 5. MAZE B: SYNTHESIS & VALUE
# ============================================================================
add_heading('5. Maze B: Synthesis & Value', 1)

doc.add_paragraph('Objective: Assess comprehension and perceived value of synthesis outputs')
p = doc.add_paragraph()
p.add_run('Duration: ').bold = True
p.add_run('~7 minutes')
p = doc.add_paragraph()
p.add_run('Testers: ').bold = True
p.add_run('8-10 (self-recruited)')
p = doc.add_paragraph()
p.add_run('Run: ').bold = True
p.add_run('Week 1, Days 4-5 (after Maze A closes)')

add_heading('Block Structure', 2)
add_table(
    ['Block #', 'Type', 'Content', 'Key Metric'],
    [
        ['1', 'Welcome Block', 'Context: Research complete, reviewing synthesis', '-'],
        ['2', '5-Second Test', 'Strategic Opportunity Map (2x2 matrix)', 'Recall accuracy'],
        ['3', 'Open Question', 'What does this visualisation show?', 'Comprehension themes'],
        ['4', 'Multiple Choice', 'What does the INVEST quadrant mean?', '% correct'],
        ['5', 'Prototype Test (Goal)', 'Mission: Find top opportunity details', 'Direct Success, Time'],
        ['6', 'Opinion Scale (0-10)', 'How actionable is this output?', 'Mean score'],
        ['7', 'Opinion Scale (0-10)', 'How valuable is the evidence trail?', 'Mean score'],
        ['8', 'Multiple Choice', 'Could you act on this information?', 'Response distribution'],
        ['9', 'Opinion Scale (0-10)', 'Would you share this with your exec team?', 'Mean score'],
        ['10', 'Open Question', 'What is missing that you would need?', 'Qualitative themes'],
    ]
)

add_heading('Block Details', 2)

quote('Block 2: 5-Second Test — Opportunity Map')
p = doc.add_paragraph()
p.add_run('Image: ').bold = True
p.add_run('Strategic Opportunity Map (2x2 matrix) with 4 plotted opportunity bubbles')
doc.add_paragraph()
add_table(
    ['Element', 'Detail'],
    [
        ['X-Axis', 'Capability Fit (Low → High)'],
        ['Y-Axis', 'Market Attractiveness (Low → High)'],
        ['Quadrants', 'INVEST (top-right), EXPLORE (top-left), HARVEST (bottom-right), DIVEST (bottom-left)'],
        ['Bubbles', '4 opportunities plotted by scores'],
    ]
)
p = doc.add_paragraph()
p.add_run('What to look for:\n').bold = True
bullets([
    'Can testers describe the visualisation purpose after 5 seconds?',
    'Do they recognise it as a prioritisation framework?',
    'Do responses mention "opportunity" or "priority" language?',
])

quote('Block 4: Multiple Choice — Quadrant Comprehension')
doc.add_paragraph('Question: "What does an opportunity in the top-right quadrant (INVEST) mean?"')
doc.add_paragraph('Options:')
bullets([
    'High priority - strong market fit and capability alignment (Correct)',
    'Interesting but needs more research',
    'Low priority - deprioritise',
    'I\'m not sure',
])
p = doc.add_paragraph()
p.add_run('Target: ').bold = True
p.add_run('>75% select the correct answer')

quote('Block 5: Prototype Test (Goal-Based) — Find Opportunity Details')
p = doc.add_paragraph()
p.add_run('Mission: ').bold = True
p.add_run('"Find the details for the highest-priority strategic opportunity."')
doc.add_paragraph()
add_table(
    ['Setting', 'Value'],
    [
        ['Expected Path', 'Map → Click INVEST bubble → Opportunity Card → View Details'],
        ['Goal Screen', 'Expanded opportunity card with evidence trail'],
        ['Hotspots', 'INVEST quadrant bubbles, opportunity card expand button'],
    ]
)
p = doc.add_paragraph()
p.add_run('What to look for:\n').bold = True
bullets([
    'Do testers understand the map is interactive (clickable)?',
    'Do they navigate from map → card → details intuitively?',
    'Time on Task: How long to find detail view?',
])

quote('Block 8: Multiple Choice — Actionability Assessment')
doc.add_paragraph('Question: "Based on this opportunity card, could you take action?"')
doc.add_paragraph('Options:')
bullets([
    'Yes, I have enough information to act',
    'Partially, but I need more detail on some areas',
    'No, I need significantly more information',
    'I\'m not sure what action to take',
])

doc.add_page_break()

# ============================================================================
# 6. BALLPARK STUDY
# ============================================================================
add_heading('6. Ballpark Study: Coach Perception & First Impressions', 1)

doc.add_paragraph('Objective: Capture attitudinal data on first impressions and coaching perception')
p = doc.add_paragraph()
p.add_run('Duration: ').bold = True
p.add_run('~5 minutes')
p = doc.add_paragraph()
p.add_run('Testers: ').bold = True
p.add_run('10-12 (self-recruited)')
p = doc.add_paragraph()
p.add_run('Run: ').bold = True
p.add_run('Week 1, parallel with Maze A')
p = doc.add_paragraph()
p.add_run('Est. Cost: ').bold = True
p.add_run('~$50-80 (credits-based)')

add_heading('Study Structure', 2)
add_table(
    ['Section', 'Type', 'Content', 'Data Collected'],
    [
        ['1', '5-Second Test', 'Full interface screenshot (coach + canvas)', 'First impression recall'],
        ['2', 'Open Text', '"What do you think this product does?"', 'Comprehension accuracy'],
        ['3', 'Rating (1-5)', '"How clear is the purpose of this tool?"', 'Value clarity score'],
        ['4', 'Rating (1-5)', '"How relevant does this seem for your work?"', 'Relevance score'],
        ['5', 'Video Review', 'Watch 60s coach interaction recording', 'Observation'],
        ['6', 'Rating (1-5)', '"How helpful was the coach\'s response?"', 'Helpfulness score'],
        ['7', 'Rating (1-5)', '"How much would you trust this AI coach?"', 'Trust score'],
        ['8', 'Multiple Choice', '"How would you describe the interaction style?"', 'Perception category'],
        ['9', 'Multiple Choice (Multi)', '"What would increase your trust in the coach?"', 'Trust driver ranking'],
        ['10', 'Open Text', '"What would make the coaching more valuable?"', 'Qualitative themes'],
    ]
)

add_heading('Section Details', 2)

quote('Section 1: 5-Second Test — First Interface Impression')
p = doc.add_paragraph()
p.add_run('Image: ').bold = True
p.add_run('Full Product Strategy Coach interface showing coach sidebar (left) and canvas (right) '
          'with Discovery phase visible.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Why Ballpark, not Maze: ').bold = True
p.add_run('No prototype interaction needed. Ballpark\'s 5-Second Test is cheaper and simpler for '
          'capturing recall from a static image.')

quote('Section 5: Video Review — Coach Interaction')
p = doc.add_paragraph()
p.add_run('Video: ').bold = True
p.add_run('60-second screen recording showing:')
bullets([
    'User types a question about competitive positioning',
    'Coach provides a contextual, probing response',
    'Coach references earlier context from the session',
    'Coach ends with a follow-up question',
])
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Why Ballpark, not Maze: ').bold = True
p.add_run('This tests perception of coaching quality, not navigation. Testers watch a video '
          'and rate it. No prototype interaction needed, so Ballpark\'s per-credit pricing '
          'is significantly cheaper.')

quote('Section 8: Multiple Choice — Interaction Style Perception')
doc.add_paragraph('Question: "How would you describe the interaction between the user and the coach?"')
doc.add_paragraph('Options:')
bullets([
    'Like talking to a knowledgeable consultant',
    'Like using a sophisticated chatbot',
    'Like following a guided tutorial',
    'Like having a conversation with a colleague',
    'Like receiving automated suggestions',
])
p = doc.add_paragraph()
p.add_run('Key insight: ').bold = True
p.add_run('We want "consultant" or "colleague" responses. "Chatbot" or "automated" signals a trust problem.')

quote('Section 9: Multiple Choice (Multi-Select) — Trust Drivers')
doc.add_paragraph('Question: "What would make you trust the AI coach more? Select all that apply."')
doc.add_paragraph('Options:')
bullets([
    'Showing its reasoning and sources',
    'Asking clarifying questions before advising',
    'Admitting when it doesn\'t know something',
    'Providing options rather than single recommendations',
    'Referencing industry benchmarks or data',
    'Remembering context from earlier in the session',
])
p = doc.add_paragraph()
p.add_run('Key insight: ').bold = True
p.add_run('Ranking of trust drivers directly informs product roadmap priorities for the coaching agent.')

doc.add_page_break()

# ============================================================================
# 7. MODERATED SESSIONS
# ============================================================================
add_heading('7. Moderated Sessions: Full Journey & Coaching Quality', 1)

doc.add_paragraph(
    'Moderated sessions are the highest-value research activity in this plan. They answer '
    'questions that unmoderated tools cannot: Why do testers struggle? Do they trust the coach? '
    'Would they use this for a real project?'
)

add_heading('Session Configuration', 2)
add_table(
    ['Element', 'Detail'],
    [
        ['Duration', '45-60 minutes per session'],
        ['Format', 'Video call with screen sharing via Zoom (free tier)'],
        ['Recording', 'Zoom recording (with consent)'],
        ['Testers', '5-6 product/strategy leaders'],
        ['Recruitment', 'Self-recruited (see Section 8)'],
        ['Incentive', '$100 Amazon gift card per tester'],
        ['Run', 'Week 2'],
    ]
)

add_heading('Session Structure', 2)
add_table(
    ['Section', 'Duration', 'Purpose'],
    [
        ['1. Introduction', '5 min', 'Set expectations, consent to record'],
        ['2. Context Setting', '5 min', 'Tester background and current challenges'],
        ['3. First Impressions', '5 min', 'Initial reaction (no guidance)'],
        ['4. Discovery Phase', '10 min', 'Upload materials, engage coach'],
        ['5. Research Phase', '15 min', 'Complete 1-2 territories'],
        ['6. Synthesis Review', '5-10 min', 'Review generated outputs'],
        ['7. Debrief Interview', '10 min', 'Structured feedback'],
    ]
)

add_heading('Facilitator Script: Introduction', 3)
script = doc.add_paragraph()
script.add_run(
    '"Thank you for joining today. I\'m [name], and I\'ll be facilitating this session.\n\n'
    'We\'re developing an AI-powered tool to help product leaders develop strategy. Today, I\'d like '
    'you to use the tool and share your thoughts as you go.\n\n'
    'A few things before we start:\n'
    '• Please think aloud - share what you\'re looking at, thinking, and feeling\n'
    '• There are no wrong answers - if something is confusing, that\'s valuable feedback\n'
    '• We\'re testing the product, not you\n'
    '• Feel free to be candid - honest criticism helps us improve\n\n'
    'Do you have any questions before we begin?"'
)
script.italic = True

add_heading('During-Session Observation Prompts', 2)
doc.add_paragraph('Use these when the tester goes silent or shows non-verbal reactions:')
add_table(
    ['Situation', 'Prompt'],
    [
        ['Tester pauses', '"What are you looking at right now?"'],
        ['Tester seems confused', '"What were you expecting to see?"'],
        ['Tester clicks unexpectedly', '"What made you click there?"'],
        ['Tester skips content', '"I noticed you scrolled past that - was it not relevant?"'],
        ['Tester shows frustration', '"What\'s causing that reaction?"'],
        ['Tester shows satisfaction', '"What do you like about this?"'],
        ['Tester ignores coach', '"I noticed you didn\'t engage with the coach message - any reason?"'],
        ['Tester speeds up', '"You seem more confident now - what changed?"'],
    ]
)

add_heading('Debrief Interview Questions', 2)

quote('Overall Experience (2 min)')
bullets([
    '"In a few words, how would you describe this experience?"',
    '"What was the single most valuable moment?"',
    '"What was the single most frustrating moment?"',
])

quote('Coaching Quality (3 min)')
bullets([
    '"How would you rate the AI coach on a scale of 1-10?"',
    '"Did the coach feel more like a partner, a tool, or something else?"',
    '"Were there moments when the coach was particularly helpful?"',
    '"What would make you trust the coach more?"',
    '"How does this compare to working with a human strategy consultant?"',
])

quote('Output Value (2 min)')
bullets([
    '"How actionable were the strategic insights?"',
    '"Would you present this synthesis to your leadership team? Why or why not?"',
    '"What additional outputs would you want?"',
])

quote('Adoption & Value (3 min)')
bullets([
    '"Would you use this tool for a real strategic project?"',
    '"What would need to change for you to use it regularly?"',
    '"Would your company pay for a tool like this? What would justify the cost?"',
    '"Who else in your organisation would benefit from this?"',
])

add_heading('Observation Template', 2)
doc.add_paragraph('Complete within 24 hours of each session:')
items = [
    'Tester ID / Date / Role / Company type',
    'Initial reaction (first 30 seconds)',
    'Discovery: Upload ease, coach engagement, friction points',
    'Research: Territory selection reasoning, question comprehension, input depth',
    'Synthesis: Comprehension level, trust signals, gaps identified',
    'Coach: Helpfulness moments, unhelpfulness moments, trust level (1-5)',
    'Outputs: Actionability perception, shareability, missing elements',
    'Overall: Primary value perceived, primary concern, adoption likelihood (1-5)',
    'Key quotes (verbatim, 3-5 per session)',
]
for item in items:
    doc.add_paragraph(f'☐ {item}')

doc.add_page_break()

# ============================================================================
# 8. PARTICIPANT RECRUITMENT
# ============================================================================
add_heading('8. Participant Recruitment', 1)

add_heading('Strategy: Self-Recruit to Save $500+', 2)
doc.add_paragraph(
    'The most expensive part of research is recruitment. Panel services (Maze Panel, UserTesting, '
    'Respondent) charge $3-30 per tester. By self-recruiting from your network, you redirect that '
    'budget to incentives, which improves completion rates and tester quality.'
)

add_heading('Target Profile', 2)
add_table(
    ['Attribute', 'Requirement'],
    [
        ['Job Title', 'VP Product, Head of Strategy, Product Director, Senior PM'],
        ['Company Size', '200+ employees (or funded growth-stage startup)'],
        ['Experience', '5+ years in product or strategy roles'],
        ['Current Context', 'Involved in or planning strategic initiatives'],
        ['AI Comfort', 'At least "somewhat comfortable" with AI tools'],
    ]
)

add_heading('Recruitment Channels (Free)', 2)

quote('Warm Network (Highest Conversion)')
bullets([
    'Existing Frontera contacts (onboarding pipeline, waitlist, beta interest)',
    'LinkedIn connections in product/strategy roles',
    'Ask existing testers to refer one peer (snowball sampling)',
    'Advisory board or investor network introductions',
])

quote('Communities (Medium Conversion)')
bullets([
    'LinkedIn posts targeting Product/Strategy leaders',
    'Mind the Product Slack community',
    'Product School community',
    'Lenny\'s Newsletter Slack/Discord',
    'Product-Led Alliance community',
    'Industry-specific Slack groups (fintech, SaaS, healthtech)',
])

quote('Recruitment Message Template')
msg = doc.add_paragraph()
msg.add_run(
    '"Hi [Name],\n\n'
    'We\'re building an AI-powered strategy coaching tool for product leaders and looking for '
    'feedback from experienced practitioners.\n\n'
    'Would you have [10 minutes / 45 minutes] to try the tool and share your thoughts? '
    'We\'re offering a [$20 / $100] Amazon gift card as a thank you.\n\n'
    'No preparation needed - just your experience and honest opinion.\n\n'
    'Interested? [Link to schedule / reply to confirm]"'
)
msg.italic = True

add_heading('Screening Questions (Quick)', 2)
doc.add_paragraph('Ask via email or brief form before scheduling:')
bullets([
    'What is your current role? (Must include Product or Strategy in title)',
    'Have you contributed to product strategy in the past 12 months? (Must be Yes)',
    'Company size? (Prefer 200+)',
])

add_heading('Sample Size Summary', 2)
add_table(
    ['Study', 'Minimum', 'Target', 'Recruitment Method'],
    [
        ['Maze A (Discovery)', '8', '10', 'Network + communities'],
        ['Maze B (Synthesis)', '8', '10', 'Network + communities'],
        ['Ballpark (Coach)', '10', '12', 'Network + communities'],
        ['Moderated Sessions', '5', '6', 'Warm network (priority)'],
    ]
)

add_heading('Incentive Structure', 2)
add_table(
    ['Study', 'Duration', 'Incentive', 'Delivery'],
    [
        ['Maze A (unmoderated)', '~8 min', '$20 Amazon gift card', 'Email after completion'],
        ['Maze B (unmoderated)', '~7 min', '$20 Amazon gift card', 'Email after completion'],
        ['Ballpark (unmoderated)', '~5 min', '$20 Amazon gift card', 'Email after completion'],
        ['Moderated Session', '~45-60 min', '$100 Amazon gift card', 'Email within 24 hours'],
    ]
)

doc.add_page_break()

# ============================================================================
# 9. THREE-WEEK TIMELINE
# ============================================================================
add_heading('9. Three-Week Timeline', 1)

add_heading('Overview', 2)
add_table(
    ['Week', 'Focus', 'Activities', 'Deliverable'],
    [
        ['1', 'Unmoderated Testing', 'Maze A → Maze B + Ballpark', 'Quantitative data + heatmaps'],
        ['2', 'Moderated Sessions', '5-6 think-aloud sessions', 'Recordings + observation notes'],
        ['3', 'Analysis & Report', 'Synthesise all findings', 'Final report + recommendations'],
    ]
)

add_heading('Week 1: Unmoderated Testing', 2)
add_table(
    ['Day', 'Activity', 'Owner'],
    [
        ['Mon', 'Launch Maze A: Discovery & Navigation', 'UX Researcher'],
        ['Mon', 'Launch Ballpark Study: Coach & First Impressions (parallel)', 'UX Researcher'],
        ['Mon-Wed', 'Maze A collects responses (8-10 testers)', '-'],
        ['Mon-Wed', 'Ballpark collects responses (10-12 testers)', '-'],
        ['Wed', 'Close Maze A, preliminary review of results', 'UX Researcher'],
        ['Thu', 'Launch Maze B: Synthesis & Value', 'UX Researcher'],
        ['Thu-Fri', 'Maze B collects responses (8-10 testers)', '-'],
        ['Fri', 'Close Ballpark, Close Maze B', 'UX Researcher'],
        ['Fri', 'Review all Week 1 data, prepare moderated guide', 'UX Researcher'],
    ]
)

add_heading('Week 2: Moderated Sessions', 2)
add_table(
    ['Day', 'Activity', 'Owner'],
    [
        ['Mon', 'Sessions 1-2 (back-to-back, 45 min each)', 'UX Researcher'],
        ['Tue', 'Session 3', 'UX Researcher'],
        ['Wed', 'Sessions 4-5', 'UX Researcher'],
        ['Thu', 'Session 6 (if needed)', 'UX Researcher'],
        ['Thu-Fri', 'Complete observation notes, tag key moments', 'UX Researcher'],
        ['Fri', 'Identify emerging themes across sessions', 'UX Researcher'],
    ]
)

add_heading('Week 3: Analysis & Reporting', 2)
add_table(
    ['Day', 'Activity', 'Owner'],
    [
        ['Mon', 'Compile quantitative data (Maze + Ballpark)', 'UX Researcher'],
        ['Tue', 'Affinity mapping of qualitative themes', 'UX Researcher'],
        ['Wed', 'Draft findings report with prioritised recommendations', 'UX Researcher'],
        ['Thu', 'Stakeholder presentation', 'Product + UX'],
        ['Fri', 'Prioritisation workshop and sprint planning', 'Product + Design + Engineering'],
    ]
)

doc.add_page_break()

# ============================================================================
# 10. SUCCESS METRICS
# ============================================================================
add_heading('10. Success Metrics', 1)

add_heading('Quantitative Targets', 2)
add_table(
    ['Metric', 'Target', 'Source', 'Priority'],
    [
        ['Usability Score (Maze)', '≥75', 'Maze A + B', 'P0'],
        ['Direct Success Rate', '≥70%', 'Maze missions', 'P0'],
        ['Bounce Rate', '<10%', 'Maze missions', 'P0'],
        ['Misclick Rate', '<15%', 'Maze heatmaps', 'P1'],
        ['Value Clarity (Ballpark)', '≥4.0/5', 'Ballpark rating', 'P0'],
        ['Coach Trust (Ballpark)', '≥4.0/5', 'Ballpark rating', 'P0'],
        ['Coach Helpfulness (Ballpark)', '≥4.0/5', 'Ballpark rating', 'P0'],
        ['Synthesis Actionability (Maze)', '≥7.0/10', 'Maze opinion scale', 'P0'],
        ['Stakeholder Shareability (Maze)', '≥7.0/10', 'Maze opinion scale', 'P1'],
    ]
)

add_heading('Risk Thresholds', 2)
add_table(
    ['Metric', 'Acceptable', 'Warning', 'Critical Action'],
    [
        ['Usability Score', '≥75', '60-74', '<60 → Redesign navigation flow'],
        ['Direct Success', '≥70%', '50-69%', '<50% → Rework information architecture'],
        ['Bounce Rate', '<10%', '10-25%', '>25% → Simplify task flow urgently'],
        ['Misclick Rate', '<15%', '15-25%', '>25% → Revise hotspot sizing/placement'],
        ['Coach Trust', '≥4.0/5', '3.0-3.9', '<3.0 → Rethink coaching interaction model'],
    ]
)

add_heading('Qualitative Success Criteria', 2)
add_table(
    ['Criterion', 'Evidence Required', 'Source'],
    [
        ['Clear value proposition', 'Testers articulate value unprompted', 'Ballpark open text + Moderated'],
        ['Intuitive navigation', '<2 "lost" moments per moderated session', 'Moderated observation'],
        ['Coach perceived as partner', '"Consultant" or "colleague" in >50% of responses', 'Ballpark multiple choice'],
        ['Outputs are actionable', 'Testers describe specific next steps', 'Moderated debrief'],
        ['Trust established', 'Testers willing to share outputs externally', 'Maze opinion scale + Moderated'],
    ]
)

doc.add_page_break()

# ============================================================================
# 11. BUDGET BREAKDOWN
# ============================================================================
add_heading('11. Budget Breakdown', 1)

add_heading('Tool Costs', 2)
add_table(
    ['Item', 'Cost', 'Notes'],
    [
        ['Maze (Free plan)', '$0', '1 maze at a time, 10 blocks, full metrics'],
        ['Ballpark HQ (credits)', '$50-80', '~50-80 credits for 10-12 testers × 5 min study'],
        ['Zoom (Free tier)', '$0', '40-min limit per call (restart if needed)'],
        ['Tool subtotal', '$50-80', ''],
    ]
)

add_heading('Tester Incentives', 2)
add_table(
    ['Study', 'Testers', 'Incentive Each', 'Subtotal'],
    [
        ['Maze A: Discovery & Navigation', '10', '$20', '$200'],
        ['Maze B: Synthesis & Value', '10', '$20', '$200'],
        ['Ballpark: Coach & Impressions', '12', '$0*', '$0'],
        ['Moderated Sessions', '5', '$100', '$500'],
        ['Incentive subtotal', '', '', '$900'],
    ]
)
doc.add_paragraph(
    '* Ballpark testers overlap with Maze testers where possible. '
    'If separate testers needed, add $20 × 12 = $240.'
)

add_heading('Total Budget', 2)
add_table(
    ['Category', 'Low Estimate', 'High Estimate'],
    [
        ['Tools (Maze + Ballpark)', '$50', '$80'],
        ['Incentives (overlapping testers)', '$900', '$900'],
        ['Incentives (separate Ballpark testers)', '-', '+$240'],
        ['TOTAL', '$950', '$1,220'],
    ]
)

add_heading('Budget Optimisation Tips', 2)
bullets([
    'Overlap testers: Ask Maze testers to also complete the Ballpark study (saves 12 × $20 = $240)',
    'Warm network first: Colleagues and contacts may participate without incentives',
    'Zoom free tier: 40-minute limit is tight for 45-min sessions. Use Zoom Pro ($13/mo) or Google Meet (free, no limit) as backup',
    'Record with consent: Free Zoom recording avoids needing separate tools',
    'Defer expert review: Ask 1-2 strategy advisors to review outputs informally (free)',
])

doc.add_page_break()

# ============================================================================
# 12. APPENDICES
# ============================================================================
add_heading('12. Appendices', 1)

add_heading('Appendix A: Terminology Quick Reference', 2)
add_table(
    ['Term', 'Tool', 'Meaning'],
    [
        ['Maze', 'Maze', 'A complete test containing multiple blocks'],
        ['Block', 'Maze', 'Individual test component (prototype test, question, etc.)'],
        ['Mission', 'Maze', 'A specific task for testers to complete in a prototype'],
        ['Expected Path', 'Maze', 'The ideal route through screens to complete a mission'],
        ['Goal Screen', 'Maze', 'The target destination screen for a mission'],
        ['Hotspot', 'Maze', 'Defined clickable area on a screen'],
        ['Direct Success', 'Maze', 'Completed mission via expected path'],
        ['Indirect Success', 'Maze', 'Completed mission via unexpected path'],
        ['Bounce', 'Maze', 'Tester abandoned the mission'],
        ['Usability Score', 'Maze', 'Composite score (0-100)'],
        ['Misclick', 'Maze', 'Click outside defined hotspots'],
        ['Heatmap', 'Maze', 'Aggregated visualisation of tester clicks'],
        ['Tester', 'Maze', 'Participant taking the maze'],
        ['Study', 'Ballpark', 'A complete test/survey in Ballpark'],
        ['Credit', 'Ballpark', 'Pricing unit (1 credit ≈ 1 tester-minute)'],
        ['5-Second Test', 'Both', 'Brief exposure to capture first impressions'],
    ]
)

add_heading('Appendix B: Moderated Session Checklist', 2)

quote('Pre-Session')
items = [
    'Live prototype URL confirmed working',
    'Zoom meeting link created and shared',
    'Recording permission prepared',
    'Tester background reviewed',
    'Observation template ready',
    'Backup scenario prepared (if tester has no real docs)',
    'Gift card ready to send post-session',
]
for item in items:
    doc.add_paragraph(f'☐ {item}')

quote('During Session')
items = [
    'Welcome and introductions',
    'Confirm consent to record',
    'Start Zoom recording',
    'Context-setting questions (5 min)',
    'First impressions - no guidance (5 min)',
    'Discovery phase walkthrough (10 min)',
    'Research phase walkthrough (15 min)',
    'Synthesis review (5-10 min)',
    'Debrief interview (10 min)',
    'Thank tester, confirm gift card delivery',
    'Stop recording',
]
for item in items:
    doc.add_paragraph(f'☐ {item}')

quote('Post-Session (within 24 hours)')
items = [
    'Complete observation template',
    'Highlight 3-5 key quotes (verbatim)',
    'Tag top 3 friction points observed',
    'Rate tester\'s coach trust level (1-5)',
    'Rate tester\'s adoption likelihood (1-5)',
    'Send gift card',
    'Share top insights with team',
]
for item in items:
    doc.add_paragraph(f'☐ {item}')

add_heading('Appendix C: Analysis Framework', 2)
doc.add_paragraph('After all testing is complete, organise findings into these categories:')
add_table(
    ['Category', 'Data Sources', 'Output'],
    [
        ['Navigation friction', 'Maze heatmaps, misclick rates, moderated observation', 'Prioritised list of UI fixes'],
        ['Comprehension gaps', 'Maze success rates, Ballpark open text, 5-Second Tests', 'Labelling/copy changes needed'],
        ['Coach perception', 'Ballpark ratings, moderated interviews', 'Trust improvement priorities'],
        ['Output value', 'Maze opinion scales, moderated debrief', 'Output enhancement backlog'],
        ['Journey drop-off', 'Maze bounce rates, moderated observation', 'Phase transition improvements'],
        ['Feature requests', 'All open text responses', 'Backlog items for future sprints'],
    ]
)

doc.add_paragraph()

quote('Severity Rating Scale')
add_table(
    ['Severity', 'Definition', 'Example'],
    [
        ['Critical', 'Blocks task completion; no workaround', 'Cannot find upload function'],
        ['High', 'Significant confusion; workaround difficult', 'Coach response feels completely irrelevant'],
        ['Medium', 'Causes frustration; workaround exists', 'Progress indicator unclear'],
        ['Low', 'Minor annoyance; easily overlooked', 'Awkward label text'],
    ]
)

# Document control
doc.add_page_break()
add_heading('Document Control', 1)
add_table(
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ['1.0', 'February 2026', 'Frontera Product Team', 'Lean testing plan ($1k budget)'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('— End of Document —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_path = r'c:\Users\deeks\frontera-platform\Background\Product_Strategy_Coach_Lean_Testing_Plan.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
