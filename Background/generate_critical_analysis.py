"""
Generate a critical analysis of the Frontera Product Strategy Coach
vs. the pain points cited by product leaders.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

NAVY = RGBColor(0x1a, 0x1f, 0x3a)


def set_cell_shading(cell, color_hex):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(s)


def H(text, level=1):
    h = doc.add_heading(text, level)
    for r in h.runs:
        r.font.color.rgb = NAVY
    return h


def T(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hc = t.rows[0].cells
    for i, header in enumerate(headers):
        hc[i].text = header
        hc[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hc[i], '1a1f3a')
        hc[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for rd in rows:
        r = t.add_row()
        for i, ct in enumerate(rd):
            r.cells[i].text = str(ct)
    doc.add_paragraph()
    return t


def B(items):
    for i in items:
        doc.add_paragraph(i, style='List Bullet')


def Q(text):
    doc.add_paragraph(text, style='Intense Quote')


def P(label, text):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)
    return p


def para(text):
    doc.add_paragraph(text)


# ============================================================================
# TITLE
# ============================================================================
title = doc.add_heading('Product Strategy Coach', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    r.font.size = Pt(32)
    r.font.color.rgb = NAVY

sub = doc.add_paragraph('Critical Analysis vs. Product Leader Pain Points')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(20)
sub.runs[0].font.color.rgb = RGBColor(0x08, 0x91, 0xb2)

doc.add_paragraph()
tag = doc.add_paragraph('Journey, Coaching & Artefact Assessment with Suggested Enhancements')
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.runs[0].font.size = Pt(14)
tag.runs[0].italic = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Date: February 2026\n')
info.add_run('Sources: Frontera Strategy Challenges Research, Lenny\'s Newsletter,\n')
info.add_run('Marty Cagan/SVPG, Atlassian State of Product 2026, Mind the Product 2025\n')
info.add_run('Assessment: Independent product design & coaching review')

doc.add_page_break()

# ============================================================================
# TOC
# ============================================================================
H('Table of Contents', 1)
toc = [
    ('1.', 'Executive Assessment'),
    ('2.', 'Challenge-by-Challenge Analysis'),
    ('3.', 'Journey Assessment'),
    ('4.', 'Coaching Assessment'),
    ('5.', 'Artefact Assessment'),
    ('6.', 'Enhancement Roadmap'),
]
for n, i in toc:
    p = doc.add_paragraph()
    p.add_run(f'{n} ').bold = True
    p.add_run(i)

doc.add_page_break()

# ============================================================================
# 1. EXECUTIVE ASSESSMENT
# ============================================================================
H('1. Executive Assessment', 1)

H('Overall Verdict', 2)
para(
    'The Frontera Product Strategy Coach addresses 5 of the 8 core challenges cited by product leaders '
    'with genuine depth and rigour. However, it has material gaps in 3 areas that represent the difference '
    'between a useful tool and an indispensable strategic platform. The strongest aspects \u2014 structured research, '
    'evidence-linked synthesis, and disciplined strategic bet framing \u2014 are genuinely differentiated. '
    'The weakest aspects \u2014 strategy-to-execution bridging, continuous adaptation, and stakeholder alignment '
    '\u2014 are precisely the areas where product leaders report the most pain.'
)

H('Scorecard', 2)
T(
    ['Challenge', 'Research Priority', 'PSC Coverage', 'Gap Severity'],
    [
        ['1. No time for strategic thinking', 'Critical (49%)', 'Strong', 'Low'],
        ['2. Strategy disconnected from execution', 'Critical (84%)', 'Weak', 'High'],
        ['3. Inability to prioritise', 'High (40%)', 'Strong', 'Low-Medium'],
        ['4. Lack of research rigour', 'High (49%)', 'Very Strong', 'Low'],
        ['5. Cross-functional misalignment', 'High (80%)', 'Moderate', 'Medium-High'],
        ['6. Strategy is static, not adaptive', 'High (39%)', 'Weak', 'High'],
        ['7. Competency gaps in strategy', 'Medium (59%)', 'Strong', 'Low-Medium'],
        ['8. Measuring strategic impact', 'Medium (12%)', 'Moderate', 'Medium'],
    ]
)

H('The Fundamental Tension', 2)
para(
    'The research reveals a paradox: 85% of product teams feel empowered to lead strategy, yet 49% '
    'cannot find time for it and 84% worry their products will fail. The PSC addresses the '
    '"how to do strategy" problem effectively but underserves the "how to make strategy stick" problem. '
    'The tool produces excellent strategic artefacts but does not yet ensure those artefacts '
    'change organisational behaviour. This is the critical gap between a strategy formulation tool '
    'and a strategy execution platform.'
)

doc.add_page_break()

# ============================================================================
# 2. CHALLENGE-BY-CHALLENGE ANALYSIS
# ============================================================================
H('2. Challenge-by-Challenge Analysis', 1)

# Challenge 1
H('Challenge 1: No Time for Strategic Thinking', 2)
Q('49% of product teams lack time for strategic planning (Atlassian 2026); 21% cite strategic-vs-tactical balance as primary challenge (Mind the Product 2025)')

H('What the PSC Does Well', 3)
B([
    'Breaks strategic thinking into modular phases (Discovery \u2192 Research \u2192 Synthesis \u2192 Bets) that can be progressed in focused sessions',
    'AI-powered synthesis compresses what would be weeks of manual analysis into minutes',
    'Structured research questions eliminate the blank-page problem \u2014 leaders don\'t have to figure out what to research',
    'Save-and-resume architecture respects interrupted schedules',
    'AI Research Assistant can generate synthetic documents when leaders lack time to gather materials',
])

H('Where Gaps Remain', 3)
B([
    'Full completion still requires ~2.5 hours of focused writing (36 open-ended questions across 9 areas)',
    'No "quick start" mode for time-constrained leaders who need a strategy direction in 30 minutes',
    'No session timer or time-boxing to help leaders stay focused within defined windows',
    'The Research phase is the bottleneck \u2014 it demands the most time but provides the least immediate gratification',
])

H('Suggested Enhancements', 3)
T(
    ['Enhancement', 'Impact', 'Effort'],
    [
        ['Reduce to 3 questions per area (36\u219227)', 'Saves ~45 min; reduces abandonment', 'Low'],
        ['Add "Express Mode" \u2014 coach-led interview that populates research areas conversationally', 'Dramatic time reduction; more natural input', 'High'],
        ['Show estimated time remaining per phase', 'Sets expectations; reduces frustration', 'Low'],
        ['Add AI pre-fill suggestions based on uploaded documents', 'Reduces manual writing; accelerates research', 'Medium'],
    ]
)

doc.add_page_break()

# Challenge 2
H('Challenge 2: Strategy Disconnected from Execution', 2)
Q('84% of teams worry products won\'t succeed (Atlassian 2026); strategy increasingly top-down with 5% YoY increase (ProductPlan 2025)')

H('What the PSC Does Well', 3)
B([
    'Strategic Bets include explicit success metrics and kill criteria \u2014 these are execution-ready',
    'Evidence trail from research \u2192 synthesis \u2192 bets creates traceable logic from insight to action',
    'Playing to Win framework inherently bridges strategy (Where to Play) with execution (Capabilities Required)',
    'PDF exports create shareable artefacts that can reach teams',
])

H('Where Gaps Remain \u2014 CRITICAL', 3)
B([
    'No mechanism to cascade strategy to team-level objectives (no OKR generation, no team briefs)',
    'Strategy lives as a static PDF export, not a living system teams can reference daily',
    'No "Strategic Guardrails" output that tells teams what NOT to do \u2014 a key Cagan principle',
    'No decision-making framework output that teams can use for daily prioritisation',
    'No connection between strategic bets and product backlog/roadmap tools',
    'The tool ends at "here are your bets" \u2014 it does not help with "how do I brief my teams?"',
    'Marty Cagan emphasises the strategy-to-team communication gap is where most strategies die; the PSC does not address this',
])

H('Suggested Enhancements', 3)
T(
    ['Enhancement', 'Impact', 'Effort'],
    [
        ['Add "Team Brief Generator" \u2014 translates bets into team-readable context (problems to solve, guardrails, success metrics)', 'Directly bridges strategy\u2192execution gap', 'Medium'],
        ['Generate Strategic Guardrails from synthesis tensions ("We will NOT compete on price")', 'Gives teams daily decision support', 'Medium'],
        ['Add Decision Framework output ("When choosing between X and Y, prioritise...")', 'Embeds strategy into team decisions', 'Medium'],
        ['Integrate with Jira/Linear/Notion to push strategic context into workflow tools', 'Makes strategy visible where teams work', 'High'],
        ['Create shareable "Strategy on a Page" view (not PDF, but live link)', 'Keeps strategy living and accessible', 'Medium'],
    ]
)

doc.add_page_break()

# Challenge 3
H('Challenge 3: Inability to Prioritise Effectively', 2)
Q('40% cite feature prioritisation as biggest challenge (Atlassian 2026); companies allocate capacity by business unit rather than strategic importance (Cagan)')

H('What the PSC Does Well', 3)
B([
    '2\u00d72 Strategic Opportunity Map (Market Attractiveness \u00d7 Capability Fit) creates visual prioritisation',
    'INVEST/EXPLORE/HARVEST/DIVEST quadrant assignment gives clear directional guidance',
    'Multi-dimensional scoring (Market 40% + Capability 35% + Competitive 25%) resists single-axis bias',
    'Strategic Bets use 4-dimension Janakiraman scoring (Impact, Certainty, Clarity, Uniqueness)',
    'Evidence linking means priorities are defensible, not political',
    'Kill criteria force pre-commitment to resource reallocation decisions',
])

H('Where Gaps Remain', 3)
B([
    'No resource allocation guidance \u2014 the tool prioritises opportunities but doesn\'t help allocate team capacity',
    'No "portfolio view" comparing effort vs. impact across all bets simultaneously',
    'No sequencing recommendation beyond coaching prompts ("Which bets are prerequisite?")',
    'Scoring relies on AI judgment from research inputs \u2014 no calibration mechanism where leadership validates scores',
    'The tool doesn\'t address the structural problem Cagan identifies: companies fund projects, not problems',
])

H('Suggested Enhancements', 3)
T(
    ['Enhancement', 'Impact', 'Effort'],
    [
        ['Add interactive "Effort vs. Impact" scatter plot for bets', 'Visual portfolio prioritisation', 'Medium'],
        ['Add bet sequencing timeline view', 'Shows dependencies and critical path', 'Medium'],
        ['Add "Leadership Calibration" step where users adjust AI-generated scores with rationale', 'Increases ownership and accuracy', 'Low-Medium'],
        ['Generate resource allocation recommendation based on bet priorities', 'Directly addresses allocation challenge', 'High'],
    ]
)

doc.add_page_break()

# Challenge 4
H('Challenge 4: Lack of Strategic Research Rigour', 2)
Q('49% can\'t fit in data analysis (Atlassian 2026); shift from data-driven to intuition-based decisions; 40% do little or no experimentation')

H('What the PSC Does Well \u2014 STRONGEST AREA', 3)
B([
    'Structured 3-territory research (Company, Customer, Competitor) ensures no dimension is skipped',
    '9 research areas with guided questions eliminate "where do I start?" paralysis',
    'AI Research Assistant generates synthetic documents from web sources when primary research is unavailable',
    'Evidence grading through synthesis \u2014 each opportunity cites specific research with quotes',
    'Forced triangulation: synthesis requires evidence from multiple territories for each opportunity',
    'Minimum 4 of 9 areas required before synthesis \u2014 prevents premature conclusions',
    'Coach challenges surface-level responses and probes for depth',
    'Industry-specific guidance adapts research prompts to sector context (fintech, healthcare, etc.)',
])

H('Where Gaps Remain', 3)
B([
    'All 36 questions are open-ended text \u2014 no data input (market size figures, NPS scores, revenue)',
    'No way to attach quantitative data to research responses (spreadsheets, metrics dashboards)',
    'Research quality depends entirely on user knowledge \u2014 no validation against external data',
    'No "confidence rating" for individual research responses (is this a fact, an estimate, or a guess?)',
    'No methodology guidance within questions (e.g., how to actually conduct customer research)',
])

H('Suggested Enhancements', 3)
T(
    ['Enhancement', 'Impact', 'Effort'],
    [
        ['Add confidence rating per response (High/Medium/Low/Guess)', 'Surfaces evidence quality; feeds synthesis confidence', 'Low'],
        ['Allow numeric data fields alongside text (market size, growth rate, NPS)', 'Grounds strategy in data, not just narrative', 'Medium'],
        ['Add "Research Method" hints per question ("Try: 5 customer interviews, win/loss analysis")', 'Builds research capability while guiding', 'Low'],
        ['Enable file attachment per research area (attach spreadsheets, reports)', 'Incorporates quantitative evidence directly', 'Medium'],
    ]
)

doc.add_page_break()

# Challenge 5
H('Challenge 5: Cross-Functional Misalignment', 2)
Q('80% don\'t involve engineers early (Atlassian 2026); 35% cite stakeholder influence as primary obstacle (Mind the Product 2025)')

H('What the PSC Does Well', 3)
B([
    'Company Territory research captures internal capabilities, resources, and constraints',
    'Synthesis identifies Strategic Tensions \u2014 conflicts in the research that surface organisational disagreements',
    'Evidence linking creates a shared, traceable logic that stakeholders can review',
    'PDF exports are stakeholder-ready (professional formatting, executive summary)',
    'Coach acknowledges organisational politics ("Transform Recovery Awareness" for orgs with failed past initiatives)',
])

H('Where Gaps Remain', 3)
B([
    'No explicit cross-functional input mechanism \u2014 the tool is single-user, not collaborative',
    'No "Colleague Research" module as described in the research document\'s recommended solution \u2014 gathering input from Sales, Engineering, Support, Marketing',
    'No Alignment Gap Detector that surfaces contradictions between what leadership says and what frontline teams experience',
    'No Risk Radar visualisation showing organisational risks',
    'No way for multiple stakeholders to contribute to or validate territory research',
    'Strategy is developed in isolation by one leader, then presented \u2014 the opposite of collaborative strategy development',
    'Cagan emphasises strategy should be developed WITH senior leaders, not presented TO them',
])

H('Suggested Enhancements', 3)
T(
    ['Enhancement', 'Impact', 'Effort'],
    [
        ['Add "Stakeholder Input" mode \u2014 share specific research areas with colleagues for input', 'Makes strategy collaborative, not isolated', 'High'],
        ['Add Alignment Gap Detector \u2014 flag contradictions between territory responses (e.g., "Leadership says customer-first but no customer research budget")', 'Proactively surfaces misalignment', 'Medium'],
        ['Generate "Stakeholder Communication Pack" from synthesis (tailored views for CEO, CTO, Sales VP)', 'Bridges alignment gap with targeted communication', 'Medium'],
        ['Add team commenting/annotation on synthesis outputs', 'Creates shared ownership of strategy', 'Medium'],
    ]
)

doc.add_page_break()

# Challenge 6
H('Challenge 6: Strategy Is Static, Not Adaptive', 2)
Q('Cagan: good strategies are "constantly updated based on learnings"; 39% of product work doesn\'t lead to expected outcomes (Productboard)')

H('What the PSC Does Well', 3)
B([
    'Conversation-based interaction means users can re-engage and evolve thinking over time',
    'Strategic Bets include kill criteria and kill dates \u2014 building in review triggers',
    'Assumption testing (WWHBT) creates checkpoints for validating or invalidating strategy',
    'Framework state persists across sessions, allowing iterative refinement',
])

H('Where Gaps Remain \u2014 CRITICAL', 3)
B([
    'Strategy is treated as a linear journey with a terminal endpoint (Discovery \u2192 Research \u2192 Synthesis \u2192 Bets \u2192 Done)',
    'No mechanism to update synthesis when new information emerges ("We lost a major deal" / "Competitor launched X")',
    'No Signals Radar or market monitoring that prompts strategy review',
    'No version history \u2014 cannot compare current strategy to 3-months-ago strategy',
    'No "assumption tracker" that monitors which WWHBT assumptions have been validated or invalidated',
    'Kill dates are set but nothing triggers a review when they arrive',
    'The PDF export creates a point-in-time document that immediately begins aging',
    'No learning loop: outcomes of bets don\'t feed back into revised territory research',
    'This is the largest philosophical gap \u2014 the tool builds strategy as a document, not as a living system',
])

H('Suggested Enhancements', 3)
T(
    ['Enhancement', 'Impact', 'Effort'],
    [
        ['Add "Strategy Review" mode \u2014 periodic prompted re-assessment of key assumptions', 'Makes strategy a living practice, not a one-time event', 'Medium'],
        ['Add kill-date notifications that trigger review prompts', 'Ensures bets are evaluated on schedule', 'Low'],
        ['Add assumption tracker dashboard (validated/invalidated/untested)', 'Visual progress toward strategic confidence', 'Medium'],
        ['Add "Signal Log" \u2014 users record market events that may affect strategy', 'Creates a change-trigger mechanism', 'Medium'],
        ['Add strategy versioning with diff view', 'Shows strategic evolution over time', 'High'],
        ['Create a feedback loop: bet outcomes update territory research confidence', 'Closes the strategy\u2192learning loop', 'High'],
    ]
)

doc.add_page_break()

# Challenge 7
H('Challenge 7: Competency Gaps in Strategic Thinking', 2)
Q('59% believe strategy and business acumen are most important PM skills (Productboard); only 2% prioritise talent development')

H('What the PSC Does Well', 3)
B([
    'Coaching is embedded in the flow of work \u2014 leaders learn strategy BY DOING strategy, not from workshops',
    'Coach challenges assumptions constructively, building "strategic muscle" through practice',
    'Frameworks are invisible \u2014 users experience Playing to Win without learning theory',
    'Phase-specific coaching adapts sophistication to where the user is in the journey',
    'The 14 distinct coaching behaviours in the Bets phase (altitude check, moat challenge, portfolio balance) are genuinely advanced',
    'Expert perspectives and case studies are injected contextually, not presented as curriculum',
    'The coach explicitly avoids patronising language and empty affirmations \u2014 treating leaders as capable adults',
])

H('Where Gaps Remain', 3)
B([
    'No explicit competency tracking \u2014 the coach doesn\'t measure or report on strategic skill development',
    'No progressive difficulty \u2014 a first-time user and a returning expert get the same coaching depth',
    'No "learning moments" that explain WHY a framework is being used ("I\'m asking about kill criteria because of pre-commitment theory...")',
    'No post-session reflection prompts ("What did you learn? What surprised you? What assumption changed?")',
    'The coaching is reactive (responds to user input) rather than proactive (identifies capability gaps and targets them)',
])

H('Suggested Enhancements', 3)
T(
    ['Enhancement', 'Impact', 'Effort'],
    [
        ['Add strategic maturity assessment at onboarding (adapts coaching depth)', 'Personalises coaching to capability level', 'Medium'],
        ['Add optional "Why This Matters" expandable sections explaining frameworks in context', 'Builds understanding alongside practice', 'Low'],
        ['Add post-phase reflection prompts', 'Deepens learning from each session', 'Low'],
        ['Track user engagement patterns to identify coaching opportunities ("You skipped competitive analysis \u2014 this is a common blind spot")', 'Proactive capability development', 'Medium'],
    ]
)

doc.add_page_break()

# Challenge 8
H('Challenge 8: Measuring Strategic Impact', 2)
Q('Only 12% find driving measurable business results rewarding (Atlassian 2026); teams can\'t connect work to outcomes')

H('What the PSC Does Well', 3)
B([
    'Strategic Bets require explicit success metrics with numbers and timeframes',
    'Kill criteria create measurable failure thresholds',
    'The coach demands quantitative metrics ("Add a number and timeframe")',
    'Scoring dimensions (Market Attractiveness, Capability Fit, Competitive Advantage) create a quantified strategic assessment',
    'Playing to Win\'s "Management Systems" field asks what metrics will track progress',
])

H('Where Gaps Remain', 3)
B([
    'No North Star metric generation \u2014 the tool doesn\'t help define the overarching success measure',
    'No OKR cascade from strategic bets to team-level objectives',
    'No outcome tracking \u2014 bet success metrics are defined but never measured',
    'No dashboard connecting strategic bets to actual business performance',
    'No ROI framing \u2014 the tool doesn\'t help leaders communicate strategic impact in financial terms',
    'The "Management Systems" field in PTW is the least developed output \u2014 often generic',
])

H('Suggested Enhancements', 3)
T(
    ['Enhancement', 'Impact', 'Effort'],
    [
        ['Add North Star metric workshop (guided coaching to define the primary outcome metric)', 'Anchors all strategy to one measurable outcome', 'Medium'],
        ['Generate OKR proposals from strategic bets', 'Cascades strategy into measurable team objectives', 'Medium'],
        ['Add bet outcome tracking (actual vs. predicted)', 'Closes the measurement loop', 'Medium-High'],
        ['Add "Strategic Impact Summary" for board reporting', 'Helps leaders communicate strategy ROI', 'Medium'],
    ]
)

doc.add_page_break()

# ============================================================================
# 3. JOURNEY ASSESSMENT
# ============================================================================
H('3. Journey Assessment', 1)

H('Phase Flow Analysis', 2)
T(
    ['Phase', 'Strength', 'Weakness', 'Verdict'],
    [
        ['Discovery', 'Low barrier to entry (1 doc minimum); AI Research Assistant for quick start', 'Passive \u2014 mostly file upload, not coaching-led. Coach doesn\'t actively explore the "why" behind transformation', 'Good entry, needs more active coaching'],
        ['Research', 'Structured, comprehensive; guided questions eliminate blank-page problem', 'Heavy \u2014 36 open-ended questions; uniform format (all textarea); no variety in input type', 'Core methodology is strong; execution needs streamlining'],
        ['Synthesis', 'Genuinely impressive \u2014 evidence-linked, framework-grounded, visually clear', 'One-shot generation; no iterative refinement; no leadership calibration of scores', 'Best phase; needs iteration capability'],
        ['Strategic Bets', 'Sophisticated coaching (14 behaviours); kill criteria discipline is rare and valuable', 'Terminal endpoint \u2014 strategy ends here rather than flowing to execution', 'Strong but incomplete \u2014 needs Phase 5'],
    ]
)

H('The Missing Phase 5', 2)
para(
    'The most significant journey gap is the absence of a post-Bets execution bridge. '
    'The research is clear: 84% of teams worry their products will fail, and the primary cause is '
    'strategy that never reaches the people building the product. The PSC produces excellent bets '
    'but stops at the moment when the hardest work begins: translating bets into team context, '
    'communicating strategy across the organisation, and monitoring execution against strategic intent.'
)

P('Recommended Phase 5: "Strategic Activation" \u2014 ', 'A phase that generates:')
B([
    'Team briefs per strategic bet (problem to solve, guardrails, success metrics)',
    'Stakeholder communication pack (tailored messaging for different audiences)',
    'Strategic guardrails document ("We will / We will not")',
    'Decision framework for daily prioritisation',
    'OKR proposals cascaded from bets',
    'Assumption tracking dashboard',
])

H('Journey Pacing', 2)
T(
    ['Phase', 'Estimated Time', 'User Energy', 'Gratification'],
    [
        ['Discovery', '15-30 min', 'High (new, curious)', 'Low (uploading files)'],
        ['Research', '1.5-2.5 hours', 'Declining (repetitive writing)', 'Low (no visible output yet)'],
        ['Synthesis', '5-10 min (generation)', 'Renewed (wow moment)', 'Very High (see strategy emerge)'],
        ['Bets', '30-60 min', 'Moderate (editing/refining)', 'High (actionable outputs)'],
    ]
)

para(
    'The journey has an energy problem: users invest 2+ hours in Research before receiving any '
    'strategic output. The "wow moment" (synthesis generation) comes too late. Leaders under '
    'time pressure (49% per Atlassian) are most likely to abandon during the Research phase \u2014 '
    'precisely where the tool demands the most effort for the least visible return.'
)

P('Recommendation: ', 'Introduce "micro-synthesis" moments during Research. After each territory is '
  'completed, generate a quick territorial insight summary. This gives users progressive gratification '
  'and builds momentum toward the full synthesis.')

doc.add_page_break()

# ============================================================================
# 4. COACHING ASSESSMENT
# ============================================================================
H('4. Coaching Assessment', 1)

H('What Makes the Coaching Genuinely Good', 2)

P('1. Authentic Voice \u2014 ', 'The coach avoids the "Great question!" / "Absolutely!" trap that plagues '
  'most AI products. The system prompt explicitly forbids patronising language and filler phrases. '
  'The tone is confident without being arrogant, challenging without being dismissive. '
  'This is rare and valuable \u2014 most enterprise AI tools default to sycophantic affirmation.')

P('2. Framework Invisibility \u2014 ', 'Users experience Playing to Win, Jobs-to-Be-Done, and the '
  'Janakiraman scoring model without being taught them. The frameworks are embedded in the coaching '
  'language rather than presented as theory. This is the right approach \u2014 as Cagan says, '
  '"the most important role of product leadership is coaching," not lecturing.')

P('3. Altitude Awareness \u2014 ', 'The coach explicitly separates market-level strategy from team-level '
  'product discovery. When users describe features, the coach elevates: "This sounds like a product '
  'discovery experiment. At the strategic level, what MARKET-LEVEL question are you trying to answer?" '
  'This prevents the common failure mode of strategy degenerating into a feature list.')

P('4. Kill Criteria Discipline \u2014 ', 'Forcing pre-commitment to abandonment conditions is rare in '
  'strategy tools. This directly addresses the sunk-cost fallacy that traps organisations into '
  'persisting with failing strategies. The coach demands specific kill dates, thresholds, and signals.')

P('5. Portfolio-Level Thinking \u2014 ', 'The coach examines bets collectively, not individually. '
  'It asks about balance (offensive vs. defensive), coherence (do bets test the same thesis?), '
  'sequencing (which bets are prerequisite?), and optionality (which preserve strategic flexibility?). '
  'This is genuinely sophisticated strategic coaching.')

H('What the Coaching Gets Wrong or Misses', 2)

P('1. Discovery Phase Is Under-Coached \u2014 ', 'The highest-leverage coaching moment is the beginning: '
  'understanding WHY the leader is pursuing transformation, what political dynamics exist, and what '
  'constraints are non-negotiable. The current Discovery phase is primarily a data-upload exercise. '
  'The coach waits for the user to engage rather than proactively exploring the strategic context. '
  'Lenny Rachitsky emphasises that the best strategies start with truly understanding the core problem \u2014 '
  'the PSC should spend more time here.')

P('2. Research Phase Coaching Is Generic \u2014 ', 'While the coach adapts to industry and phase, it does '
  'not provide methodological guidance for answering research questions. A user facing "What are your '
  'primary customer segments?" may not know HOW to segment customers. The coach could offer: '
  '"Try segmenting by: job-to-be-done, willingness to pay, acquisition channel, or use frequency. '
  'Which dimension seems most strategically relevant for your business?"')

P('3. No Socratic Discovery of Blind Spots \u2014 ', 'The coach responds to what users provide but does '
  'not identify what they HAVEN\'T provided. If a user completes Customer Territory without mentioning '
  'churn or retention, the coach should flag this: "I notice you haven\'t mentioned customer retention. '
  'Is churn a non-issue, or a blind spot worth exploring?"')

P('4. Coaching Doesn\'t Build on Itself \u2014 ', 'Each session starts fresh in terms of coaching insight. '
  'The coach doesn\'t remember that last session the user struggled with competitive analysis, or that '
  'they tend to think in features rather than markets. Progressive coaching should adapt based on '
  'accumulated observation of the user\'s strategic thinking patterns.')

P('5. No Challenge Escalation \u2014 ', 'The coach challenges with equal intensity throughout. '
  'Early in the journey, gentle probing is appropriate. By the Bets phase, the coach should be '
  'significantly more demanding: "This bet has no kill criteria, vague metrics, and no evidence link. '
  'I cannot recommend proceeding without addressing these gaps." The intensity should match the stakes.')

P('6. Missing the "So What?" Moment \u2014 ', 'After synthesis is generated, the coach helps explore '
  'outputs but doesn\'t force the crucial "So what does this mean for your company specifically?" '
  'conversation. The synthesis shows opportunities, but the coach should push for commitment: '
  '"Of these 5 opportunities, which one scares you the most? That\'s often where the real strategy is."')

doc.add_page_break()

# ============================================================================
# 5. ARTEFACT ASSESSMENT
# ============================================================================
H('5. Artefact Assessment', 1)

H('Artefact Inventory', 2)
T(
    ['Artefact', 'Format', 'Quality', 'Execution Readiness'],
    [
        ['Client Context Profile', 'Editable form (in-app)', 'Good \u2014 captures key dimensions', 'Low \u2014 context only, not actionable'],
        ['Territory Research Responses', 'JSONB stored Q&A (in-app)', 'Variable \u2014 depends on user depth', 'None \u2014 raw research input'],
        ['Strategic Synthesis', 'Interactive view + PDF export', 'Strong \u2014 evidence-linked, scored, visualised', 'Medium \u2014 opportunities are directional, not executable'],
        ['Strategic Opportunity Map', 'Interactive 2\u00d72 matrix (in-app)', 'Strong \u2014 clear visual prioritisation', 'Medium \u2014 shows what to focus on, not how'],
        ['Strategic Bets', 'Interactive cards + PDF export', 'Very Strong \u2014 hypothesis-driven with kill criteria', 'High \u2014 closest to execution-ready'],
        ['Strategy Document', 'Generated narrative + PDF', 'Good \u2014 coherent prose from bets', 'Medium-High \u2014 shareable but needs customisation'],
    ]
)

H('What\'s Missing from the Artefact Suite', 2)

P('1. Strategy on a Page \u2014 ', 'No single-view summary that captures the entire strategy in a format '
  'that can be pinned to a wall, shared in Slack, or referenced in a meeting. The best strategies '
  'are communicable in 60 seconds. The PSC produces detailed artefacts but not a concise summary.')

P('2. Strategic Guardrails Document \u2014 ', '"We will / We will not" statements are among the most '
  'useful strategic artefacts for teams. They prevent scope creep and provide daily decision support. '
  'The synthesis identifies tensions but doesn\'t convert them into clear boundary statements.')

P('3. Assumption Register \u2014 ', 'WWHBT assumptions are embedded in opportunity cards but not '
  'extracted into a trackable, updatable register. An assumption register is the backbone of '
  'adaptive strategy \u2014 it tells you what to test and what would invalidate your strategy.')

P('4. Stakeholder Communication Pack \u2014 ', 'Different audiences need different views of the same '
  'strategy. The CEO needs a 2-minute narrative. The CTO needs capability requirements. Sales needs '
  'positioning guidance. The current tool produces one-size-fits-all output.')

P('5. Competitive Response Scenarios \u2014 ', 'The synthesis identifies competitive dynamics but '
  'doesn\'t model "If we execute Bet A, how will Competitor X likely respond?" This wargaming '
  'perspective is standard in consulting engagements but absent here.')

P('6. Research Evidence Summary \u2014 ', 'No artefact summarises the raw research into a digestible '
  'briefing document. Before synthesis, leaders may want to share "here\'s what we learned" with '
  'their team. Currently, research lives only as Q&A responses in the tool.')

H('Artefact Quality vs. Consulting Benchmark', 2)
T(
    ['Dimension', 'PSC Current', 'Strategy Consulting', 'Gap'],
    [
        ['Research depth', 'Structured but self-reported', 'Primary + secondary research, interviews, data analysis', 'PSC lacks primary research methods'],
        ['Framework rigour', 'Strong (PTW, JTBD, DHM)', 'Strong (varies by firm)', 'Comparable'],
        ['Evidence quality', 'Good (user-sourced, AI-synthesised)', 'Strong (validated, multi-source)', 'PSC evidence quality depends on user input'],
        ['Visual communication', 'Good (2\u00d72 map, scored cards)', 'Excellent (custom charts, diagrams)', 'PSC needs richer visualisation'],
        ['Actionability', 'High (kill criteria, success metrics)', 'Variable (often abstract recommendations)', 'PSC advantage \u2014 more execution-ready'],
        ['Stakeholder readiness', 'Moderate (PDF export)', 'Excellent (tailored decks per audience)', 'PSC needs audience-specific outputs'],
        ['Cost', '~$X/month', '$200K-500K per engagement', 'PSC massive advantage'],
        ['Time to value', 'Days', 'Weeks to months', 'PSC massive advantage'],
    ]
)

doc.add_page_break()

# ============================================================================
# 6. ENHANCEMENT ROADMAP
# ============================================================================
H('6. Enhancement Roadmap', 1)

H('Priority 1: Quick Wins (Low Effort, High Impact)', 2)
T(
    ['#', 'Enhancement', 'Challenge Addressed', 'Effort'],
    [
        ['1', 'Reduce questions to 3 per area (36\u219227)', 'Time (#1), Abandonment', '1-2 days'],
        ['2', 'Add confidence rating per research response', 'Research rigour (#4)', '1 day'],
        ['3', 'Add post-phase reflection prompts', 'Competency (#7)', '1 day'],
        ['4', 'Add estimated time remaining per phase', 'Time (#1)', '0.5 days'],
        ['5', 'Add optional "Why This Matters" expandable sections in coaching', 'Competency (#7)', '2 days'],
        ['6', 'Add kill-date notification triggers', 'Adaptive (#6)', '1 day'],
    ]
)

H('Priority 2: Strategic Enhancements (Medium Effort, High Impact)', 2)
T(
    ['#', 'Enhancement', 'Challenge Addressed', 'Effort'],
    [
        ['7', 'Micro-synthesis after each territory completion', 'Time (#1), Journey pacing', '3-5 days'],
        ['8', 'Strategy on a Page generator', 'Execution gap (#2)', '3 days'],
        ['9', 'Strategic Guardrails generator from synthesis tensions', 'Execution gap (#2), Prioritisation (#3)', '3 days'],
        ['10', 'Assumption tracker dashboard', 'Adaptive (#6), Measurement (#8)', '5 days'],
        ['11', 'Leadership score calibration step in synthesis', 'Prioritisation (#3)', '3 days'],
        ['12', 'Stakeholder communication pack generator (CEO/CTO/Sales views)', 'Alignment (#5), Execution gap (#2)', '5 days'],
        ['13', 'Blind spot detection in coaching ("I notice you haven\'t mentioned...")', 'Coaching quality, Research rigour (#4)', '3-5 days'],
        ['14', 'Research method hints per question', 'Competency (#7), Research rigour (#4)', '2 days'],
    ]
)

H('Priority 3: Transformative Features (High Effort, Highest Impact)', 2)
T(
    ['#', 'Enhancement', 'Challenge Addressed', 'Effort'],
    [
        ['15', 'Phase 5: Strategic Activation (team briefs, OKRs, decision frameworks)', 'Execution gap (#2) \u2014 CRITICAL', '2-3 weeks'],
        ['16', 'Express Mode \u2014 coach-led conversational research', 'Time (#1)', '2 weeks'],
        ['17', 'Strategy Review mode with versioning', 'Adaptive (#6) \u2014 CRITICAL', '2 weeks'],
        ['18', 'Stakeholder input/collaboration features', 'Alignment (#5)', '3-4 weeks'],
        ['19', 'Bet outcome tracking and strategy\u2192learning feedback loop', 'Adaptive (#6), Measurement (#8)', '2-3 weeks'],
        ['20', 'Workflow tool integration (Jira/Linear/Notion)', 'Execution gap (#2)', '3-4 weeks'],
    ]
)

H('Recommended Implementation Sequence', 2)
para(
    'Sequence should follow a "strengthen core, then extend" pattern. '
    'The first priority is reducing friction in the existing journey (Quick Wins), '
    'then enriching the artefact suite (Strategic Enhancements), and finally '
    'extending the platform into execution territory (Transformative Features).'
)

T(
    ['Sprint', 'Focus', 'Items'],
    [
        ['Sprint 1 (1 week)', 'Reduce friction', '#1-6 (Quick Wins)'],
        ['Sprint 2-3 (2 weeks)', 'Enrich journey', '#7-9 (Micro-synthesis, Strategy on a Page, Guardrails)'],
        ['Sprint 4-5 (2 weeks)', 'Deepen coaching', '#10, 11, 13, 14 (Assumption tracker, Calibration, Blind spots)'],
        ['Sprint 6-8 (3 weeks)', 'Bridge to execution', '#15 (Phase 5: Strategic Activation)'],
        ['Sprint 9-10 (2 weeks)', 'Enable adaptation', '#17 (Strategy Review mode)'],
        ['Sprint 11+ (ongoing)', 'Scale collaboration', '#18, 19, 20 (Stakeholder input, Outcome tracking, Integrations)'],
    ]
)

# Document control
doc.add_page_break()
H('Document Control', 1)
T(
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ['1.0', 'February 2026', 'Frontera Product Team', 'Initial critical analysis'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('\u2014 End of Document \u2014').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_path = r'c:\Users\deeks\frontera-platform\Background\Product_Strategy_Coach_Critical_Analysis.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
