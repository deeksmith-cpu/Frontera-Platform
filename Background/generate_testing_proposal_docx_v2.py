"""
Generate a Word document for the Product Strategy Coach User Testing Proposal
Version 2: Updated with Maze terminology and best practices
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Define colors
NAVY = RGBColor(0x1a, 0x1f, 0x3a)
GOLD = RGBColor(0xfb, 0xbf, 0x24)
CYAN = RGBColor(0x08, 0x91, 0xb2)

# Helper functions
def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(text, level=1):
    """Add a heading with consistent styling"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = NAVY
    return heading

def add_table_from_data(headers, rows, widths=None):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '1a1f3a')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)

    doc.add_paragraph()
    return table

# ============================================================================
# TITLE PAGE
# ============================================================================
title = doc.add_heading('Product Strategy Coach', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(36)
    run.font.color.rgb = NAVY

subtitle = doc.add_paragraph('User Research & Testing Proposal')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(24)
subtitle.runs[0].font.color.rgb = CYAN

doc.add_paragraph()

tagline = doc.add_paragraph('Pilot Validation Programme')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.runs[0].font.size = Pt(18)
tagline.runs[0].italic = True

doc.add_paragraph()
doc.add_paragraph()

# Document info
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Version: 2.0\n').bold = True
info.add_run('Date: February 2026\n')
info.add_run('Author: Frontera Product Team\n')
info.add_run('Testing Platform: Maze\n')
info.add_run('Status: Draft for Review')

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
add_heading('Table of Contents', 1)
toc_items = [
    ('1.', 'Executive Summary'),
    ('2.', 'Research Objectives'),
    ('3.', 'Product Overview'),
    ('4.', 'Research Strategy'),
    ('5.', 'Maze Test Architecture'),
    ('6.', 'Maze Block Specifications'),
    ('7.', 'Moderated Interview Studies'),
    ('8.', 'Expert Evaluation Framework'),
    ('9.', 'Participant Recruitment'),
    ('10.', 'Research Timeline'),
    ('11.', 'Success Metrics & Usability Scores'),
    ('12.', 'Budget & Resources'),
    ('13.', 'Appendices'),
]
for num, item in toc_items:
    p = doc.add_paragraph()
    p.add_run(f'{num} ').bold = True
    p.add_run(item)

doc.add_page_break()

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================
add_heading('1. Executive Summary', 1)

add_heading('Purpose', 2)
doc.add_paragraph(
    'This proposal outlines a comprehensive user research programme for the Product Strategy Coach pilot, '
    'leveraging Maze as the primary research platform. The programme is designed to validate three critical dimensions:'
)

bullet_points = [
    'Experience & Usability - Can testers navigate the interface effectively?',
    'Journey Completion - Do testers progress through all methodology phases?',
    'Coaching Quality & Outputs - Is the AI coaching valuable? Are strategic outputs actionable?',
]
for point in bullet_points:
    doc.add_paragraph(point, style='List Bullet')

add_heading('Research Approach', 2)
doc.add_paragraph('A hybrid research methodology combining:')
approaches = [
    'Maze Unmoderated Testing - Scalable prototype tests with quantitative metrics',
    'Maze Moderated Interview Studies - In-depth sessions for qualitative insights',
    'Expert Panel Evaluation - Strategic output quality assessment',
]
for approach in approaches:
    doc.add_paragraph(approach, style='List Bullet')

add_heading('Why Maze?', 2)
doc.add_paragraph(
    'Maze provides an integrated research platform combining prototype testing, surveys, and participant '
    'recruitment. Key advantages for this research programme include:'
)
advantages = [
    'Automated usability scoring (0-100) based on success rate, duration, and misclicks',
    'Click heatmaps and path analysis for identifying friction points',
    'Direct/indirect success tracking for understanding user journeys',
    'Maze Panel access to 3+ million pre-screened testers across 130+ countries',
    '400+ demographic filters for precise participant targeting',
    'Integrated question blocks for mixed-methods research',
]
for adv in advantages:
    doc.add_paragraph(adv, style='List Bullet')

add_heading('Key Outcomes', 2)
outcomes = [
    'Quantified usability scores for each phase of the Product Strategy Coach',
    'Identification and prioritisation of friction points via heatmap analysis',
    'Validated coaching methodology resonance with target users',
    'Assessment of perceived value of AI-generated strategic outputs',
    'Actionable insights for pre-launch iteration',
]
for outcome in outcomes:
    doc.add_paragraph(outcome, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 2. RESEARCH OBJECTIVES
# ============================================================================
add_heading('2. Research Objectives', 1)

add_heading('Primary Objectives', 2)
add_table_from_data(
    ['Objective', 'Success Criteria', 'Maze Measurement'],
    [
        ['Validate navigation clarity', 'Usability Score ≥75', 'Prototype test blocks'],
        ['Confirm methodology comprehension', '>80% direct success on key missions', 'Mission success rates'],
        ['Assess coach interaction value', 'Trust score ≥4.0/5.0', 'Opinion scale blocks'],
        ['Evaluate synthesis output quality', 'Actionability score ≥4.0/5.0', 'Opinion scale + open questions'],
        ['Identify critical friction points', 'Misclick rate <15%', 'Heatmap analysis'],
    ]
)

add_heading('Research Questions', 2)

doc.add_paragraph('Experience & Usability', style='Intense Quote')
questions = [
    'How intuitive is the two-panel layout (coach + canvas)?',
    'Do testers understand the relationship between chat coaching and canvas activities?',
    'Can testers locate key actions (upload materials, explore territories, generate synthesis)?',
    'Is the progress stepper effective at communicating journey state?',
    'What is the misclick rate on primary navigation elements?',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph('Journey & Methodology', style='Intense Quote')
questions = [
    'Do testers understand the 4-phase methodology without explanation?',
    'What paths do testers take through the research territories?',
    'Where do testers bounce or give up on missions?',
    'Is the "terrain mapping" metaphor meaningful?',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph('Coaching Quality', style='Intense Quote')
questions = [
    'Do testers perceive the coach as helpful vs. intrusive?',
    'Is the coaching personalised enough to feel relevant?',
    'Do testers trust AI-generated strategic guidance?',
    'What coaching behaviours increase or decrease trust?',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph('Output Value', style='Intense Quote')
questions = [
    'Are synthesis outputs immediately understandable?',
    'Would testers share these outputs with stakeholders?',
    'Do outputs feel generic or specifically tailored?',
    'What additional outputs would testers want?',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

add_heading('Hypotheses', 2)
doc.add_paragraph(
    'The following hypotheses will be validated through the research programme:'
)
hypotheses = [
    'H1: Testers will achieve >75% direct success rate on Discovery phase missions',
    'H2: The two-panel layout will be understood without instruction (Opinion Scale ≥4/5)',
    'H3: Testers will trust AI coaching recommendations (Trust Score ≥4/5)',
    'H4: Synthesis outputs will be perceived as actionable (Actionability ≥4/5)',
    'H5: Testers will complete the full journey in <45 minutes',
]
for h in hypotheses:
    doc.add_paragraph(h, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 3. PRODUCT OVERVIEW
# ============================================================================
add_heading('3. Product Overview', 1)

add_heading('Product Description', 2)
doc.add_paragraph(
    'The Product Strategy Coach is an AI-powered strategic coaching platform that guides enterprise '
    'leaders through a structured methodology to develop product transformation strategies. The experience combines:'
)
features = [
    'Persistent AI Coach Sidebar for contextual guidance and Q&A',
    'Scrollable Canvas with phased, card-based workflow',
    'Progressive Disclosure through "terrain mapping" that unlocks as users complete phases',
    'Document Upload to ground strategy in real source materials',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

add_heading('Methodology Phases', 2)
add_table_from_data(
    ['Phase', 'Name', 'Purpose', 'Key Activities'],
    [
        ['1', 'Discovery', 'Context Setting', 'Upload materials, establish strategic baseline'],
        ['2', 'Research', 'Terrain Mapping', 'Explore Company, Customer, Competitor territories'],
        ['3', 'Synthesis', 'Strategy Formation', 'Review AI-generated opportunities and insights'],
        ['4', 'Strategic Bets', 'Route Planning', 'Define hypothesis-driven strategic bets'],
    ]
)

add_heading('Target Testers', 2)
doc.add_paragraph('Primary Persona: Strategic Product Leader', style='Intense Quote')
add_table_from_data(
    ['Attribute', 'Requirement'],
    [
        ['Job Title', 'VP Product, Head of Strategy, Product Director, CPO'],
        ['Company Size', '200+ employees (enterprise focus)'],
        ['Experience', '5+ years in product or strategy roles'],
        ['Current Context', 'Actively involved in strategic initiatives'],
        ['Industry', 'Technology, Financial Services, Healthcare, B2B SaaS'],
    ]
)

doc.add_page_break()

# ============================================================================
# 4. RESEARCH STRATEGY
# ============================================================================
add_heading('4. Research Strategy', 1)

add_heading('Research Layers', 2)

doc.add_paragraph('LAYER 1: Unmoderated Maze Testing', style='Intense Quote')
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Scalable quantitative data collection with automated usability scoring')
layer1 = [
    '6 mazes targeting specific phases and interactions',
    '15-20 testers per maze via Maze Panel',
    'Prototype test blocks with goal-based missions',
    'Question blocks for attitudinal data (Opinion Scales, Multiple Choice, Open Questions)',
    'Automated metrics: Usability Score, Direct/Indirect Success, Misclick Rate, Time on Task',
    'Heatmaps and path analysis for identifying friction',
]
for item in layer1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('LAYER 2: Moderated Interview Studies', style='Intense Quote')
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Deep qualitative insights through guided exploration')
layer2 = [
    '6-8 moderated sessions via Maze Interview Studies',
    '45-60 minute think-aloud sessions',
    'Real context (tester\'s actual company)',
    'Deep observation of coaching interaction',
    'Rich qualitative insights on trust and value perception',
]
for item in layer2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('LAYER 3: Expert Evaluation', style='Intense Quote')
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Strategic output quality assessment by domain experts')
layer3 = [
    'Panel of 3-5 strategy experts',
    'Blind review of synthesis outputs',
    'Structured scoring rubric for coaching and output quality',
    'Comparative assessment vs. human consultant outputs',
]
for item in layer3:
    doc.add_paragraph(item, style='List Bullet')

add_heading('Layer Comparison', 2)
add_table_from_data(
    ['Research Need', 'Unmoderated Maze', 'Moderated Studies', 'Expert Review'],
    [
        ['Navigation accuracy', 'Excellent', 'Good', 'N/A'],
        ['Task completion rates', 'Excellent', 'Good', 'N/A'],
        ['Click patterns & heatmaps', 'Excellent', 'Limited', 'N/A'],
        ['Emotional response', 'Limited', 'Excellent', 'N/A'],
        ['Trust assessment', 'Moderate', 'Excellent', 'Moderate'],
        ['Coaching quality', 'Limited', 'Excellent', 'Excellent'],
        ['Output actionability', 'Moderate', 'Good', 'Excellent'],
        ['Long-form journey', 'Limited', 'Excellent', 'N/A'],
        ['Scale (many testers)', 'Excellent', 'Limited', 'N/A'],
        ['Cost efficiency', 'High', 'Moderate', 'Moderate'],
    ]
)

doc.add_page_break()

# ============================================================================
# 5. MAZE TEST ARCHITECTURE
# ============================================================================
add_heading('5. Maze Test Architecture', 1)

add_heading('Maze Overview', 2)
add_table_from_data(
    ['Maze #', 'Name', 'Focus Area', 'Est. Duration', 'Testers'],
    [
        ['1', 'First Impressions', 'Overall comprehension & emotional response', '5 min', '20'],
        ['2', 'Discovery Phase', 'Context setting UX & coach interaction', '8 min', '20'],
        ['3', 'Coach Evaluation', 'AI coaching perception & trust', '6 min', '20'],
        ['4', 'Research Navigation', 'Territory exploration UX', '8 min', '20'],
        ['5', 'Synthesis Review', 'Output comprehension & value', '7 min', '20'],
        ['6', 'End-to-End Journey', 'Full flow (abbreviated prototype)', '12 min', '15'],
    ]
)

add_heading('Block Types Used', 2)
doc.add_paragraph(
    'Each maze combines multiple block types to gather both behavioural and attitudinal data:'
)
add_table_from_data(
    ['Block Type', 'Purpose', 'Data Collected'],
    [
        ['Welcome Block', 'Introduce maze and set context', 'N/A'],
        ['Prototype Test (Goal-based)', 'Test specific task flows with success criteria', 'Usability Score, Direct/Indirect Success, Misclick Rate, Time on Task, Heatmaps'],
        ['Prototype Test (Free Explore)', 'Observe natural navigation patterns', 'Click paths, Time on Task, Heatmaps'],
        ['5-Second Test', 'Capture first impressions of screens', 'Recall accuracy, Emotional response'],
        ['Opinion Scale', 'Rate agreement/satisfaction (0-10)', 'Quantitative ratings, NPS-style scores'],
        ['Multiple Choice', 'Categorise responses, single/multi-select', 'Response distribution, Preferences'],
        ['Open Question', 'Collect qualitative feedback', 'Themes, Quotes, Suggestions'],
        ['Yes/No', 'Binary comprehension checks', 'Pass/fail rates'],
    ]
)

add_heading('Maze Metrics Captured', 2)
add_table_from_data(
    ['Metric', 'Definition', 'Target'],
    [
        ['Usability Score', 'Composite score (0-100) based on success, duration, misclicks', '≥75'],
        ['Direct Success', 'Completed mission via expected path', '≥70%'],
        ['Indirect Success', 'Completed mission via unexpected path', '<20%'],
        ['Bounce Rate', 'Gave up or abandoned mission', '<10%'],
        ['Misclick Rate', 'Clicks outside hotspots / total clicks', '<15%'],
        ['Time on Task', 'Duration to complete mission', 'Varies by task'],
        ['Opinion Scale Avg', 'Mean rating on 0-10 scale', '≥7.0 (scaled to 4/5)'],
    ]
)

doc.add_page_break()

# ============================================================================
# 6. MAZE BLOCK SPECIFICATIONS
# ============================================================================
add_heading('6. Maze Block Specifications', 1)

# MAZE 1
add_heading('Maze 1: First Impressions', 2)
doc.add_paragraph('Objective: Assess initial comprehension and emotional response to the interface')
doc.add_paragraph('Duration: ~5 minutes | Testers: 20')

add_heading('Block Structure', 3)
add_table_from_data(
    ['Block #', 'Type', 'Content'],
    [
        ['1', 'Welcome Block', 'Introduction to the maze and tester instructions'],
        ['2', '5-Second Test', 'Show full interface screenshot'],
        ['3', 'Open Question', 'What do you think this product does?'],
        ['4', 'Multiple Choice', 'Emotional response selection'],
        ['5', 'Multiple Choice', 'Product purpose comprehension check'],
        ['6', 'Open Question', 'What is the left sidebar for?'],
        ['7', 'Prototype Test (Goal)', 'First-click: Where would you start?'],
        ['8', 'Multiple Choice', 'Progress stepper comprehension'],
        ['9', 'Opinion Scale', 'Value clarity (0-10)'],
        ['10', 'Opinion Scale', 'Relevance to work (0-10)'],
        ['11', 'Open Question', 'What questions do you have?'],
    ]
)

add_heading('Block Details', 3)

doc.add_paragraph('Block 2: 5-Second Test', style='Intense Quote')
doc.add_paragraph('Image: Full interface screenshot (coach sidebar + canvas with Discovery phase)')
doc.add_paragraph('Purpose: Capture immediate impression before cognitive processing')

doc.add_paragraph('Block 3: Open Question', style='Intense Quote')
doc.add_paragraph('Question: "Based on what you just saw, what do you think this product does?"')
doc.add_paragraph('Purpose: Validate product concept clarity')

doc.add_paragraph('Block 4: Multiple Choice (Single Select)', style='Intense Quote')
doc.add_paragraph('Question: "What\'s your immediate reaction to this interface?"')
doc.add_paragraph('Options:')
options = [
    'Excited - I want to explore this',
    'Curious - I\'d like to learn more',
    'Neutral - No strong feeling',
    'Confused - I\'m not sure what I\'m looking at',
    'Overwhelmed - This looks complex',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Block 5: Multiple Choice (Single Select)', style='Intense Quote')
doc.add_paragraph('Question: "Based on what you see, what is the main purpose of this tool?"')
doc.add_paragraph('Options:')
options = [
    'A project management dashboard',
    'An AI-powered strategy coaching tool (Correct)',
    'A document collaboration platform',
    'A customer analytics system',
    'I\'m not sure',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Block 7: Prototype Test (Goal-Based)', style='Intense Quote')
doc.add_paragraph('Mission: "If you wanted to start working on your strategy, where would you click first?"')
doc.add_paragraph('Goal Screen: Discovery section / Coach input / Upload materials button')
doc.add_paragraph('Success Criteria: Click within defined hotspots within 30 seconds')
doc.add_paragraph('Metrics: Direct Success Rate, Misclick Rate, Time to First Click, Heatmap')

doc.add_paragraph('Block 8: Multiple Choice (Single Select)', style='Intense Quote')
doc.add_paragraph('Question: "Looking at the circles at the top of the main panel, what do they represent?"')
doc.add_paragraph('Options:')
options = [
    'Steps in a process I need to complete (Correct)',
    'Different features I can access',
    'My team members',
    'Notification indicators',
    'Not sure',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Block 9 & 10: Opinion Scale', style='Intense Quote')
doc.add_paragraph('Question 9: "How clear is it what value this tool would provide to you?"')
doc.add_paragraph('Question 10: "How relevant does this tool seem for your work?"')
doc.add_paragraph('Scale: 0 (Not at all) to 10 (Extremely)')
doc.add_paragraph('Labels: Left: "Not at all clear/relevant" | Right: "Extremely clear/relevant"')

doc.add_page_break()

# MAZE 2
add_heading('Maze 2: Discovery Phase Deep Dive', 2)
doc.add_paragraph('Objective: Validate Discovery phase usability and coach interaction')
doc.add_paragraph('Duration: ~8 minutes | Testers: 20')

add_heading('Block Structure', 3)
add_table_from_data(
    ['Block #', 'Type', 'Content'],
    [
        ['1', 'Welcome Block', 'Scenario: VP Product at fintech company'],
        ['2', 'Prototype Test (Goal)', 'Mission: Locate document upload function'],
        ['3', 'Opinion Scale', 'Ease of finding upload (0-10)'],
        ['4', 'Multiple Choice (Multi)', 'What documents would you upload?'],
        ['5', 'Prototype Test (Goal)', 'Mission: Respond to coach question'],
        ['6', 'Multiple Choice', 'Coach tone perception'],
        ['7', 'Opinion Scale', 'Coach question relevance (0-10)'],
        ['8', 'Open Question', 'What do you need to do to proceed?'],
        ['9', 'Opinion Scale', 'Progress requirement clarity (0-10)'],
        ['10', 'Opinion Scale', 'Overall Discovery satisfaction (0-10)'],
        ['11', 'Open Question', 'What was confusing or frustrating?'],
        ['12', 'Open Question', 'What guidance would help?'],
    ]
)

add_heading('Block Details', 3)

doc.add_paragraph('Block 1: Welcome Block', style='Intense Quote')
scenario = (
    'Scenario: You\'re the VP of Product at a mid-sized fintech company. Your CEO has asked you '
    'to develop a product strategy for the next 18 months. You\'ve decided to use this AI coaching '
    'tool to help structure your thinking. You\'ve just logged in for the first time.'
)
doc.add_paragraph(scenario)

doc.add_paragraph('Block 2: Prototype Test (Goal-Based)', style='Intense Quote')
doc.add_paragraph('Mission: "Find where you would upload your company\'s strategic documents."')
doc.add_paragraph('Expected Path: Welcome → Discovery Section → Upload Materials Card → Upload Button')
doc.add_paragraph('Goal Screen: Upload button or drag-drop zone')
doc.add_paragraph('Time Limit: 30 seconds')
doc.add_paragraph('Metrics: Direct Success, Indirect Success, Bounce, Misclick Rate, Time on Task')

doc.add_paragraph('Block 4: Multiple Choice (Multi-Select)', style='Intense Quote')
doc.add_paragraph('Question: "What documents would you upload if using this for your company?"')
doc.add_paragraph('Options:')
options = [
    'Annual report / financial statements',
    'Strategic plan / roadmap',
    'Market research reports',
    'Competitor analysis',
    'Customer feedback / NPS data',
    'Product specs / PRDs',
    'Board presentations',
    'I wouldn\'t know what to upload',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Block 5: Prototype Test (Goal-Based)', style='Intense Quote')
doc.add_paragraph('Mission: "The AI coach has asked you a question. Respond to it."')
doc.add_paragraph('Context Screen: Coach message asking "What competitive dynamics are making product transformation urgent for your company?"')
doc.add_paragraph('Goal Screen: Input field focused or response submitted')
doc.add_paragraph('Metrics: Engagement rate, Time to engage')

doc.add_paragraph('Block 6: Multiple Choice (Single Select)', style='Intense Quote')
doc.add_paragraph('Question: "How would you describe the tone of the AI coach?"')
doc.add_paragraph('Options:')
options = [
    'Professional and authoritative',
    'Friendly and supportive',
    'Neutral and factual',
    'Pushy or demanding',
    'Robotic and impersonal',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_page_break()

# MAZE 3
add_heading('Maze 3: Coach Evaluation', 2)
doc.add_paragraph('Objective: Assess AI coaching perception, value, and trust')
doc.add_paragraph('Duration: ~6 minutes | Testers: 20')

add_heading('Block Structure', 3)
add_table_from_data(
    ['Block #', 'Type', 'Content'],
    [
        ['1', 'Welcome Block', 'Introduction to coaching evaluation'],
        ['2', '5-Second Test', 'Coach interaction screenshot'],
        ['3', 'Opinion Scale', 'Helpfulness rating (0-10)'],
        ['4', 'Opinion Scale', 'Relevance rating (0-10)'],
        ['5', 'Multiple Choice', 'Interaction style perception'],
        ['6', '5-Second Test', 'Scenario A: Proactive insight'],
        ['7', 'Opinion Scale', 'Proactive coaching value (0-10)'],
        ['8', 'Opinion Scale', 'Trust in insight (0-10)'],
        ['9', '5-Second Test', 'Scenario B: Assumption challenge'],
        ['10', 'Opinion Scale', 'Comfort with challenge (0-10)'],
        ['11', 'Multiple Choice', 'Challenge perception'],
        ['12', 'Opinion Scale', 'Trust vs. human consultant (0-10)'],
        ['13', 'Multiple Choice (Multi)', 'Behaviours that increase trust'],
        ['14', 'Open Question', 'What would make coaching more valuable?'],
    ]
)

add_heading('Coaching Scenarios', 3)

doc.add_paragraph('Scenario A: Proactive Insight (Block 6)', style='Intense Quote')
doc.add_paragraph(
    'Coach message: "I notice you mentioned \'integration depth\' as a customer priority, but your '
    'competitor analysis suggests most alternatives struggle with this. This could be a significant '
    '\'How to Win\' opportunity. Would you like to explore what capabilities you\'d need to excel here?"'
)

doc.add_paragraph('Scenario B: Assumption Challenge (Block 9)', style='Intense Quote')
doc.add_paragraph(
    'Coach message: "You mentioned that \'price is not a key differentiator\' for your customers. '
    'However, your earlier notes indicated several lost deals due to pricing. Can you help me '
    'understand this apparent tension?"'
)

doc.add_paragraph('Block 11: Multiple Choice (Single Select)', style='Intense Quote')
doc.add_paragraph('Question: "When the coach challenges your assumptions, it feels:"')
doc.add_paragraph('Options:')
options = [
    'Helpful - it makes me think more deeply',
    'Neutral - neither good nor bad',
    'Uncomfortable - I don\'t like being challenged by AI',
    'Annoying - the AI doesn\'t understand my context',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Block 12: Opinion Scale', style='Intense Quote')
doc.add_paragraph('Question: "Compared to a human strategy consultant, how much would you trust this AI coach?"')
doc.add_paragraph('Scale: 0-10')
doc.add_paragraph('Labels: 0 = "Much less than human" | 5 = "About the same" | 10 = "Much more than human"')

doc.add_paragraph('Block 13: Multiple Choice (Multi-Select)', style='Intense Quote')
doc.add_paragraph('Question: "What behaviours would make you trust the AI coach more?"')
doc.add_paragraph('Options:')
options = [
    'Showing its reasoning/sources',
    'Asking clarifying questions before advising',
    'Admitting when it doesn\'t know something',
    'Providing options rather than single recommendations',
    'Referencing industry benchmarks or data',
    'Remembering context from earlier in the session',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_page_break()

# MAZE 4
add_heading('Maze 4: Research Navigation', 2)
doc.add_paragraph('Objective: Validate Research phase navigation and territory exploration')
doc.add_paragraph('Duration: ~8 minutes | Testers: 20')

add_heading('Block Structure', 3)
add_table_from_data(
    ['Block #', 'Type', 'Content'],
    [
        ['1', 'Welcome Block', 'Scenario: Ready to map strategic terrain'],
        ['2', 'Prototype Test (Goal)', 'Mission: Navigate to Customer territory'],
        ['3', 'Opinion Scale', 'Territory selection clarity (0-10)'],
        ['4', 'Prototype Test (Goal)', 'Mission: Find Segments & Needs area'],
        ['5', 'Opinion Scale', 'Navigation ease (0-10)'],
        ['6', 'Prototype Test (Free Explore)', 'Explore research questions'],
        ['7', 'Opinion Scale', 'Question clarity (0-10)'],
        ['8', 'Opinion Scale', 'Question relevance (0-10)'],
        ['9', 'Multiple Choice', 'Input format preference'],
        ['10', 'Open Question', 'How would you track progress?'],
        ['11', 'Opinion Scale', 'Progress visibility (0-10)'],
        ['12', 'Opinion Scale', '"Terrain mapping" metaphor clarity (0-10)'],
        ['13', 'Multiple Choice', '3Cs category assessment'],
        ['14', 'Open Question', 'What other territories would you explore?'],
        ['15', 'Opinion Scale', 'Research phase satisfaction (0-10)'],
        ['16', 'Open Question', 'What would make research more engaging?'],
    ]
)

add_heading('Block Details', 3)

doc.add_paragraph('Block 2: Prototype Test (Goal-Based)', style='Intense Quote')
doc.add_paragraph('Mission: "You want to understand your customers better. Navigate to the Customer territory."')
doc.add_paragraph('Expected Path: Research Section → Customer Territory Card → Click')
doc.add_paragraph('Goal Screen: Customer territory deep dive view')
doc.add_paragraph('Metrics: Direct Success, Heatmap of clicks on territory cards')

doc.add_paragraph('Block 4: Prototype Test (Goal-Based)', style='Intense Quote')
doc.add_paragraph('Mission: "Find where you would answer questions about customer segments."')
doc.add_paragraph('Expected Path: Customer Territory → Segments & Needs Research Area')
doc.add_paragraph('Goal Screen: Segments & Needs question view')

doc.add_paragraph('Block 6: Prototype Test (Free Explore)', style='Intense Quote')
doc.add_paragraph('Instruction: "Take 60 seconds to explore the research questions. Click around naturally."')
doc.add_paragraph('Purpose: Observe natural navigation patterns without defined goals')
doc.add_paragraph('Metrics: Click paths, Areas of interest, Time distribution')

doc.add_paragraph('Block 9: Multiple Choice (Single Select)', style='Intense Quote')
doc.add_paragraph('Question: "The text input format for answering research questions is:"')
doc.add_paragraph('Options:')
options = [
    'Perfect - text input works well',
    'Acceptable - but I\'d prefer bullet points',
    'Acceptable - but I\'d prefer voice input',
    'Frustrating - I want a different method',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Block 13: Multiple Choice (Single Select)', style='Intense Quote')
doc.add_paragraph('Question: "The three territories are Company, Customer, and Competitor (the \'3Cs\'). These categories are:"')
doc.add_paragraph('Options:')
options = [
    'Exactly right - these are the key areas to explore',
    'Mostly right - but I\'d add or change something',
    'Somewhat relevant - but missing important areas',
    'Not how I think about strategy',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_page_break()

# MAZE 5
add_heading('Maze 5: Synthesis Review', 2)
doc.add_paragraph('Objective: Assess comprehension and perceived value of synthesis outputs')
doc.add_paragraph('Duration: ~7 minutes | Testers: 20')

add_heading('Block Structure', 3)
add_table_from_data(
    ['Block #', 'Type', 'Content'],
    [
        ['1', 'Welcome Block', 'Introduction to synthesis outputs'],
        ['2', '5-Second Test', 'Strategic Opportunity Map (2x2)'],
        ['3', 'Open Question', 'What does this visualisation show?'],
        ['4', 'Multiple Choice', 'INVEST quadrant meaning'],
        ['5', 'Opinion Scale', 'Visualisation usefulness (0-10)'],
        ['6', 'Opinion Scale', 'Quadrant clarity (0-10)'],
        ['7', '5-Second Test', 'Opportunity card detail'],
        ['8', 'Open Question', 'What action is recommended?'],
        ['9', 'Opinion Scale', 'Evidence trail value (0-10)'],
        ['10', 'Opinion Scale', 'WWHBT usefulness (0-10)'],
        ['11', 'Multiple Choice', 'Actionability assessment'],
        ['12', 'Open Question', 'What additional info needed?'],
        ['13', 'Multiple Choice', 'Personalisation perception'],
        ['14', 'Opinion Scale', 'Stakeholder shareability (0-10)'],
        ['15', 'Open Question', 'Why would/wouldn\'t you share?'],
        ['16', 'Multiple Choice (Multi)', 'Additional outputs wanted'],
        ['17', 'Open Question', 'What would make outputs more valuable?'],
    ]
)

add_heading('Block Details', 3)

doc.add_paragraph('Block 2: 5-Second Test', style='Intense Quote')
doc.add_paragraph('Image: Strategic Opportunity Map (2x2 matrix) with 4 plotted opportunities')
doc.add_paragraph('Quadrants: INVEST (top-right), EXPLORE (top-left), HARVEST (bottom-right), DIVEST (bottom-left)')
doc.add_paragraph('Axes: X = Capability Fit, Y = Market Attractiveness')

doc.add_paragraph('Block 4: Multiple Choice (Single Select)', style='Intense Quote')
doc.add_paragraph('Question: "What does an opportunity in the top-right quadrant (INVEST) mean?"')
doc.add_paragraph('Options:')
options = [
    'High priority - execute immediately (Correct)',
    'Interesting but needs more research',
    'Low priority - deprioritise',
    'I\'m not sure',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Block 7: 5-Second Test', style='Intense Quote')
doc.add_paragraph('Image: Detailed opportunity card showing:')
details = [
    'Title: "Enterprise Integration Platform"',
    'Scores: Market Attractiveness 8/10, Capability Fit 7/10',
    'Where to Play: "Mid-market consolidators (£5-50bn AUM)"',
    'How to Win: "Deep platform integration + compliance automation"',
    'Evidence: 3 linked sources from research',
    'WWHBT: "What Would Have to Be True" assumptions',
]
for d in details:
    doc.add_paragraph(d, style='List Bullet')

doc.add_paragraph('Block 11: Multiple Choice (Single Select)', style='Intense Quote')
doc.add_paragraph('Question: "Could you take action on this strategic opportunity based on the information provided?"')
doc.add_paragraph('Options:')
options = [
    'Yes, I have enough information to act',
    'Partially, but I need more detail on some areas',
    'No, I need significantly more information',
    'I\'m not sure what action to take',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Block 14: Opinion Scale (NPS-style)', style='Intense Quote')
doc.add_paragraph('Question: "Would you share this synthesis output with your executive team?"')
doc.add_paragraph('Scale: 0 (Definitely not) to 10 (Definitely yes)')

doc.add_paragraph('Block 16: Multiple Choice (Multi-Select)', style='Intense Quote')
doc.add_paragraph('Question: "What other outputs would you want from the synthesis phase?"')
doc.add_paragraph('Options:')
options = [
    'SWOT analysis',
    'Competitive positioning map',
    'Risk assessment',
    'Implementation timeline',
    'Resource requirements',
    'Financial projections',
    'Stakeholder communication plan',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_page_break()

# MAZE 6
add_heading('Maze 6: End-to-End Journey', 2)
doc.add_paragraph('Objective: Validate full journey flow with abbreviated prototype')
doc.add_paragraph('Duration: ~12 minutes | Testers: 15')

add_heading('Block Structure', 3)
add_table_from_data(
    ['Block #', 'Type', 'Content'],
    [
        ['1', 'Welcome Block', 'Full journey introduction'],
        ['2', 'Prototype Test (Goal)', 'Phase 1: Complete Discovery'],
        ['3', 'Prototype Test (Goal)', 'Phase 2: Complete one territory'],
        ['4', 'Prototype Test (Goal)', 'Phase 3: Review synthesis'],
        ['5', 'Prototype Test (Goal)', 'Phase 4: Create strategic bet'],
        ['6', 'Opinion Scale', 'Journey clarity (0-10)'],
        ['7', 'Opinion Scale', 'Phase transition smoothness (0-10)'],
        ['8', 'Multiple Choice', 'Time investment value'],
        ['9', 'Opinion Scale', 'Completion likelihood (0-10)'],
        ['10', 'Opinion Scale', 'Overall satisfaction (0-10)'],
        ['11', 'Opinion Scale', 'NPS: Likelihood to recommend (0-10)'],
        ['12', 'Open Question', 'Most valuable aspect'],
        ['13', 'Open Question', 'Biggest concern/hesitation'],
        ['14', 'Open Question', 'One thing to change'],
        ['15', 'Multiple Choice', 'Willingness to pay'],
    ]
)

add_heading('Mission Specifications', 3)

doc.add_paragraph('Mission 1: Discovery Phase (Block 2)', style='Intense Quote')
doc.add_paragraph('Task: "Complete the Discovery phase by uploading a document and answering the coach\'s question."')
doc.add_paragraph('Expected Path: Upload → Coach Response → Continue')
doc.add_paragraph('Time Allowance: 2 minutes')

doc.add_paragraph('Mission 2: Research Phase (Block 3)', style='Intense Quote')
doc.add_paragraph('Task: "Complete one territory (Customer) by answering at least one research area."')
doc.add_paragraph('Expected Path: Customer Card → Research Area → Save Response')
doc.add_paragraph('Time Allowance: 3 minutes')

doc.add_paragraph('Mission 3: Synthesis Phase (Block 4)', style='Intense Quote')
doc.add_paragraph('Task: "Review the synthesis and identify the top strategic opportunity."')
doc.add_paragraph('Expected Path: View Map → Click Opportunity → Review Details')
doc.add_paragraph('Time Allowance: 2 minutes')

doc.add_paragraph('Mission 4: Strategic Bets Phase (Block 5)', style='Intense Quote')
doc.add_paragraph('Task: "Create a strategic bet based on the synthesis."')
doc.add_paragraph('Expected Path: Create Bet → Fill Details → Save')
doc.add_paragraph('Time Allowance: 2 minutes')

doc.add_paragraph('Block 8: Multiple Choice (Single Select)', style='Intense Quote')
doc.add_paragraph('Question: "For the value provided, the time investment in this process is:"')
doc.add_paragraph('Options:')
options = [
    'Worth much more than the time required',
    'Worth the time required',
    'Worth somewhat less than the time required',
    'Not worth the time required',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Block 15: Multiple Choice (Single Select)', style='Intense Quote')
doc.add_paragraph('Question: "Would you or your company pay for a tool like this?"')
doc.add_paragraph('Options:')
options = [
    'Yes, definitely',
    'Yes, probably',
    'Not sure',
    'Probably not',
    'Definitely not',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 7. MODERATED INTERVIEW STUDIES
# ============================================================================
add_heading('7. Moderated Interview Studies', 1)

add_heading('Overview', 2)
doc.add_paragraph(
    'Maze\'s Moderated Interview Studies (previously "Live Sessions") enable real-time, guided research '
    'with screen sharing and recording. These complement unmoderated mazes by providing:'
)
benefits = [
    'Deep qualitative insights through think-aloud observation',
    'Ability to probe unexpected behaviours and emotional responses',
    'Testing with real company context (tester\'s actual data)',
    'Rich understanding of trust and value perception',
]
for b in benefits:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Session Configuration', 2)
add_table_from_data(
    ['Element', 'Details'],
    [
        ['Duration', '45-60 minutes'],
        ['Format', 'Video call with screen sharing (via Maze)'],
        ['Recording', 'Video + audio + screen (with consent)'],
        ['Testers', '6-8 strategic product leaders'],
        ['Recruitment', 'Maze Panel with custom screening'],
    ]
)

add_heading('Screening Questions (Required)', 2)
doc.add_paragraph(
    'Maze requires at least 2 screening questions for moderated panel recruitment. '
    'Recommended screening criteria:'
)

doc.add_paragraph('Question 1: Role Confirmation', style='Intense Quote')
doc.add_paragraph('What is your current job title?')
options = [
    'VP/Head/Director of Product or Strategy → Qualify',
    'Senior Product Manager → Qualify',
    'Product Manager → Qualify with caution',
    'Other → Screen out',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Question 2: Strategic Involvement', style='Intense Quote')
doc.add_paragraph('In the past 12 months, have you led or significantly contributed to product strategy development?')
options = [
    'Yes, I led the process → Ideal',
    'Yes, I was a key contributor → Qualify',
    'No → Screen out',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_paragraph('Question 3: Company Size', style='Intense Quote')
doc.add_paragraph('How many employees does your company have?')
options = [
    '1-49 → Screen out (unless funded startup)',
    '50-199 → Qualify with caution',
    '200-999 → Qualify',
    '1000+ → Ideal',
]
for opt in options:
    doc.add_paragraph(opt, style='List Bullet')

add_heading('Session Structure', 2)
add_table_from_data(
    ['Section', 'Duration', 'Purpose'],
    [
        ['1. Introduction', '5 min', 'Set expectations, obtain consent'],
        ['2. Context Setting', '5 min', 'Understand tester background'],
        ['3. First Impressions', '5 min', 'Initial interface reaction'],
        ['4. Discovery Phase', '10-15 min', 'Walkthrough with real context'],
        ['5. Research Phase', '15-20 min', 'Territory exploration'],
        ['6. Synthesis Review', '10 min', 'Output evaluation'],
        ['7. Debrief Interview', '10-15 min', 'Structured feedback'],
    ]
)

add_heading('Facilitator Script: Introduction', 3)
script = doc.add_paragraph()
script.add_run(
    '"Thank you for joining today. I\'m [name], and I\'ll be facilitating this session.\n\n'
    'We\'re developing an AI-powered tool to help product leaders develop strategy. Today, I\'d like '
    'you to use the tool and share your thoughts as you go.\n\n'
    'A few things before we start:\n'
    '• Please think aloud - share what you\'re looking at, thinking, and feeling\n'
    '• There are no wrong answers - if something is confusing, that\'s valuable feedback\n'
    '• We\'re testing the product, not you\n'
    '• Feel free to be candid - criticism helps us improve\n\n'
    'Do you have any questions before we begin?"'
)
script.italic = True

add_heading('During-Session Observation Prompts', 3)
add_table_from_data(
    ['Situation', 'Prompt'],
    [
        ['Tester pauses', '"What are you looking at right now?"'],
        ['Tester seems confused', '"What are you expecting to see here?"'],
        ['Tester clicks unexpectedly', '"What made you click there?"'],
        ['Tester skips content', '"I noticed you scrolled past that - was it not relevant?"'],
        ['Tester sighs/frustrated', '"What\'s causing that reaction?"'],
        ['Tester smiles/satisfied', '"What do you like about this?"'],
        ['Tester ignores coach', '"I noticed you didn\'t engage with the coach - any reason?"'],
    ]
)

add_heading('Debrief Interview Questions', 3)

doc.add_paragraph('Overall Experience', style='Intense Quote')
questions = [
    '"How would you describe this experience in a few words?"',
    '"What was the most valuable part?"',
    '"What was the most frustrating part?"',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph('Coaching Quality', style='Intense Quote')
questions = [
    '"How did the AI coach compare to working with a human consultant?"',
    '"Did you trust the coach\'s guidance? Why or why not?"',
    '"What would make you trust it more?"',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph('Output Value', style='Intense Quote')
questions = [
    '"How actionable were the strategic outputs?"',
    '"What would you do next with this synthesis?"',
    '"What outputs are missing that you\'d want?"',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph('Adoption Likelihood', style='Intense Quote')
questions = [
    '"Would you use this tool for a real strategic project?"',
    '"What would need to change for you to use it?"',
    '"What would you be willing to pay for this?"',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 8. EXPERT EVALUATION FRAMEWORK
# ============================================================================
add_heading('8. Expert Evaluation Framework', 1)

add_heading('Purpose', 2)
doc.add_paragraph(
    'Expert evaluation assesses the quality of AI-generated coaching and synthesis outputs through '
    'blind review by strategy domain experts. This layer cannot be replicated through user testing alone.'
)

add_heading('Panel Composition', 2)
add_table_from_data(
    ['Attribute', 'Requirement'],
    [
        ['Panel Size', '3-5 experts'],
        ['Background Mix', 'Management consulting, corporate strategy, product leadership'],
        ['Experience', '10+ years in strategy development'],
        ['Prior Exposure', 'None (blind review)'],
    ]
)

add_heading('Scoring Rubric', 2)

doc.add_paragraph('Dimension 1: Coaching Relevance (1-5)', style='Intense Quote')
add_table_from_data(
    ['Score', 'Description'],
    [
        ['1 - Poor', 'Coaching completely misses the point; generic or off-topic'],
        ['2 - Below Average', 'Some relevant points but largely unhelpful'],
        ['3 - Adequate', 'Reasonably relevant but not insightful'],
        ['4 - Good', 'Relevant and occasionally insightful'],
        ['5 - Excellent', 'Highly relevant, insightful, and well-timed'],
    ]
)

doc.add_paragraph('Dimension 2: Strategic Actionability (1-5)', style='Intense Quote')
add_table_from_data(
    ['Score', 'Description'],
    [
        ['1 - Poor', 'Outputs are too vague to act upon'],
        ['2 - Below Average', 'Some actionable elements but mostly abstract'],
        ['3 - Adequate', 'Actionable but requires significant interpretation'],
        ['4 - Good', 'Clear actions with reasonable specificity'],
        ['5 - Excellent', 'Highly actionable with clear next steps'],
    ]
)

doc.add_paragraph('Dimension 3: Insight Quality (1-5)', style='Intense Quote')
add_table_from_data(
    ['Score', 'Description'],
    [
        ['1 - Poor', 'No insights; purely descriptive'],
        ['2 - Below Average', 'Obvious insights only'],
        ['3 - Adequate', 'Some non-obvious connections'],
        ['4 - Good', 'Multiple meaningful insights'],
        ['5 - Excellent', 'Breakthrough insights that reframe the problem'],
    ]
)

doc.add_paragraph('Dimension 4: Stakeholder Readiness (1-5)', style='Intense Quote')
add_table_from_data(
    ['Score', 'Description'],
    [
        ['1 - Poor', 'Would embarrass the user if shared'],
        ['2 - Below Average', 'Needs significant editing before sharing'],
        ['3 - Adequate', 'Shareable with caveats'],
        ['4 - Good', 'Ready to share with minor polish'],
        ['5 - Excellent', 'Could be presented to the board immediately'],
    ]
)

add_heading('Comparative Assessment', 2)
doc.add_paragraph(
    'For 3-5 sessions, create a "gold standard" synthesis manually. Present both AI and human-generated '
    'outputs to experts without labels.'
)
doc.add_paragraph('Questions for comparative review:', style='Intense Quote')
questions = [
    '"Which synthesis would you prefer to receive?"',
    '"Which feels more trustworthy?"',
    '"Which is more actionable?"',
    '"Can you tell which is AI-generated?"',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 9. PARTICIPANT RECRUITMENT
# ============================================================================
add_heading('9. Participant Recruitment', 1)

add_heading('Maze Panel Overview', 2)
doc.add_paragraph(
    'Maze Panel provides access to 3+ million pre-screened testers across 130+ countries. '
    'Key features for this research programme:'
)
features = [
    '400+ demographic and professional filters',
    'Custom screening questions for precise targeting',
    'Order between 1-250 testers per maze',
    'Partners: Prolific, User Interviews, Bilendi',
    'Integrated payment and scheduling',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

add_heading('Target Profile', 2)
add_table_from_data(
    ['Attribute', 'Filter/Requirement'],
    [
        ['Job Function', 'Product Management, Strategy, Executive'],
        ['Seniority', 'Director, VP, C-Level'],
        ['Company Size', '200+ employees'],
        ['Industry', 'Technology, Financial Services, Healthcare, B2B'],
        ['Experience', '5+ years in role'],
        ['Geography', 'UK, US, EU (English-speaking)'],
    ]
)

add_heading('Sample Size Recommendations', 2)
doc.add_paragraph(
    'Maze sample size recommendations based on research type and statistical confidence:'
)
add_table_from_data(
    ['Research Type', 'Minimum', 'Recommended', 'Notes'],
    [
        ['Unmoderated Prototype Test', '15', '20', '85% confidence for usability issues'],
        ['Unmoderated Survey/Opinion', '20', '30', 'Statistical significance for quantitative data'],
        ['Moderated Interview Study', '5', '6-8', 'Diminishing returns after 8 sessions'],
        ['Expert Panel Review', '3', '5', 'Inter-rater reliability'],
    ]
)

add_heading('Incentives', 2)
add_table_from_data(
    ['Test Type', 'Duration', 'Suggested Incentive'],
    [
        ['Unmoderated Maze (short)', '5-8 min', '$15-25'],
        ['Unmoderated Maze (medium)', '8-12 min', '$25-40'],
        ['Moderated Interview Study', '45-60 min', '$150-200'],
        ['Expert Panel Review', '30 min per review', '$100-150 per review'],
    ]
)

doc.add_page_break()

# ============================================================================
# 10. RESEARCH TIMELINE
# ============================================================================
add_heading('10. Research Timeline', 1)

add_heading('6-Week Programme Overview', 2)
add_table_from_data(
    ['Week', 'Activity', 'Deliverable'],
    [
        ['1', 'Preparation: Build mazes, configure panel', 'Ready-to-launch mazes'],
        ['2', 'Unmoderated Mazes 1-3', 'Live tests, initial data'],
        ['3', 'Unmoderated Mazes 4-6', 'Complete quantitative data'],
        ['4', 'Moderated Interview Studies', 'Recordings, transcripts'],
        ['5', 'Expert Review + Analysis', 'Integrated findings'],
        ['6', 'Reporting + Planning', 'Final report, recommendations'],
    ]
)

add_heading('Detailed Schedule', 2)

doc.add_paragraph('Week 1: Preparation', style='Intense Quote')
tasks = [
    'Finalise research plan and get stakeholder sign-off',
    'Build Mazes 1-3 in Maze platform',
    'Build Mazes 4-6 in Maze platform',
    'Configure Maze Panel screening criteria',
    'QA all mazes with internal pilot',
    'Schedule moderated interview slots',
]
for t in tasks:
    doc.add_paragraph(t, style='List Bullet')

doc.add_paragraph('Week 2: Unmoderated Mazes (Batch 1)', style='Intense Quote')
tasks = [
    'Launch Maze 1: First Impressions (20 testers)',
    'Launch Maze 2: Discovery Phase (20 testers)',
    'Launch Maze 3: Coach Evaluation (20 testers)',
    'Monitor completion rates and data quality',
    'Preliminary analysis of early results',
]
for t in tasks:
    doc.add_paragraph(t, style='List Bullet')

doc.add_paragraph('Week 3: Unmoderated Mazes (Batch 2)', style='Intense Quote')
tasks = [
    'Launch Maze 4: Research Navigation (20 testers)',
    'Launch Maze 5: Synthesis Review (20 testers)',
    'Launch Maze 6: End-to-End Journey (15 testers)',
    'Close Mazes 1-3, full analysis',
    'Prepare moderated session guide based on findings',
]
for t in tasks:
    doc.add_paragraph(t, style='List Bullet')

doc.add_paragraph('Week 4: Moderated Interview Studies', style='Intense Quote')
tasks = [
    'Conduct 6-8 moderated sessions via Maze',
    'Real-time note-taking and tagging',
    'Identify emerging themes',
    'Share preliminary insights with team',
]
for t in tasks:
    doc.add_paragraph(t, style='List Bullet')

doc.add_paragraph('Week 5: Expert Review + Analysis', style='Intense Quote')
tasks = [
    'Prepare synthesis outputs for expert review',
    'Conduct expert panel evaluation (3-5 experts)',
    'Analyse expert feedback',
    'Integrate all data sources (quantitative + qualitative + expert)',
]
for t in tasks:
    doc.add_paragraph(t, style='List Bullet')

doc.add_paragraph('Week 6: Reporting + Planning', style='Intense Quote')
tasks = [
    'Complete final research report',
    'Stakeholder presentation',
    'Prioritisation workshop',
    'Plan iteration sprint',
]
for t in tasks:
    doc.add_paragraph(t, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 11. SUCCESS METRICS & USABILITY SCORES
# ============================================================================
add_heading('11. Success Metrics & Usability Scores', 1)

add_heading('Maze Usability Score', 2)
doc.add_paragraph(
    'The Maze Usability Score is a composite metric (0-100) calculated from three factors:'
)
add_table_from_data(
    ['Factor', 'Weight', 'Description'],
    [
        ['Success Rate', '+1 point per % direct success', 'Testers who completed via expected path'],
        ['Misclick Penalty', '-0.5 points per % misclick', 'Clicks outside defined hotspots'],
        ['Duration Penalty', 'Variable', 'Time vs. expected benchmark'],
    ]
)

add_heading('Target Metrics', 2)
add_table_from_data(
    ['Metric', 'Target', 'Source', 'Priority'],
    [
        ['Usability Score (overall)', '≥75', 'Maze prototype tests', 'P0'],
        ['Direct Success Rate', '≥70%', 'Maze missions', 'P0'],
        ['Indirect Success Rate', '<20%', 'Maze missions', 'P1'],
        ['Bounce Rate', '<10%', 'Maze missions', 'P0'],
        ['Misclick Rate', '<15%', 'Maze heatmaps', 'P1'],
        ['Opinion Scale Average', '≥7.0/10', 'Maze question blocks', 'P0'],
        ['NPS Score', '≥30', 'End-to-end maze', 'P0'],
        ['Coach Trust Score', '≥7.0/10 (scaled 4/5)', 'Coach evaluation maze', 'P0'],
        ['Synthesis Actionability', '≥7.0/10 (scaled 4/5)', 'Synthesis maze + expert review', 'P0'],
    ]
)

add_heading('Risk Thresholds', 2)
add_table_from_data(
    ['Metric', 'Acceptable', 'Warning', 'Critical'],
    [
        ['Usability Score', '≥75', '60-74', '<60'],
        ['Direct Success', '≥70%', '50-69%', '<50%'],
        ['Bounce Rate', '<10%', '10-25%', '>25%'],
        ['Misclick Rate', '<15%', '15-25%', '>25%'],
        ['Opinion Scale Avg', '≥7.0', '5.0-6.9', '<5.0'],
        ['NPS', '≥30', '0-29', '<0'],
    ]
)

add_heading('Qualitative Success Criteria', 2)
add_table_from_data(
    ['Criterion', 'Evidence Required'],
    [
        ['Clear value proposition', 'Testers articulate value in open questions'],
        ['Intuitive navigation', '<2 "lost" moments per moderated session'],
        ['Coach adds value', 'Testers voluntarily engage with coach in free explore'],
        ['Outputs are actionable', 'Testers describe specific next steps'],
        ['Trust is established', 'Testers indicate willingness to share outputs externally'],
    ]
)

doc.add_page_break()

# ============================================================================
# 12. BUDGET & RESOURCES
# ============================================================================
add_heading('12. Budget & Resources', 1)

add_heading('Maze Platform Costs', 2)
add_table_from_data(
    ['Item', 'Quantity', 'Unit Cost', 'Total'],
    [
        ['Maze subscription (Team plan)', '1 month', '$150', '$150'],
        ['Panel testers (unmoderated)', '115 testers', '$3-5/tester', '$345-575'],
        ['Panel testers (moderated)', '8 testers', '$20-30/tester', '$160-240'],
        ['Subtotal (Platform + Panel)', '', '', '$655-965'],
    ]
)

add_heading('Tester Incentives', 2)
add_table_from_data(
    ['Test Type', 'Testers', 'Incentive', 'Total'],
    [
        ['Unmoderated Mazes (5-8 min)', '60', '$20', '$1,200'],
        ['Unmoderated Mazes (8-12 min)', '55', '$35', '$1,925'],
        ['Moderated Sessions (60 min)', '8', '$175', '$1,400'],
        ['Subtotal (Incentives)', '', '', '$4,525'],
    ]
)

add_heading('Expert Review Costs', 2)
add_table_from_data(
    ['Item', 'Quantity', 'Unit Cost', 'Total'],
    [
        ['Expert reviewer fee', '5 experts × 3 reviews', '$125/review', '$1,875'],
        ['Subtotal (Expert Review)', '', '', '$1,875'],
    ]
)

add_heading('Total Budget Summary', 2)
add_table_from_data(
    ['Category', 'Low Estimate', 'High Estimate'],
    [
        ['Maze Platform + Panel', '$655', '$965'],
        ['Tester Incentives', '$4,525', '$4,525'],
        ['Expert Review', '$1,875', '$1,875'],
        ['Contingency (15%)', '$1,058', '$1,105'],
        ['TOTAL', '$8,113', '$8,470'],
    ]
)

add_heading('Internal Resources', 2)
add_table_from_data(
    ['Role', 'Hours', 'Activities'],
    [
        ['UX Researcher', '50-60', 'Build mazes, moderate sessions, analyse data, report'],
        ['Product Manager', '15-20', 'Define objectives, review findings, prioritise backlog'],
        ['Designer', '10-15', 'Prepare prototypes, refine based on findings'],
    ]
)

doc.add_page_break()

# ============================================================================
# 13. APPENDICES
# ============================================================================
add_heading('13. Appendices', 1)

add_heading('Appendix A: Maze Terminology Glossary', 2)
add_table_from_data(
    ['Term', 'Definition'],
    [
        ['Maze', 'A complete test containing multiple blocks'],
        ['Block', 'Individual test component (prototype test, question, etc.)'],
        ['Mission', 'A specific task for testers to complete in a prototype'],
        ['Expected Path', 'The ideal route through screens to complete a mission'],
        ['Goal Screen', 'The target destination screen for a mission'],
        ['Hotspot', 'Defined clickable area on a screen'],
        ['Direct Success', 'Completed mission via expected path'],
        ['Indirect Success', 'Completed mission via unexpected path'],
        ['Bounce/Give Up', 'Tester abandoned the mission'],
        ['Usability Score', 'Composite score (0-100) based on success, duration, misclicks'],
        ['Misclick', 'Click outside defined hotspots'],
        ['Heatmap', 'Aggregated visualisation of tester clicks'],
        ['Click Path', 'Sequence of screens visited by a tester'],
        ['Tester', 'Participant taking the maze (Maze terminology)'],
        ['Panel', 'Maze\'s pool of pre-screened testers'],
        ['Opinion Scale', 'Rating question (0-10 scale)'],
        ['5-Second Test', 'Brief exposure to capture first impressions'],
    ]
)

add_heading('Appendix B: Consent Template', 2)
consent = doc.add_paragraph()
consent.add_run('PARTICIPANT CONSENT\n\n').bold = True
consent.add_run(
    'Study: Product Strategy Coach User Research\n'
    'Organisation: Frontera\n\n'
    'Purpose:\n'
    'We are conducting research to evaluate and improve our product strategy coaching tool. '
    'Your participation will help us understand how users interact with the tool.\n\n'
    'What to Expect:\n'
    '• Duration: Approximately [X] minutes\n'
    '• Activities: Using a prototype, answering questions, sharing feedback\n'
    '• Recording: Sessions may be recorded for research purposes\n\n'
    'Your Rights:\n'
    '• Participation is voluntary\n'
    '• You may withdraw at any time\n'
    '• Your responses will be kept confidential\n\n'
    'By proceeding, you confirm that you understand and agree to participate.'
)

add_heading('Appendix C: Moderated Session Checklist', 2)

doc.add_paragraph('Pre-Session', style='Intense Quote')
items = [
    'Prototype URL confirmed working',
    'Maze recording enabled',
    'Tester background reviewed',
    'Observation template ready',
    'Backup scenarios prepared',
]
for item in items:
    doc.add_paragraph(f'☐ {item}')

doc.add_paragraph('During Session', style='Intense Quote')
items = [
    'Welcome and introductions',
    'Confirm consent to record',
    'Start Maze recording',
    'Context-setting questions',
    'First impressions observation',
    'Discovery phase walkthrough',
    'Research phase walkthrough',
    'Synthesis review',
    'Debrief interview',
    'Thank tester',
    'Stop recording',
]
for item in items:
    doc.add_paragraph(f'☐ {item}')

doc.add_paragraph('Post-Session', style='Intense Quote')
items = [
    'Complete observation notes within 24 hours',
    'Tag key moments in Maze',
    'Update tester tracker',
    'Send thank-you message',
    'Process incentive',
    'Share preliminary insights',
]
for item in items:
    doc.add_paragraph(f'☐ {item}')

# Document control
doc.add_page_break()
add_heading('Document Control', 1)
add_table_from_data(
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ['1.0', 'February 2026', 'Frontera Product Team', 'Initial draft'],
        ['2.0', 'February 2026', 'Frontera Product Team', 'Updated with Maze terminology and best practices'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('— End of Document —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
output_path = r'c:\Users\deeks\frontera-platform\Background\Product_Strategy_Coach_User_Testing_Proposal_v2.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
